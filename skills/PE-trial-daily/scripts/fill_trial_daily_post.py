# -*- coding: utf-8 -*-
"""
fill_trial_daily_post.py v3
===========================

「体育试讲设计每日一练」帖子生成脚本（v3：教学质量优先的门禁版）。

流程（成功提示只在全部门禁通过后出现）：
  1. 读取 pending_trial_daily.json（draft@3：内容与渲染共用一份受检 JSON）
  2. 视图门：source_view_entry 可生成、无误收选题 blocker
  3. 评分门：ptd_core.score_draft（事实锁定 + 100 分量表 + v3 硬门）→ release
  4. 评审门：review_trial_daily.json 与当前草稿同版本（draft_sha），五项内容评审全过
  5. 构建：品牌化 DOCX（语义分段逐字稿，无固定字符截断）
  6. 格式验证 + 渲染检查（soffice 转 PDF 逐页检查）
  7. 快照 → 原子提交 → 删 JSON

输入：scripts/pending_trial_daily.json（draft@3）
评审：scripts/review_trial_daily.json（review@1，由内容评审产出，须与草稿同版本）
输出：desktop-attachments/2 体育试讲每日一练-帖子内容编辑模板.docx（原地覆盖）
"""

import json
import os
import re
import shutil
import struct
import sys
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
sys.path.insert(0, str(SCRIPT_DIR))
import ptd_core as core  # noqa: E402

def _resolve_workspace() -> Path:
    """工作区优先级：--workspace 参数 > TRIAL_DAILY_WORKSPACE > 技能目录。"""
    if "--workspace" in sys.argv:
        return Path(sys.argv[sys.argv.index("--workspace") + 1]).expanduser().resolve()
    env_ws = os.environ.get("TRIAL_DAILY_WORKSPACE")
    if env_ws:
        return Path(env_ws).expanduser().resolve()
    return SKILL_DIR


WORKSPACE = _resolve_workspace()
PROJECT_SCRIPT_DIR = WORKSPACE / "scripts"
SOURCE_TEMPLATE = WORKSPACE / "模板文件" / "2 体育试讲每日一练-帖子内容编辑模板.docx"
TEMPLATE_PATH = WORKSPACE / "desktop-attachments" / "2 体育试讲每日一练-帖子内容编辑模板.docx"
PENDING_JSON = PROJECT_SCRIPT_DIR / "pending_trial_daily.json"
REVIEW_JSON = PROJECT_SCRIPT_DIR / "review_trial_daily.json"
SNAPSHOT_DIR = PROJECT_SCRIPT_DIR / "_snapshots_trial"
MAX_SNAPSHOTS = 10

COVER_TITLE = "体育试讲设计每日一练"
DRAIN_TEXT = "关注我，每天一个体育试讲设计，帮你备考上岸"
COVER_BG = SCRIPT_DIR / "cover_bg.png"   # 封面整页底层背景图（从用户模板提取）
WATERMARK_TEXT = "世豪老师"              # 页眉水印文字（与用户模板一致）
NAVY = RGBColor(0x0B, 0x32, 0x89)
CYAN = RGBColor(0x9F, 0xD8, 0xE8)
DARK = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)

# ---- 3:4 手机版版式契约（可被 config.default.json 覆盖） ----
CFG_PAGE = {"size_cm": [15.0, 20.0], "ratio": [3, 4]}
CFG_FONT = ["Hiragino Sans GB", "Heiti SC", "Songti SC", "PingFang SC", "Microsoft YaHei"]
BODY_PT = 18          # 正文/表格 ≥18
SECTION_PT = 26       # 栏目标题 24–28
LABEL_PT = 16         # 图注/标签 ≥16
CTA_PT = 18           # CTA ≥18
TABLE_COLS_PCT = [42, 58]
WATERMARK_OPACITY = 0.10   # 8%–12%
PARA_MAX_CHARS = 30   # 语义分段：整句成段的近似上限（小块少浪费页尾，绝不截断句子）


def load_page_config():
    """从 config.default.json 读取页面尺寸与字体契约（不存在时用默认值）。"""
    cfg_path = SKILL_DIR / "config.default.json"
    page = dict(CFG_PAGE)
    fonts = list(CFG_FONT)
    if cfg_path.exists():
        try:
            cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
            pg = cfg.get("page") or {}
            if pg.get("size_cm"):
                page["size_cm"] = pg["size_cm"]
            if pg.get("ratio"):
                page["ratio"] = pg["ratio"]
            if pg.get("font_contract"):
                fonts = pg["font_contract"]
        except Exception:
            pass
    return page, fonts


def _first_available_font(fonts):
    """取本机可渲染的契约字体（fc-match 不可用时回退首个契约字体）。"""
    import subprocess

    for f in fonts:
        try:
            r = subprocess.run(["fc-match", f], capture_output=True, text=True)
        except FileNotFoundError:
            break
        if r.returncode == 0 and r.stdout.strip() and "not found" not in r.stdout.lower():
            return f
    return fonts[0]


PAGE_SIZE, FONT_CONTRACT = load_page_config()
FONT = _first_available_font(FONT_CONTRACT)

SEGMENT_TYPES = {
    "game": "游戏 / 比赛环节",
    "practice": "练习环节",
    "fitness": "体能游戏环节",
}


def take_snapshot():
    SNAPSHOT_DIR.mkdir(exist_ok=True)
    if not TEMPLATE_PATH.exists():
        return None
    ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    snap = SNAPSHOT_DIR / f"snapshot_{ts}.docx"
    shutil.copy(TEMPLATE_PATH, snap)
    snaps = sorted(SNAPSHOT_DIR.glob("snapshot_*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    for old in snaps[MAX_SNAPSHOTS:]:
        old.unlink()
    return snap


# ---------------------------------------------------------------------------
# 读取与门禁
# ---------------------------------------------------------------------------


def _require_text(value, label):
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{label} 必须是非空字符串")
    return value


def _check_banned(value, label):
    for ch in ("：", ":", "——"):
        if ch in value:
            raise ValueError(f"{label} 不得含「{ch}」（去 AI 味铁律）")


def load_draft():
    """读取并结构校验 draft@3。"""
    with open(PENDING_JSON, encoding="utf-8") as f:
        draft = json.load(f)
    if not isinstance(draft, dict):
        raise ValueError("pending_trial_daily.json 顶层必须是对象")
    if draft.get("schema") != core.SCHEMA_DRAFT:
        raise ValueError(f"schema 须为 {core.SCHEMA_DRAFT}（收到 {draft.get('schema')}）")
    _require_text(draft.get("id"), "id")

    seg = draft.get("segment") or {}
    if seg.get("type") not in SEGMENT_TYPES:
        raise ValueError(f"segment.type 必须是 {list(SEGMENT_TYPES)} 之一")
    meta = seg.get("meta") or {}
    for k in core.SEGMENT_FIELDS:
        _require_text(meta.get(k), f"segment.meta.{k}")
    _check_banned(meta.get("时长", ""), "segment.meta.时长")

    render = draft.get("render") or {}
    for k in ("sport", "chapter", "segment_name"):
        _require_text(render.get(k), f"render.{k}")
    _require_text(render.get("difficulty_display"), "render.difficulty_display")
    if not isinstance(render.get("figure"), str):
        raise ValueError("render.figure 必须是字符串")
    if not isinstance(render.get("figure_images"), list):
        raise ValueError("render.figure_images 必须是列表")
    for p in render["figure_images"]:
        if not os.path.isfile(p):
            raise ValueError(f"图例图片不存在：{p}")
    if (render.get("cta") or "").strip() != DRAIN_TEXT:
        raise ValueError("render.cta 与固定引流段不一致（不许改引流文案）")
    _require_text(render.get("hashtags"), "render.hashtags")

    fields = draft.get("fields") or {}
    for k in ("method", "rules", "intent", "organization"):
        blk = fields.get(k)
        if not isinstance(blk, dict) or not (blk.get("text") or "").strip():
            raise ValueError(f"fields.{k} 缺失或 text 为空")
        _check_banned(blk["text"], f"fields.{k}.text")
    rows = (fields.get("errors") or {}).get("rows") or []
    if seg["type"] == "practice" and not rows:
        raise ValueError("practice 环节必须提供易犯错误与纠正(errors.rows)")
    for i, row in enumerate(rows):
        for side in ("error", "fix"):
            blk = row.get(side)
            if not isinstance(blk, dict) or not (blk.get("text") or "").strip():
                raise ValueError(f"errors.rows[{i}].{side} 缺失或 text 为空")

    flow = draft.get("flow") or []
    if len(flow) < 3:
        raise ValueError("flow 至少 3 个教学阶段")
    for i, st in enumerate(flow):
        _require_text(st.get("stage"), f"flow[{i}].stage")
        _require_text(st.get("script"), f"flow[{i}].script")
        _check_banned(st["script"], f"flow[{i}].script")
        for k in ("demo_sec", "pause_sec"):
            v = st.get(k, 0)
            if not isinstance(v, (int, float)) or v < 0:
                raise ValueError(f"flow[{i}].{k} 必须是非负数字（示范/停顿秒数显式登记）")

    view_entry = draft.get("source_view_entry")
    if not isinstance(view_entry, dict):
        raise ValueError("缺 source_view_entry（先跑 build_generatable_view.py 取本条视图记录）")
    if view_entry.get("id") != draft.get("id"):
        raise ValueError("source_view_entry.id 与草稿 id 不一致")
    return draft


def gate_view(view_entry: dict):
    """视图门：可生成、无误收选题。"""
    if view_entry.get("generatable_blockers"):
        raise RuntimeError(
            f"视图门未过：blockers={view_entry['generatable_blockers']}（本选题不可生成）"
        )
    if not view_entry.get("generatable", True):
        raise RuntimeError("视图门未过：generatable=false")


def gate_score(draft: dict, view_entry: dict, lib) -> dict:
    """评分门：事实锁定 + 100 分量表 + v3 硬门。"""
    result = core.score_draft(draft, view_entry, lib)
    print(f"      评分：总分 {result['total']}（{result['scores']}）")
    if result["factlock"]["violations"]:
        for v in result["factlock"]["violations"][:6]:
            print(f"      · factlock {v['where']} {v['type']} {v.get('token') or ''}")
    for det_k, det_vs in result["detail"].items():
        for dv in det_vs:
            print(f"      · [{det_k}] {dv}")
    if result["hard_gates"]:
        print(f"      硬门：{result['hard_gates']}")
    if not result["release"]:
        raise RuntimeError(
            f"评分门未过：总分 {result['total'] / 100:.2f} 不足或硬门非零 "
            f"{result['hard_gates']}（问题稿不得生成正式成品）"
        )
    return result


def gate_review(draft: dict):
    """评审门：内容评审记录与当前草稿同版本、五项内容评审全过。"""
    if not REVIEW_JSON.exists():
        raise RuntimeError(
            f"评审门未过：缺少内容评审记录 {REVIEW_JSON.name}（内容评审实责，无记录不放行）"
        )
    review = json.loads(REVIEW_JSON.read_text(encoding="utf-8"))
    ok, errors = core.validate_review(review, draft)
    if not ok:
        raise RuntimeError("评审门未过：" + "；".join(errors))
    return review


# ---------------------------------------------------------------------------
# 底层样式工具
# ---------------------------------------------------------------------------


def _set_font(run, size=14, bold=False, color=DARK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top="200", left="200", bottom="200", right="200"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = tcMar.find(qn("w:" + name))
        if el is None:
            el = OxmlElement("w:" + name)
            tcMar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")


def _cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def _add_para(cell, text, size=14, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_after=4, space_before=0, line=None):
    p = cell if isinstance(cell, str) else cell.add_paragraph()
    if isinstance(cell, str):
        p = cell
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if line:
        pf.line_spacing = line
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def _body_para(doc, text, size=BODY_PT, bold=False, color=DARK, space_after=4, line=1.15,
               align=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pf.widow_control = False  # 长段允许跨页续排，减少页底空白
    if keep_with_next:
        pf.keep_with_next = True  # False 时不写属性，避免 keepNext val=0 干扰
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def _add_watermark(header, text=WATERMARK_TEXT):
    """在页眉追加 VML 斜向水印（灰 #C0C0C0，透明度 8%–12%）。

    要点：不用 DrawingML（LibreOffice 会把整页锚定形状翻页成多页空白）；
    VML 必须去掉 mso-position-*-relative 等相对定位（否则 LibreOffice 会把页眉撑高 ~4cm）。
    """
    op = hex(int(WATERMARK_OPACITY * 0xFFFFFF))[2:].zfill(6)  # 24bit 分数
    wm = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:v="urn:schemas-microsoft-com:vml">'
        '<w:pPr><w:spacing w:line="2" w:lineRule="exact"/></w:pPr>'
        '<w:r><w:rPr><w:sz w:val="2"/></w:rPr><w:pict>'
        '<v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049" o:spt="136" '
        'type="#_x0000_t136" '
        'style="position:absolute;left:12pt;top:150pt;height:30pt;width:396pt;'
        'rotation:-2949120f;z-index:-251657216;" fillcolor="#C0C0C0" filled="t" '
        'stroked="f" coordsize="21600,21600" adj="10800">'
        '<v:path/><v:fill on="t" opacity="' + op + '" focussize="0,0"/>'
        '<v:stroke on="f"/><v:imagedata o:title=""/>'
        '<o:lock v:ext="edit" aspectratio="t"/>'
        '<v:textpath on="t" fitshape="t" fitpath="t" trim="t" xscale="f" '
        'string="__WM__" style="font-family:Hiragino Sans GB;font-size:36pt;'
        'v-same-letter-heights:f;v-text-align:center;"/>'
        '</v:shape></w:pict></w:r></w:p>'
    )
    header._element.append(parse_xml(wm.replace("__WM__", text)))


def _anchor_cover_bg(para, image_path):
    """插入整页底层背景图，按页面精确宽高显示且不裁切品牌区。

    图片左上角锚定页边，layoutInCell=0 使锚点不受封面表格单元格裁切。
    禁止放大出血，否则底部标语与右侧 SHTr 会被页面边界裁掉。
    """
    w_cm, h_cm = PAGE_SIZE["size_cm"]
    page_w = int(w_cm * 360000)   # EMU
    page_h = int(h_cm * 360000)
    run = para.add_run()
    run.add_picture(str(image_path), width=Emu(page_w), height=Emu(page_h))
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    anchor = OxmlElement("wp:anchor")
    for attr, val in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("simplePos", "0"), ("relativeHeight", "251660288"), ("behindDoc", "1"),
        ("locked", "1"), ("layoutInCell", "0"), ("allowOverlap", "1"),
    ):
        anchor.set(attr, val)
    simple_pos = OxmlElement("wp:simplePos"); simple_pos.set("x", "0"); simple_pos.set("y", "0")
    pos_h = OxmlElement("wp:positionH"); pos_h.set("relativeFrom", "page")
    off_x = OxmlElement("wp:posOffset"); off_x.text = "0"; pos_h.append(off_x)
    pos_v = OxmlElement("wp:positionV"); pos_v.set("relativeFrom", "page")
    off_y = OxmlElement("wp:posOffset"); off_y.text = "0"; pos_v.append(off_y)
    wrap_none = OxmlElement("wp:wrapNone")
    kept = {}
    for child in list(inline):
        tag = child.tag
        for key in ("wp:extent", "wp:effectExtent", "wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
            if tag == qn(key):
                kept[key] = child
    for key in ("wp:extent", "wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        if key not in kept:
            raise RuntimeError(f"封面背景图缺少 {key}")
    if "wp:effectExtent" not in kept:
        kept["wp:effectExtent"] = OxmlElement("wp:effectExtent")
    for child in (simple_pos, pos_h, pos_v, kept["wp:extent"], kept["wp:effectExtent"],
                  wrap_none, kept["wp:docPr"], kept["wp:cNvGraphicFramePr"], kept["a:graphic"]):
        anchor.append(child)
    drawing.remove(inline)
    drawing.append(anchor)
    return anchor


def _section_title(doc, text, color=NAVY, page_break_before=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(4)
    pf.page_break_before = page_break_before
    pf.keep_with_next = True  # 标题与正文相随，不孤立在页尾
    run = p.add_run(text)
    _set_font(run, size=SECTION_PT, bold=True, color=color)
    # 底部加粗下边框
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "0B3289")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# 构建
# ---------------------------------------------------------------------------


def _apply_page_size(section):
    """把节页面设为精确 3:4 手机版（15cm × 20cm，可配置）。"""
    w_cm, h_cm = PAGE_SIZE["size_cm"]
    section.page_width = Cm(w_cm)
    section.page_height = Cm(h_cm)


def build_cover(doc, draft):
    """封面：整页深蓝底 + 浅青标题 + 项目标签 + 环节名 + 难度（3:4 单页）。"""
    render = draft["render"]
    section = doc.sections[0]
    _apply_page_size(section)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    # 封面页不显示页眉页脚品牌文字，同时叠加斜向水印
    try:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for hp in section.header.paragraphs:
            for r in list(hp.runs):
                r._element.getparent().remove(r._element)
        for fp in section.footer.paragraphs:
            for r in list(fp.runs):
                r._element.getparent().remove(r._element)
    except Exception:
        pass
    _add_watermark(section.header)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 封面表格浮动锚定整页（与用户模板一致），保证放大字号后仍恰好一页
    tblPr = table._tbl.tblPr
    for tag in ("w:tblW", "w:jc", "w:tblInd", "w:tblLayout", "w:tblCellMar", "w:tblLook"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    tblpPr = OxmlElement("w:tblpPr")
    for attr, val in (("leftFromText", "0"), ("rightFromText", "0"), ("vertAnchor", "page"),
                      ("horzAnchor", "page"), ("tblpX", "0"), ("tblpY", "1")):
        tblpPr.set(qn("w:" + attr), val)
    tblPr.append(tblpPr)
    overlap = OxmlElement("w:tblOverlap"); overlap.set(qn("w:val"), "never"); tblPr.append(overlap)
    PAGE_W_DXA = int(PAGE_SIZE["size_cm"][0] / 2.54 * 1440)  # 3:4 页宽（15cm=8505）
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(PAGE_W_DXA)); tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "autofit"); tblPr.append(layout)
    cell_mar = OxmlElement("w:tblCellMar")
    for name, w in (("top", "0"), ("left", "108"), ("bottom", "0"), ("right", "108")):
        m = OxmlElement("w:" + name); m.set(qn("w:w"), w); m.set(qn("w:type"), "dxa"); cell_mar.append(m)
    tblPr.append(cell_mar)
    for gc in table._tbl.tblGrid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), str(PAGE_W_DXA))

    cell = table.cell(0, 0)
    # 底层背景图锚定封面单元格段落（behindDoc 铺满整页）；
    # 有底图时保持单元格透明，避免底色遮住底部标语；无图时才回退纯深蓝底。
    if COVER_BG.exists():
        _anchor_cover_bg(cell.paragraphs[0], COVER_BG)
    else:
        _shade_cell(cell, "0B3289")
    _set_cell_margins(cell, top="200", left="108", bottom="200", right="108")
    _cell_vertical_center(cell)

    # 顶部：固定小标题
    _add_para(cell, "体师备考知识库 · 教师编面试", size=18, bold=True, color=CYAN,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=16)
    # 大标题（48pt 九字超 15cm 行宽，按品牌拆两行）
    _add_para(cell, "体育试讲设计", size=48, bold=True, color=CYAN,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _add_para(cell, "每日一练", size=48, bold=True, color=CYAN,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    # 项目标签
    _add_para(cell, f"【{render['sport']}】{SEGMENT_TYPES[draft['segment']['type']]}",
              size=22, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    # 环节名
    _add_para(cell, render["segment_name"], size=30, bold=True, color=WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    # 难度（诚实标注：教材未标难度时明示，不出现编造星级）
    _add_para(cell, f"难度：{render['difficulty_display']}", size=18, bold=True,
              color=RGBColor(0xF0, 0xC8, 0x6C), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 行高撑满大部分页面，底部留出底层背景图的品牌信息条（与模板一致：15290 atLeast）
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trh = OxmlElement("w:trHeight")
    trh.set(qn("w:val"), "10400")
    trh.set(qn("w:hRule"), "atLeast")
    trPr.append(trh)


def split_script_paragraphs(script, max_chars=PARA_MAX_CHARS):
    """语义分段：整句成段，相邻整句凑段，绝不按固定字符数截断句子。

    孤立标点片段（旧稿"独立成段的句号"病灶）直接丢弃。
    """
    parts = re.findall(r"[^。！？]+[。！？]?", script)
    paras: list[str] = []
    buf = ""
    for raw in parts:
        s = raw.strip()
        if not s:
            continue
        if re.fullmatch(r"[。！？，、；：:]+", s):
            continue  # 孤立标点不渲染
        if buf and len(buf) + len(s) > max_chars:
            paras.append(buf)
            buf = ""
        buf += s
    if buf:
        paras.append(buf)
    return paras


def _stage_sec(st: dict, rate: int) -> int:
    speech = len(st.get("script", "")) / rate * 60.0
    return round(speech + float(st.get("demo_sec") or 0) + float(st.get("pause_sec") or 0))


def _png_dims(path):
    """读 PNG 宽高（纯 stdlib）。"""
    with open(path, "rb") as f:
        head = f.read(24)
    if head[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError(f"非 PNG: {path}")
    w, h = struct.unpack(">II", head[16:24])
    return w, h


def _figure_fit_cm(path, avail_w_cm, avail_h_cm, pref_w_cm=None):
    """按图像宽高比在可用区内等比适配，保证标题+图注+图同页不溢出。"""
    w_px, h_px = _png_dims(path)
    if pref_w_cm is None:
        pref_w_cm = avail_w_cm  # 图例按可用区满宽放大，兼顾占高
    fw = min(pref_w_cm, avail_w_cm)
    fh = fw * h_px / w_px
    if fh <= avail_h_cm:
        return fw, fh
    fh = avail_h_cm
    fw = fh * w_px / h_px
    return fw, fh


def _fixed_table(doc, rows, cols, width_cm, col_pct):
    """固定总宽、42/58 列、禁 autofit、行不跨页 的错误纠正表。"""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    for child in list(tblPr):
        tblPr.remove(child)
    total_dxa = int(width_cm / 2.54 * 1440)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(total_dxa)); tblW.set(qn("w:type"), "dxa")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "8EAADB"); el.set(qn("w:space"), "0")
        borders.append(el)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    cell_mar = OxmlElement("w:tblCellMar")
    for name, w in (("top", "40"), ("left", "108"), ("bottom", "40"), ("right", "108")):
        m = OxmlElement("w:" + name); m.set(qn("w:w"), w); m.set(qn("w:type"), "dxa")
        cell_mar.append(m)
    for el in (tblW, jc, borders, layout, cell_mar):
        tblPr.append(el)
    # 固定列宽 42/58
    grid = tbl._tbl.tblGrid
    for gc in grid.findall(qn("w:gridCol")):
        grid.remove(gc)
    for pct in col_pct:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(total_dxa * pct / 100)))
        grid.append(gc)
    for ri, row in enumerate(tbl.rows):
        trPr = row._tr.get_or_add_trPr()
        trh = OxmlElement("w:trHeight"); trh.set(qn("w:val"), "0"); trh.set(qn("w:hRule"), "atLeast")
        trPr.append(trh)
        cant = OxmlElement("w:cantSplit")  # 行不跨页
        trPr.append(cant)
        if ri == 0:
            trPr.append(OxmlElement("w:tblHeader"))
        for ci, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(int(total_dxa * col_pct[ci] / 100)))
            vAlign = OxmlElement("w:vAlign"); vAlign.set(qn("w:val"), "center"); tcPr.append(vAlign)
    return tbl


def _fill_table_cell(cell, text, header=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        p.runs[0].text = str(text)
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(str(text))
    run = p.runs[0]
    if header:
        _set_font(run, size=BODY_PT, bold=True, color=WHITE)
        _shade_cell(cell, "0B3289")
    else:
        _set_font(run, size=BODY_PT, bold=False, color=DARK)


def build_content(doc, draft):
    """正文：页眉 + 图例 + 环节拆解（含片段定位）+ 教学要点 + 纠错表 + 分段逐字稿 + 引流。"""
    render = draft["render"]
    meta = draft["segment"]["meta"]
    rate = int((draft.get("config") or {}).get("speech_rate_chars_per_min", 230))

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_page_size(section)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    usable_w_cm = PAGE_SIZE["size_cm"][0] - 2 * 1.0   # 13.0cm

    # 页眉（必须先断开与前节链接，否则品牌文字会写进封面节）
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("体师备考知识库")
    _set_font(hr, size=10, bold=True, color=NAVY)
    _add_watermark(header)

    # 页脚：禁用页码（产品要求不显示页码，清空默认页脚文本）
    footer = section.footer
    footer.is_linked_to_previous = False
    for fpar in footer.paragraphs:
        for r in list(fpar.runs):
            r._element.getparent().remove(r._element)

    # ===== 图例区 =====
    if render["figure_images"]:
        _section_title(doc, "图例直观")
        # 图可占高（保守，避免 LibreOffice 把高图推下页）：页高-上下边距-页眉带-标题-图注-留白
        avail_h_cm = 11.0
        for img in render["figure_images"]:
            fw, fh = _figure_fit_cm(img, usable_w_cm, avail_h_cm)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img, width=Cm(fw), height=Cm(fh))
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
        if render["figure"]:
            _body_para(doc, render["figure"], size=LABEL_PT, bold=True, color=GREEN,
                       space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ===== 环节拆解（含片段定位：适用学段/嵌入位置/时长） =====
    # 图例放大铺满首屏后，环节拆解标题另起一页置于最上端；
    # 用 page_break_before 把页断挂在标题上，避免独立 add_page_break 段落产生空白页
    _section_title(doc, "环节拆解", page_break_before=bool(render["figure_images"]))

    def _field(label, value):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.15
        pf.widow_control = False
        r1 = p.add_run(f"{label}｜")
        _set_font(r1, size=LABEL_PT, bold=True, color=NAVY)
        r2 = p.add_run(value)
        _set_font(r2, size=BODY_PT, bold=False, color=DARK)

    _field("环节名称", render["segment_name"])
    _field("难度", render["difficulty_display"])
    _field("活动类型", SEGMENT_TYPES[draft["segment"]["type"]])
    _field("适用学段", meta["学段"])
    _field("片段位置", meta["片段位置"])
    _field("时长", meta["时长"])
    _field("活动方法", (draft["fields"]["method"])["text"])
    _field("规则", (draft["fields"]["rules"])["text"])
    _field("活动设计意图", (draft["fields"]["intent"])["text"])
    _field("活动组织形式", (draft["fields"]["organization"])["text"])

    # ===== 教学要点（怎么教 + 安全安排） =====
    _section_title(doc, "教学要点")
    _field("本期重点", meta["重点"])
    _field("器材与场地", meta["器材"])
    _field("安全安排", meta["安全"])
    _field("分层", meta["分层"])
    _field("评价与观察", meta["评价"])

    # 易犯错误与纠正（仅 practice）
    rows = (draft["fields"].get("errors") or {}).get("rows") or []
    if rows:
        _section_title(doc, "易犯错误与纠正")
        tbl = _fixed_table(doc, rows=len(rows) + 1, cols=2,
                           width_cm=usable_w_cm, col_pct=TABLE_COLS_PCT)
        _fill_table_cell(tbl.cell(0, 0), "易犯错误", header=True)
        _fill_table_cell(tbl.cell(0, 1), "纠正方法", header=True)
        for i, row in enumerate(rows, 1):
            _fill_table_cell(tbl.cell(i, 0), row["error"]["text"], header=False)
            _fill_table_cell(tbl.cell(i, 1), row["fix"]["text"], header=False)

    # ===== 试讲逐字稿（按教学阶段语义分段；阶段标注含口播+示范停顿秒数） =====
    _section_title(doc, "试讲逐字稿")
    seen_paras: set[str] = set()
    stage_blocks = []
    for st in draft["flow"]:
        label = f"{st['stage']}（约{_stage_sec(st, rate)}秒）"
        lp = _body_para(doc, label, size=LABEL_PT, bold=True, color=GREEN,
                        space_after=3, keep_with_next=True)
        stage_blocks.append(lp)
        for para in split_script_paragraphs(st["script"]):
            if para in seen_paras:
                continue  # 内容去重：重复整段不重复渲染
            seen_paras.add(para)
            stage_blocks.append(
                _body_para(doc, para, size=BODY_PT, color=DARK, space_after=2, line=1.15)
            )
    # 引流与结尾成组：最后一段逐字稿与引流链成组（整段不拆行）。
    # 链条只保留最小组，过长的 keep 链 LibreOffice 会整段放弃，反而产生近空页
    if stage_blocks:
        last = stage_blocks[-1]
        last.paragraph_format.keep_with_next = True
        last.paragraph_format.keep_together = True

    # ===== 引流（紧接逐字稿结尾，空一行直接写，不另起页） =====
    hs = _body_para(doc, render["hashtags"], size=LABEL_PT, bold=True,
                    color=RGBColor(0x80, 0x80, 0x80),
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    hs.paragraph_format.keep_with_next = True
    hs.paragraph_format.keep_together = True  # 标签行整段不拆
    hs.paragraph_format.widow_control = True
    cta_p = _body_para(doc, render["cta"], size=CTA_PT, bold=True, color=NAVY,
                       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    cta_p.paragraph_format.keep_together = True  # 引流段整段不拆，杜绝跨页劈字
    cta_p.paragraph_format.widow_control = True


def build_doc(draft):
    doc = Document()
    # 默认样式字体（3:4 版式契约）
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    # 默认页面尺寸（封面节）
    _apply_page_size(doc.sections[0])

    build_cover(doc, draft)
    build_content(doc, draft)
    return doc


# ---------------------------------------------------------------------------
# 验证
# ---------------------------------------------------------------------------


def validate_output(doc, draft):
    errors = []
    render = draft["render"]
    meta = draft["segment"]["meta"]
    # 3:4 页面（所有节）
    for si, sec in enumerate(doc.sections):
        w, h = sec.page_width, sec.page_height
        if w and h:
            ratio = w / h
            if abs(ratio - 0.75) / 0.75 > 0.01:
                errors.append(f"第 {si + 1} 节页面非 3:4（{ratio:.4f}）")
    all_text = []
    for p in doc.paragraphs:
        all_text.append(p.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text or "")
    text = "\n".join(all_text)
    # 内容命中：渲染字段 + 片段定位要素 + 逐字稿全文
    for label, value in [
        ("sport", render["sport"]), ("segment_name", render["segment_name"]),
        ("difficulty_display", render["difficulty_display"]),
        ("学段", meta["学段"]), ("片段位置", meta["片段位置"]), ("时长", meta["时长"]),
        ("重点", meta["重点"]), ("安全", meta["安全"]),
        ("method", draft["fields"]["method"]["text"]),
        ("rules", draft["fields"]["rules"]["text"]),
        ("intent", draft["fields"]["intent"]["text"]),
        ("organization", draft["fields"]["organization"]["text"]),
        ("hashtags", render["hashtags"]),
        ("cta", render["cta"]),
    ]:
        if value not in text:
            errors.append(f"关键词未命中：{label}")
    # 逐字稿按语义段落渲染，逐段校验存在性（内容不丢）
    for st in draft["flow"]:
        if f"{st['stage']}（约" not in text:
            errors.append(f"阶段标注缺失：{st['stage']}")
        for para in split_script_paragraphs(st["script"]):
            if para not in text:
                errors.append(f"逐字稿段未命中：{para[:14]}…")
    # 断句防护：不允许纯标点段
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and re.fullmatch(r"[。！？，、；：:]+", t):
            errors.append(f"出现孤立标点段：{t!r}")
    # 纠错表
    rows = (draft["fields"].get("errors") or {}).get("rows") or []
    if rows:
        if len(doc.tables) < 2:
            errors.append("practice 环节应有易犯错误表格")
        elif doc.tables[-1].cell(0, 0).text.strip() != "易犯错误":
            errors.append("易犯错误表头异常")
    # 图例（仅统计行内 figure，封面底层背景图为锚定图不计入）
    img_count = sum(len(p._element.findall(".//" + qn("wp:inline"))) for p in doc.paragraphs)
    if render["figure_images"] and img_count != len(render["figure_images"]):
        errors.append(f"图例图片数异常：{img_count}（期望 {len(render['figure_images'])}）")
    # 占位符残留
    leftovers = re.findall(r"[ＭＭ]?\{[A-Z_]+\}", text)
    if leftovers:
        errors.append(f"占位符残留：{leftovers}")
    if render["cta"] not in text:
        errors.append("固定引流段未命中")
    # 引流页样式：标签行居中+灰，引流段居中+深蓝（品牌铁律）
    body = [p for p in doc.paragraphs if p.text.strip()]
    if body:
        tag_p = body[-2] if body[-1].text.strip() == render["cta"] else None
        cta_p = body[-1] if body[-1].text.strip() == render["cta"] else None
        if cta_p is None or cta_p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            errors.append("固定引流段必须居中")
        elif tag_p is None or tag_p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            errors.append("#标签行必须居中")
        elif tag_p is not None and tag_p.runs and tag_p.runs[0].font.color.rgb != RGBColor(0x80, 0x80, 0x80):
            errors.append("#标签行颜色必须为灰色 #808080")
        elif cta_p is not None and cta_p.runs and cta_p.runs[0].font.color.rgb != NAVY:
            errors.append("固定引流段颜色必须为深蓝 #0B3289")
        # 引流两行不另起一页：hashtags 段之前紧邻的段落不得含分页符（应空一行直连逐字稿）
        hs_elem = tag_p._element
        prev = hs_elem.getprevious() if tag_p is not None else None
        if prev is not None and prev.tag == qn("w:p"):
            for br in prev.findall(".//" + qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    errors.append("引流两行不得另起一页，须紧接逐字稿后空一行")
                    break

    # 页眉斜向水印（品牌铁律：VML，透明度 8%–12%）
    for si, sec in enumerate(doc.sections):
        hx = sec.header._element.xml
        if "PowerPlusWaterMarkObject" not in hx:
            errors.append(f"第 {si + 1} 节页眉缺少水印")
        if "z-index:-" not in hx.replace(" ", ""):
            errors.append(f"第 {si + 1} 节水印未锚定到内容之下")

    # 封面整页底层背景图（behindDoc 锚定，精确贴页且不受单元格裁切）
    anchors = doc.element.body.findall(".//" + qn("wp:anchor"))
    cover_anchors = [a for a in anchors if a.get("behindDoc") == "1"]
    if not cover_anchors:
        errors.append("封面底层背景图缺失")
    else:
        anchor = cover_anchors[0]
        extent = anchor.find(qn("wp:extent"))
        pos_h = anchor.find(qn("wp:positionH"))
        pos_v = anchor.find(qn("wp:positionV"))
        expected_w = int(PAGE_SIZE["size_cm"][0] * 360000)
        expected_h = int(PAGE_SIZE["size_cm"][1] * 360000)
        if anchor.get("layoutInCell") != "0":
            errors.append("封面底图不得受表格单元格裁切（layoutInCell 须为 0）")
        if extent is None or int(extent.get("cx") or 0) != expected_w or int(extent.get("cy") or 0) != expected_h:
            errors.append("封面底图须按页面精确宽高显示，禁止放大出血或裁切")
        if (pos_h is None or pos_h.get("relativeFrom") != "page" or
                pos_h.find(qn("wp:posOffset")) is None or
                pos_h.find(qn("wp:posOffset")).text != "0"):
            errors.append("封面底图水平位置须贴齐页面左边界")
        if (pos_v is None or pos_v.get("relativeFrom") != "page" or
                pos_v.find(qn("wp:posOffset")) is None or
                pos_v.find(qn("wp:posOffset")).text != "0"):
            errors.append("封面底图垂直位置须贴齐页面上边界")
        if COVER_BG.exists() and doc.tables:
            cover_shading = doc.tables[0].cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd"))
            if cover_shading is not None:
                errors.append("封面底图存在时单元格必须透明，避免遮住底部标语和 SHTr")

    # 封面大标题字号（48pt）
    title_ok = False
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text.strip() in COVER_TITLE and run.font.size and run.font.size.pt >= 47:
                title_ok = True
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip() in COVER_TITLE and run.font.size and run.font.size.pt >= 47:
                            title_ok = True
    if not title_ok:
        errors.append("封面大标题字号须为 48pt")

    # 有图例时：图例放大（宽 ≥ 正文85% 或高填满可用区）+ 环节拆解另起一页
    if render["figure_images"]:
        usable_cm = PAGE_SIZE["size_cm"][0] - 2 * 1.0   # 13.0cm
        min_w = 0.85 * usable_cm * 360000
        avail_h_cm = 11.0
        min_h = 0.95 * avail_h_cm * 360000
        big_img = False
        for dw in doc.element.body.findall(".//" + qn("w:drawing")):
            if dw.find(qn("wp:inline")) is None:
                continue  # 封面底层锚定图不计入图例
            ext = dw.find(".//" + qn("wp:extent"))
            if ext is None or ext.get("cx") is None:
                continue
            cx = int(ext.get("cx"))
            cy = int(ext.get("cy") or 0)
            if cx >= min_w or cy >= min_h:
                big_img = True
        if not big_img:
            errors.append("图例未放大（宽须≥正文85%，或高填满可用区）")
        titles = [i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "环节拆解"]
        if titles:
            title_p = doc.paragraphs[titles[0]]
            if not title_p.paragraph_format.page_break_before:
                errors.append("图例后环节拆解未另起一页")
    return errors


def run_render_check(docx_path: Path) -> None:
    """渲染检查（真实执行）：soffice 转 PDF + 逐页版式检查。工具缺失即失败，不跳过。

    进程内调用 render_docx（runpy，不走 shell、不拼命令行字符串）。
    通过后在 desktop-attachments/rendered/render_check.pdf 留存证据
    （固定文件名，工作流据此核验渲染检查确实执行过）。
    """
    import runpy

    out_dir = TEMPLATE_PATH.parent / "rendered"
    out_dir.mkdir(parents=True, exist_ok=True)
    argv_backup = sys.argv
    code = 0
    try:
        sys.argv = [
            "render_docx.py", str(docx_path), "--emit_pdf", "--check",
            "--out", str(out_dir), "--cta", DRAIN_TEXT,
        ]
        try:
            runpy.run_path(str(SCRIPT_DIR / "render_docx.py"), run_name="__main__")
        except SystemExit as exc:
            code = int(exc.code or 0)
    finally:
        sys.argv = argv_backup
    if code != 0:
        raise RuntimeError(f"渲染检查未过（exit={code}，见上方明细）")
    pdfs = [p for p in out_dir.iterdir()
            if p.suffix == ".pdf" and p.name != "render_check.pdf"]
    pdfs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    if pdfs:
        evidence = out_dir / "render_check.pdf"
        if evidence.exists():
            evidence.unlink()
        os.replace(pdfs[0], evidence)
        print(f"      渲染检查证据：{evidence}")


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------


def main():
    temp_path = None
    stage = "读取待写入内容"
    try:
        print("[1/7] 读取待写入内容（draft@3）")
        draft = load_draft()
        render = draft["render"]
        print(f"      环节：{render['segment_name']}（{SEGMENT_TYPES[draft['segment']['type']]}）")

        stage = "视图门"
        print("[2/7] 视图门（可生成性 / 误收选题拦截）")
        gate_view(draft["source_view_entry"])
        print("      ✅ 视图门通过")

        stage = "评分门"
        print("[3/7] 评分门（事实锁定 + 100 分量表 + 硬门）")
        result = gate_score(draft, draft["source_view_entry"], core.BookLibrary())
        br = result["duration"]
        print(
            f"      ✅ 评分门通过（口播 {br['speech_sec']:.0f}s + 示范停顿 {br['demo_pause_sec']:.0f}s"
            f" = 合计 {br['total_sec']:.0f}s）"
        )

        stage = "评审门"
        print("[4/7] 评审门（内容评审记录与草稿同版本）")
        gate_review(draft)
        print("      ✅ 评审门通过（五项内容评审全过，版本一致）")

        stage = "创建隔离构建副本"
        print("[5/7] 构建文档")
        TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        fd, temp_name = tempfile.mkstemp(prefix=".trial_daily_build_", suffix=".docx", dir=TEMPLATE_PATH.parent)
        os.close(fd)
        temp_path = Path(temp_name)
        doc = build_doc(draft)
        doc.save(temp_path)
        print(f"      已生成 {len(render['figure_images'])} 张图例，结构构建完成")

        stage = "自动验证"
        print("[6/7] 格式验证 + 渲染检查")
        check = Document(temp_path)
        errors = validate_output(check, draft)
        if errors:
            raise RuntimeError("；".join(errors))
        print("      ✅ 构建验证通过")
        run_render_check(temp_path)

        stage = "快照并原子提交"
        print("[7/7] 快照并原子提交")
        snap = take_snapshot()
        if snap:
            print(f"      快照：{snap.resolve()}")
        os.replace(temp_path, TEMPLATE_PATH)
        temp_path = None
        print(f"      ✅ 已原子提交：{TEMPLATE_PATH}")
        print("✅ 全部通过（视图门、评分门、评审门、格式验证、渲染检查均已执行）")
    except Exception as exc:
        if temp_path is not None and temp_path.exists():
            try:
                temp_path.unlink()
            except OSError as cleanup_exc:
                print(f"      ⚠️ 临时文件清理失败：{cleanup_exc}")
        print(f"      ❌ {stage}失败：{exc}")
        if TEMPLATE_PATH.exists():
            print("      旧工作副本未改动")
        else:
            print("      首次运行失败，未留下半成品工作副本")
        if PENDING_JSON.exists():
            print("[KEEP] pending_trial_daily.json 已原样保留，供排查")
        sys.exit(1)

    print("[收尾] 删除中间 JSON")
    try:
        PENDING_JSON.unlink()
    except Exception as exc:
        print(f"      ❌ DOCX 已通过并提交，但 JSON 删除失败：{exc}")
        sys.exit(2)
    print("      ✅ pending_trial_daily.json 已删除")
    print("完成")


if __name__ == "__main__":
    main()
