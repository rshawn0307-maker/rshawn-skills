# -*- coding: utf-8 -*-
"""PE-trial-daily 统一测试入口（任务1 扩展版）。

隔离测试：在临时工作区构建 pending_trial_daily.json，运行 fill_trial_daily_post.py，
验证生成 DOCX 的结构契约（封面、图例、环节拆解、易犯错误表格、试讲逐字稿、引流页）；
并覆盖任务1 核心库 ptd_core 的可生成视图、dry-run 迁移、事实锁定、诚实性硬门、
三类片段流程与 100 分量表放行线（含故意制造坏草稿的反向用例）。

用法:
    python3 test_fill_trial_daily_post.py

退出码 0 当且仅当全部用例通过；任一失败打印 ❌ 用例名 + 原因。skipped=0。
"""
import copy
import json
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))
import ptd_core  # noqa: E402
import build_fixtures  # noqa: E402
import ptd_workflow  # noqa: E402
import render_docx  # noqa: E402
import fill_trial_daily_post as fill  # noqa: E402

CYAN_LEGACY = "关注我，每天一个体育试讲设计，帮你备考上岸"
FILL_SCRIPT = SCRIPT_DIR / "fill_trial_daily_post.py"

# 源数据哈希快照（导入时；rev_data_hashes_unchanged 用）
import hashlib as _hashlib  # noqa: E402

_SRC_HASHES = {
    p.name: _hashlib.sha256(p.read_bytes()).hexdigest()
    for p in (ptd_core.INDEX_DEFAULT, ptd_core.PROGRESS_DEFAULT) if p.exists()
}

# ---------------------------------------------------------------------------
# 用例注册器
# ---------------------------------------------------------------------------

CASES = []


def case(fn):
    CASES.append((fn.__name__, fn))
    return fn


def _assert(cond, detail=""):
    if not cond:
        raise AssertionError(detail or "断言失败")


# ---------------------------------------------------------------------------
# 只读源缓存
# ---------------------------------------------------------------------------

_cache = {}


def _view():
    if "view" not in _cache:
        _cache["view"] = ptd_core.build_generatable_view()
    return _cache["view"]


def _fixtures():
    if "fixtures" not in _cache:
        _cache["fixtures"] = build_fixtures.build_all()
    return _cache["fixtures"]


def _lib():
    if "lib" not in _cache:
        _cache["lib"] = ptd_core.BookLibrary()
    return _cache["lib"]


def _entry(view, rid):
    return next(e for e in view["entries"] if e["id"] == rid)


# ===========================================================================
# 可生成视图 / 稳定 ID
# ===========================================================================


@case
def view_total_313():
    _assert(len(_view()["entries"]) == 313, f"entries={len(_view()['entries'])}")


@case
def view_ids_unique():
    ids = [e["id"] for e in _view()["entries"]]
    _assert(len(set(ids)) == len(ids), "稳定 ID 存在重复")


@case
def view_difficulty_missing_171():
    st = _view()["stats"]
    _assert(st["difficulty_missing"] == 171, f"difficulty_missing={st['difficulty_missing']}")


@case
def view_generatable_302_blockers_11():
    st = _view()["stats"]
    # v3：误收选题（课外作业建议等 10 条）进入 blocker，可生成数 302→292
    _assert(st["generatable"] == 292, f"generatable={st['generatable']}")
    _assert(st["blockers"] == 21, f"blockers={st['blockers']}")
    _assert(st["miscollected_topic"] == 10, f"miscollected={st['miscollected_topic']}")
    _assert(st["method_cross_reference"] >= 10, f"cross_ref={st['method_cross_reference']}")


@case
def view_blockers_all_wushu_no_pdf():
    blk = [e for e in _view()["entries"] if e["generatable_blockers"]]
    _assert(len(blk) == 21, f"blocked={len(blk)}")
    for e in blk:
        if "miscollected_topic_not_teaching_activity" in e["generatable_blockers"]:
            _assert("miscollected_topic" in e["flags"], f"{e['id']} 误收未标记")
            continue
        _assert("武术" in e["book_file"], f"{e['id']} 阻塞非武术/误收")
        _assert("figure_required_but_pdf_missing" in e["generatable_blockers"], e["id"])
        _assert(not e["book_pdf_available"], f"{e['id']} 应无 PDF")


@case
def view_use_extracted_implies_pdf():
    for e in _view()["entries"]:
        if e["figure_policy"] == "use_extracted":
            _assert(e["book_pdf_available"], f"{e['id']} use_extracted 但缺 PDF")


@case
def view_needs_ocr_verify_15():
    n = sum(1 for e in _view()["entries"] if "figure_needs_ocr_verify" in e["flags"])
    _assert(n == 15, f"needs_ocr_verify={n}")


@case
def view_p1_no_figure_generatable():
    e = _entry(_view(), "PTD-000-乒乓球-平击球")
    _assert(e["figure_policy"] == "none", e["figure_policy"])
    _assert(e["generatable"], "P1 应可生成")


@case
def view_p2_figures_ok():
    e = _entry(_view(), "PTD-244-篮球-原地运球")
    _assert(e["figure_policy"] == "use_extracted", e["figure_policy"])
    _assert(e["figures"] and all(f["match"] == "ok" for f in e["figures"]), e["figures"])


@case
def view_p3_misattribution_kept_not_deleted():
    e = _entry(_view(), "PTD-048-体能-照镜子")
    _assert("figure_misattribution_suspect" in e["flags"], "照镜子误收应留证")
    _assert(e["figure_policy"] == "misattributed_treat_as_none", e["figure_policy"])
    _assert(e["generatable"], "误收按无图处理，仍可生成")


@case
def view_misattribution_suspect_total_28():
    st = _view()["stats"]
    _assert(st["figure_misattribution_suspect"] == 28, st["figure_misattribution_suspect"])


# ===========================================================================
# dry-run 迁移表（孤儿不丢）
# ===========================================================================


@case
def migration_rows_3_dry_run():
    m = _view()["migration_dryrun"]
    import json as _json
    prog = _json.loads(ptd_core.PROGRESS_DEFAULT.read_text(encoding="utf-8"))
    _assert(m["mode"] == "dry-run", m["mode"])
    _assert(len(m["rows"]) == len(prog.get("done", [])), f"rows={len(m['rows'])}")


@case
def migration_orphan_kept_classified():
    m = _view()["migration_dryrun"]
    row = next(r for r in m["rows"] if r["progress_name"] == "技巧大赛")
    _assert(row["disposition"] == "orphan_keep_classified", row["disposition"])
    _assert(row["view_id"] is None, "孤儿不应被硬匹配")


@case
def migration_other_two_migrate():
    m = _view()["migration_dryrun"]
    by = {r["progress_name"]: r for r in m["rows"]}
    _assert(by["半米字移动"]["view_id"] == "PTD-045-体能-半米字移动", by["半米字移动"])
    _assert(by["抢背后滚球"]["view_id"] == "PTD-046-体能-抢背后滚球", by["抢背后滚球"])
    _assert(by["半米字移动"]["disposition"] == by["抢背后滚球"]["disposition"] == "migrate")


# ===========================================================================
# 三类片段流程 / 片段要素 / 量表
# ===========================================================================


@case
def flows_cover_three_types():
    _assert({"practice", "game", "fitness"} <= set(ptd_core.FLOWS), set(ptd_core.FLOWS))


@case
def flow_stage_counts():
    _assert(len(ptd_core.FLOWS["practice"]) == 5, len(ptd_core.FLOWS["practice"]))
    _assert(len(ptd_core.FLOWS["game"]) == 5, len(ptd_core.FLOWS["game"]))
    _assert(len(ptd_core.FLOWS["fitness"]) == 4, len(ptd_core.FLOWS["fitness"]))


@case
def flow_sec_ranges_valid():
    for t, stages in ptd_core.FLOWS.items():
        for s in stages:
            lo, hi = s["sec_range"]
            _assert(0 < lo <= hi <= 120, f"{t} {s['stage']} {s['sec_range']}")
            _assert(s["stage"] and s["purpose"], f"{t} 缺 stage/purpose")


@case
def segment_fields_complete():
    want = ["学段", "片段位置", "时长", "重点", "器材", "安全", "分层", "评价"]
    _assert(ptd_core.SEGMENT_FIELDS == want, ptd_core.SEGMENT_FIELDS)


@case
def scale_matches_frozen_thresholds():
    _assert(ptd_core.SCALE == {
        "教材事实": 30, "考编可用": 20, "安全": 20, "教学": 15, "口语": 10, "证据": 5,
    }, ptd_core.SCALE)
    _assert(ptd_core.RELEASE == {
        "total_min": 85, "textbook_min": 27, "safety_min": 16, "hard_gate_max": 0,
    }, ptd_core.RELEASE)


# ===========================================================================
# fixture 正向（v3）：v2 存档草稿=回归反例，必须被新硬门拦截
# ===========================================================================

V2_ARCHIVE_EXPECT = {
    # fixtures9 里的 v2 草稿带哪些病灶（v3 硬门应全部命中）
    "PTD-000-乒乓球-平击球": {"broken_punctuation", "cross_ref_evidence_missing",
                          "duration_annotation_mismatch", "factlock_unclassified_gt0"},
}


@case
def fixture_v2_archive_now_blocked():
    for f in _fixtures():
        entry = _entry(_view(), f["draft"]["id"])  # v3 视图条目（含新 flags）
        res = ptd_core.score_draft(f["draft"], entry, _lib())
        got = set(res["hard_gates"])
        _assert(not res["release"], f"{f['draft']['id']} v2 存档不应放行")
        _assert("factlock_unclassified_gt0" in got,
                f"{f['draft']['id']} v2 草稿无理由登记必须被拦（got={got}）")
        if "method_cross_reference" in (entry.get("flags") or []):
            md = int(entry.get("md_line", 0))
            has_prior = any(
                (ev.get("line") is not None and int(ev["line"]) < md)
                for ev in f["draft"]["fields"]["method"].get("evidence") or []
            )
            if has_prior:
                _assert("cross_ref_evidence_missing" not in got,
                        f"{f['draft']['id']} 已补前文证据不应误拦")
            else:
                _assert("cross_ref_evidence_missing" in got,
                        f"{f['draft']['id']} 前文引用未补证据必须被拦（got={got}）")


@case
def fixture_flow_stages_match_template():
    for f in _fixtures():
        want = [s["stage"] for s in ptd_core.FLOWS[f["view"]["activity_type"]]]
        got = [st["stage"] for st in f["draft"]["flow"]]
        _assert(got == want, f"{f['draft']['id']} {got} != {want}")


@case
def fixture_covers_practice_and_fitness():
    types = {f["view"]["activity_type"] for f in _fixtures()}
    _assert("practice" in types and "fitness" in types, types)


@case
def difficulty_empty_no_fabricated_stars():
    for f in _fixtures():
        d = f["draft"]["fields"]["difficulty"]
        if not f["view"]["index_difficulty"]:
            _assert(d["kind"] != "index_stars", f"{f['draft']['id']} 空难度不得编星")
            _assert("★" not in d.get("display", ""), f"{f['draft']['id']} display 含星")
            _assert(d["kind"] == "index_empty_adapted", f"{f['draft']['id']} kind={d['kind']}")


@case
def p2_errors_textbook_legit_not_faked():
    f = _fixtures()[1]  # P2 篮球原地运球（索引 has_errors=True）
    _assert(f["view"]["index_has_errors"], "P2 应标记有教材纠错")
    rows = f["draft"]["fields"]["errors"]["rows"]
    _assert(len(rows) == 3, f"P2 纠错行数={len(rows)}")
    for i, r in enumerate(rows):
        _assert(r["error"]["provenance"] == "textbook", f"P2 纠错[{i}] 应教材原文")
        _assert(r["fix"]["provenance"] == "textbook", f"P2 纠正[{i}] 应教材原文")
    _assert(ptd_core.check_honesty(f["draft"], f["view"]) == [], "P2 教材纠错不应误判为造假")


# ===========================================================================
# 反向用例：坏草稿必须被拒绝（先红后绿已在修复期验证）
# ===========================================================================


def _neg_base():
    return copy.deepcopy(_fixtures()[0]["draft"]), _fixtures()[0]["view"]


@case
def neg_textbook_no_evidence_rejected():
    draft, view = _neg_base()
    draft["fields"]["method"]["evidence"] = []
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "教材块缺证据仍放行")
    _assert(res["scores"]["教材事实"] < 27, f"教材={res['scores']['教材事实']}")


@case
def neg_evidence_line_mismatch_rejected():
    draft, view = _neg_base()
    draft["fields"]["method"]["evidence"][0]["line"] = 999999
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "行号与摘录不符仍放行")


@case
def neg_adapted_unregistered_token_rejected():
    draft, view = _neg_base()
    draft["flow"][1]["script"] += "。每个人跳三次"
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "未归类事实 token 仍放行")
    _assert("factlock_unclassified_gt0" in res["hard_gates"], res["hard_gates"])


@case
def neg_fabricated_difficulty_rejected():
    draft, view = _neg_base()
    draft["fields"]["difficulty"] = {"kind": "index_stars", "display": "★★", "provenance": "textbook"}
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "空难度编星仍放行")
    _assert("fabricated_difficulty" in res["hard_gates"], res["hard_gates"])


@case
def neg_practice_errors_faked_rejected():
    draft, view = _neg_base()
    _assert(view["activity_type"] == "practice" and not view["index_has_errors"], "前置：P1 无教材纠错标记")
    for r in draft["fields"]["errors"]["rows"]:
        r["error"]["provenance"] = "textbook"
        r["fix"]["provenance"] = "textbook"
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "practice 纠错冒充教材原文仍放行")
    _assert(any(h.startswith("practice_errors_faked") for h in res["hard_gates"]), res["hard_gates"])


@case
def neg_duration_out_of_range_rejected():
    draft, view = _neg_base()
    draft["config"]["segment_duration_sec"] = [10, 20]
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "口播时长超范围仍放行")
    _assert("script_duration_out_of_range" in res["hard_gates"], res["hard_gates"])


@case
def neg_no_textbook_block_rejected():
    draft, view = _neg_base()
    for k in ("method", "intent"):
        draft["fields"][k]["provenance"] = "adapted"
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "无教材原文支撑仍放行")
    _assert(res["scores"]["教材事实"] < 27, f"教材={res['scores']['教材事实']}")


# ===========================================================================
# excerpt_at 行级证据校验
# ===========================================================================


@case
def excerpt_at_valid_line_true():
    book = "人教版教师用书-乒乓球.md"
    line = 952
    raw = _lib().lines(book)[line]
    _assert(_lib().excerpt_at(book, line, raw), "真实行+整行摘录应通过")


@case
def excerpt_at_empty_line_rejected():
    book = "人教版教师用书-乒乓球.md"
    lines = _lib().lines(book)
    empty = next((i for i, ln in enumerate(lines) if not ln.strip()), None)
    _assert(empty is not None, "教材中应存在空行供测试")
    _assert(not _lib().excerpt_at(book, empty, "任意内容"), "空行不得通过行级校验")


@case
def excerpt_at_wrong_line_rejected():
    book = "人教版教师用书-乒乓球.md"
    _assert(not _lib().excerpt_at(book, 100, "完全不存在的内容内容内容"), "错误行不得通过")


# ===========================================================================
# 任务2：3:4 版式契约（XML 级快查 + 官方渲染集成）
# ===========================================================================

LAYOUT_TITLES = {"环节拆解", "易犯错误与纠正", "试讲逐字稿", "图例直观"}


def _doc_runs(doc):
    """遍历段落与表格单元格里的所有 run。"""
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        yield r


@case
def layout_page_size_3_4():
    fr = _fill_result()
    _assert(fr["doc"] is not None)
    for si, sec in enumerate(fr["doc"].sections):
        ratio = sec.page_width / sec.page_height
        _assert(abs(ratio - 0.75) / 0.75 <= 0.01, f"第{si+1}节页宽高比 {ratio:.4f}")


@case
def layout_cover_brand_bg_full_page():
    fr = _fill_result()
    doc = fr["doc"]
    _assert(doc is not None)
    anchors = [a for a in doc.element.body.findall(".//" + qn("wp:anchor"))
               if a.get("behindDoc") == "1"]
    _assert(anchors, "封面底层图缺失")
    anchor = anchors[0]
    _assert(anchor.get("layoutInCell") == "0", "封面底图仍受单元格裁切")
    extent = anchor.find(qn("wp:extent"))
    _assert(extent is not None, "封面底图缺少尺寸")
    _assert(abs(int(extent.get("cx") or 0) - int(doc.sections[0].page_width)) <= 635,
            "封面底图宽度未精确贴页")
    _assert(abs(int(extent.get("cy") or 0) - int(doc.sections[0].page_height)) <= 635,
            "封面底图高度未精确贴页")
    for tag in ("wp:positionH", "wp:positionV"):
        pos = anchor.find(qn(tag))
        off = pos.find(qn("wp:posOffset")) if pos is not None else None
        _assert(pos is not None and pos.get("relativeFrom") == "page" and
                off is not None and off.text == "0", f"{tag} 未贴齐页边")
    shd = doc.tables[0].cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd"))
    _assert(shd is None, "封面单元格底色会遮住底部标语和 SHTr")


@case
def layout_font_cjk_contract():
    fr = _fill_result()
    style = fr["doc"].styles["Normal"]
    rpr = style._element.rPr
    east = rpr.rFonts.get(qn("w:eastAsia")) if rpr is not None and rpr.rFonts is not None else None
    _assert(east and any(c in east for c in
            ("Hiragino", "Heiti", "Songti", "PingFang", "Microsoft YaHei", "Noto", "WenQuanYi")),
            f"Normal eastAsia 字体 {east} 不在 CJK 契约")


@case
def layout_body_min_18pt():
    fr = _fill_result()
    for r in _doc_runs(fr["doc"]):
        t = (r.text or "").strip()
        if t in LAYOUT_TITLES or t.startswith("#") or t in ("易犯错误", "纠正方法"):
            continue
        if r.font.size and any(k in t for k in ("两腿微屈", "同学们好", "降低重心", "好收球")):
            _assert(r.font.size.pt >= 18, f"正文 {t[:12]} 字号 {r.font.size.pt}")


@case
def layout_section_title_24_28pt():
    fr = _fill_result()
    for r in _doc_runs(fr["doc"]):
        if (r.text or "").strip() in LAYOUT_TITLES:
            _assert(r.font.size and 24 <= r.font.size.pt <= 28,
                    f"栏目 {r.text} 字号 {r.font.size and r.font.size.pt}")


@case
def layout_caption_label_min_16pt():
    fr = _fill_result()
    for r in _doc_runs(fr["doc"]):
        t = (r.text or "").strip()
        if t.startswith("图") and "3-2-7" in t:
            _assert(r.font.size and r.font.size.pt >= 16, f"图注 {t[:12]} {r.font.size and r.font.size.pt}")
        if t.startswith("#"):
            _assert(r.font.size and r.font.size.pt >= 16, f"标签 {t[:12]} {r.font.size and r.font.size.pt}")


@case
def layout_cta_min_18pt():
    fr = _fill_result()
    for r in _doc_runs(fr["doc"]):
        if (r.text or "").strip() == CYAN_LEGACY:
            _assert(r.font.size and r.font.size.pt >= 18, f"CTA 字号 {r.font.size and r.font.size.pt}")


@case
def layout_script_split_short_lines():
    fr = _fill_result()
    scripts = [p.text for p in fr["doc"].paragraphs
               if p.text and any(k in p.text for k in ("同学们好", "看我示范", "跟我做", "低运球", "好收球"))]
    _assert(len(scripts) >= 3, f"逐字稿仅 {len(scripts)} 段（应按语义分段）")
    for s in scripts:
        # v3 语义分段：整句成段（≤60 字近似上限），不以逗号/顿号截断，无孤立标点
        _assert(len(s) <= 60, f"段落过长 {len(s)} 字：{s[:20]}")
        _assert(re.fullmatch(r"[。！？，、；：:]+", s) is None, f"孤立标点段：{s!r}")


@case
def layout_table_fixed_42_58():
    fr = _fill_result()
    err_tbl = None
    for t in fr["doc"].tables:
        if len(t.columns) == 2 and t.cell(0, 0).text.strip() == "易犯错误":
            err_tbl = t
            break
    _assert(err_tbl is not None, "未找到易犯错误表")
    tblPr = err_tbl._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    _assert(layout is not None and layout.get(qn("w:type")) == "fixed", "表格须固定布局(fixed)")
    grid = err_tbl._tbl.tblGrid.findall(qn("w:gridCol"))
    _assert(len(grid) == 2, f"gridCol={len(grid)}")
    w1, w2 = int(grid[0].get(qn("w:w"))), int(grid[1].get(qn("w:w")))
    _assert(abs(w1 / (w1 + w2) - 0.42) < 0.03, f"列宽比 {w1}/{w1 + w2}")
    _assert(len(err_tbl._tbl.findall(".//" + qn("w:cantSplit"))) >= 2, "表格行缺 cantSplit（行会跨页）")


@case
def layout_watermark_opacity_8_12():
    fr = _fill_result()
    for sec in fr["doc"].sections:
        x = sec.header._element.xml
        m = re.search(r'opacity="([0-9a-fA-F]{6})"', x)
        _assert(m, "水印无 opacity 属性")
        frac = int(m.group(1), 16) / 0xFFFFFF
        _assert(0.08 <= frac <= 0.12, f"水印透明度 {frac:.3f}（期望 8%–12%）")


@case
def render_full_checks_pass():
    """官方 render_docx.py --emit_pdf --check 全页检查（3:4/无方框/无裁切/CTA不孤页/水印）。"""
    fr = _fill_result()
    _assert(fr["docx_path"].exists(), "缺少输出 DOCX")
    r = subprocess.run(
        [sys.executable, str(SCRIPT_DIR / "render_docx.py"), str(fr["docx_path"]),
         "--emit_pdf", "--check", "--out", str(fr["ws"] / "rendered")],
        capture_output=True, text=True, env=dict(os.environ),
    )
    _assert(r.returncode == 0, f"渲染检查失败:\n{r.stdout}\n{r.stderr}")


# ===========================================================================
# 任务3：工作流（锁 / 状态机 / OCR缓存 / 图例STOP / IMA幂等）
# ===========================================================================


@case
def wf_deps_check_runs():
    # v3：工作流依赖预检 = python-docx 可导入（内联检查）
    try:
        import docx  # noqa: F401
    except ImportError:
        _assert(False, "python-docx 应可导入")
    _assert(hasattr(ptd_workflow, "run_workflow"), "v3 工作流应提供 run_workflow")


@case
def wf_lock_exclusive_two_processes():
    ws = Path(tempfile.mkdtemp(prefix="wf_lock_"))
    try:
        _assert(ptd_workflow.acquire_lock(ws, timeout=1), "首个持有者应拿到锁")
        # 锁被占用时其他获取者必须失败（排他）
        _assert(not ptd_workflow.acquire_lock(ws, timeout=0.5), "锁被占时应失败")
        ptd_workflow.release_lock(ws)
        _assert(ptd_workflow.acquire_lock(ws, timeout=1), "释放后应能拿到锁")
    finally:
        ptd_workflow.release_lock(ws)


@case
def wf_state_idempotent_terminal():
    ws = Path(tempfile.mkdtemp(prefix="wf_state_"))
    ptd_workflow.advance(ws, "PTD-X", "到终态", to="progress_commit")
    st = json.loads(ptd_workflow.state_path(ws).read_text(encoding="utf-8"))
    _assert(st["entries"]["PTD-X"]["stage"] == "progress_commit", "终态应被记录")
    ptd_workflow.advance(ws, "PTD-X", "追加日志")
    st = json.loads(ptd_workflow.state_path(ws).read_text(encoding="utf-8"))
    _assert("追加日志" in " ".join(st["entries"]["PTD-X"]["log"]), "日志应保留")


@case
def wf_state_write_atomic_valid():
    ws = Path(tempfile.mkdtemp(prefix="wf_atom_"))
    for i in range(5):
        ptd_workflow.advance(ws, f"PTD-{i}", "n", to="docx_commit")
        st = json.loads(ptd_workflow.state_path(ws).read_text(encoding="utf-8"))
        _assert(f"PTD-{i}" in st["entries"], f"第{i}次写入后状态应有效")
    _assert(not list(ws.glob(".wf_*")), "临时状态文件应清理")


@case
def wf_draft_hash_stable():
    a = ptd_core.draft_hash({"x": [1, 2], "s": "中文"})
    b = ptd_core.draft_hash({"s": "中文", "x": [1, 2]})
    _assert(a == b, "draft_hash 应稳定（与键序无关）")


@case
def wf_gate_select_blocks_miscollected():
    view = _entry(_view(), "PTD-016-乒乓球-课外作业建议")
    draft = {"id": view["id"], "schema": ptd_core.SCHEMA_DRAFT, "source_view_entry": view}
    try:
        ptd_workflow.gate_select(draft)
        _assert(False, "误收选题应被 select 门拦截")
    except ptd_workflow.GateStop as exc:
        _assert(exc.args[0][0] == "select", exc.args[0])


@case
def wf_gate_extract_stops_missing_figure():
    entry = dict(_entry(_view(), "PTD-045-体能-半米字移动"))
    entry["figure_policy"] = "use_extracted"
    draft = {"id": entry["id"], "schema": ptd_core.SCHEMA_DRAFT,
             "source_view_entry": entry, "render": {"figure_images": []}}
    try:
        ptd_workflow.gate_extract(draft, entry, _lib())
        _assert(False, "use_extracted 缺图应 STOP（不再跳图放行）")
    except ptd_workflow.GateStop as exc:
        _assert("图例" in exc.args[0][1], exc.args[0])


@case
def wf_v3_topic_and_repetition_units():
    """v3 新门单元：题文一致 / 近重复句 / 断裂标点 / 安全口令。"""
    draft, view = _neg_base()  # 平击球 v2 草稿 + v2 视图
    _assert(ptd_core.check_topic_match(draft, view), "平击球稿应命中题文")
    bad = copy.deepcopy(draft)
    for st in bad["flow"]:
        st["script"] = "同学们好，今天我们学队列队形。向右看齐，向前看。"
    _assert(not ptd_core.check_topic_match(bad, view), "讲别的活动应判题文错配")
    dup_text = ("大家注意看老师的示范动作一起跟着做一做。"
                "下面大家注意看老师的示范动作一起跟着做一做。")
    _assert(len(ptd_core.near_duplicate_pairs(dup_text)) >= 1, "重复句应被识别")
    _assert(ptd_core.check_broken_punctuation("动作相同。。看明白") , "双句号应被识别")
    ok_draft = {"segment": {"meta": {"安全": "检查场地平整，器材摆放到位，两人间隔一臂距离，听到哨音立即停止并退回起点，不推搡"}},
                "flow": []}
    exec_ok, cats = ptd_core.check_safety_executable(ok_draft)
    _assert(exec_ok and len(cats) >= 3, f"可执行安全应通过: {exec_ok} {cats}")
    bad_safety = {"segment": {"meta": {"安全": "注意安全，保护好器材，注意负荷"}},
                  "flow": []}
    exec_ok2, cats2 = ptd_core.check_safety_executable(bad_safety)
    _assert(not exec_ok2, f"安全套话不得通过: {exec_ok2} {cats2}")


@case
def wf_v3_review_validate_units():
    """v3 评审记录：同版本通过 / 版本不一致拦截 / 缺理由拦截。"""
    f = _fixtures()[0]
    draft = copy.deepcopy(f["draft"])
    view = _entry(_view(), draft["id"])
    draft.setdefault("notes", {})["human_rewrite_applied"] = True
    review = {
        "schema": ptd_core.SCHEMA_REVIEW, "id": draft["id"],
        "draft_sha": ptd_core.draft_hash(draft), "reviewer": "agent",
        "checked": {k: True for k in ptd_core.REVIEW_CHECKS},
        "suggestion_notes": [{"where": "flow[0]", "why": "口令化改写，动作未变"}],
        "verdict": "pass",
    }
    ok, errors = ptd_core.validate_review(review, draft, view)
    _assert(ok, f"同版本评审应通过: {errors}")
    stale = dict(review, draft_sha="0" * 16)
    ok2, errors2 = ptd_core.validate_review(stale, draft, view)
    _assert(not ok2 and any("版本不一致" in e for e in errors2), "沿用旧审核必须被拦")
    no_reason = dict(review, suggestion_notes=[])
    ok3, errors3 = ptd_core.validate_review(no_reason, draft, view)
    _assert(not ok3, "缺建议理由必须被拦")


@case
def wf_workflow_dryrun_advances_states():
    """v3：误收选题走 run_workflow 应在 select 门 STOP（退出码 4），状态留痕。"""
    ws = Path(tempfile.mkdtemp(prefix="wf_run_"))
    tmp = Path(tempfile.mkdtemp(prefix="wf_draft_"))
    view = _entry(_view(), "PTD-016-乒乓球-课外作业建议")
    draft = {"schema": ptd_core.SCHEMA_DRAFT, "id": view["id"], "source_view_entry": view}
    dp = tmp / "pending_trial_daily.json"
    dp.write_text(json.dumps(draft, ensure_ascii=False), encoding="utf-8")
    rc = ptd_workflow.run_workflow(dp, ws, tmp, dry_run=True)
    _assert(rc == 4, f"误收选题应 STOP，rc={rc}")
    st = ptd_workflow.read_state(ws)
    log = " ".join(st["entries"][view["id"]].get("log", []))
    _assert("STOP@select" in log, f"应留 STOP 痕迹: {log}")
    _assert(not (ws / "workflow.lock").exists(), "运行结束应释放锁")


# ===========================================================================
# 任务4：反向验证（故意制造坏输入 → 先红；好输入全绿；正式数据哈希不变）
# ===========================================================================


@case
def rev_short_script_rejected():
    draft, view = _neg_base()
    draft["flow"] = [{"stage": "导入与示范", "sec": 30, "provenance": "adapted",
                      "evidence": [{"book_file": "人教版教师用书-乒乓球.md", "line": 952,
                                    "excerpt": "正手平击球与正手平击发球动作相同。"}],
                      "adapted_facts": [], "script": "同学们好，今天学平击球。看示范。"}]
    res = ptd_core.score_draft(draft, view, _lib())
    _assert(not res["release"], "短稿（口播不足 2 分钟）仍放行")
    _assert("script_duration_out_of_range" in res["hard_gates"], res["hard_gates"])


@case
def rev_direction_reversed_rejected():
    """human-writing 后教材方向被静默反转（同侧→对侧），事实锁定必须拒绝。"""
    f = _fixtures()[1]  # P2 篮球原地运球（教材原文：同侧脚外侧前方）
    draft = copy.deepcopy(f["draft"])
    assert "同侧脚的外侧前方" in draft["flow"][3]["script"], "前置：原稿含同侧方向"
    draft["flow"][3]["script"] = draft["flow"][3]["script"].replace("同侧脚的外侧前方", "对侧脚的外侧前方")
    res = ptd_core.score_draft(draft, f["view"], _lib())
    _assert(not res["release"], "方向被静默反转仍放行")
    _assert(res["factlock"]["unclassified"] > 0, "反转后应有未归类事实 token")


@case
def rev_figure_width_1cm_rejected():
    fr = _fill_result()
    _assert(fr["doc"] is not None, "缺 docx")
    # 把内联图例 extent 改为 1cm×1cm（图宽 1cm，既不够宽也不够高）
    for p in fr["doc"].paragraphs:
        for dw in p._element.findall(".//" + qn("w:drawing")):
            inline = dw.find(qn("wp:inline"))
            if inline is None:
                continue
            ext = inline.find(qn("wp:extent"))
            if ext is not None:
                ext.set("cx", "360000")
                ext.set("cy", "360000")
    errors = fill.validate_output(fr["doc"], fr["pending"])
    _assert(any("图例未放大" in e for e in errors), f"1cm 图应被拒: {errors}")


@case
def rev_cover_bg_overflow_rejected():
    fr = _fill_result()
    anchor = next(a for a in fr["doc"].element.body.findall(".//" + qn("wp:anchor"))
                  if a.get("behindDoc") == "1")
    extent = anchor.find(qn("wp:extent"))
    extent.set("cx", str(round(int(extent.get("cx")) * 1.12)))
    extent.set("cy", str(round(int(extent.get("cy")) * 1.12)))
    errors = fill.validate_output(fr["doc"], fr["pending"])
    _assert(any("禁止放大出血" in e for e in errors), f"出血底图应被拒: {errors}")


@case
def rev_cover_bg_cell_clipping_rejected():
    fr = _fill_result()
    anchor = next(a for a in fr["doc"].element.body.findall(".//" + qn("wp:anchor"))
                  if a.get("behindDoc") == "1")
    anchor.set("layoutInCell", "1")
    tc_pr = fr["doc"].tables[0].cell(0, 0)._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "0B3289")
    tc_pr.append(shd)
    errors = fill.validate_output(fr["doc"], fr["pending"])
    _assert(any("layoutInCell" in e for e in errors), f"单元格裁切应被拒: {errors}")
    _assert(any("单元格必须透明" in e for e in errors), f"不透明底色应被拒: {errors}")


@case
def rev_cta_color_changed_rejected():
    fr = _fill_result()
    _assert(fr["doc"] is not None)
    changed = False
    for p in fr["doc"].paragraphs:
        if p.text.strip() == CYAN_LEGACY and p.runs:
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            changed = True
    _assert(changed, "未找到 CTA 段落")
    errors = fill.validate_output(fr["doc"], fr["pending"])
    _assert(any("固定引流段颜色" in e for e in errors), f"CTA 改色应被拒: {errors}")


@case
def rev_cta_orphan_page_rejected():
    pages = ["逐字稿第一行\n逐字稿第二行\n#标签\n" + CYAN_LEGACY,
             "只有 CTA 的孤页\n" + CYAN_LEGACY]
    bad = render_docx.check_cta_not_alone(pages, CYAN_LEGACY)
    _assert(bad, f"CTA 孤页应被拒: {bad}")
    _assert(any("仅" in b for b in bad), f"应指出正文行数不足: {bad}")


@case
def rev_ocr_partial_failure_not_verifiable():
    """v3：图例策略由视图驱动，needs_ocr_verify 缺图必须在 extract 门 STOP。"""
    entry = dict(_entry(_view(), "PTD-045-体能-半米字移动"))
    entry["figure_policy"] = "needs_ocr_verify"
    draft = {"id": entry["id"], "schema": ptd_core.SCHEMA_DRAFT,
             "source_view_entry": entry, "render": {"figure_images": []}}
    try:
        ptd_workflow.gate_extract(draft, entry, _lib())
        _assert(False, "needs_ocr_verify 缺图例不应放行")
    except ptd_workflow.GateStop:
        pass


@case
def rev_data_hashes_unchanged():
    """源数据在套件运行期间不得被改动（与导入时快照比对）。"""
    import hashlib

    def _h(p):
        return hashlib.sha256(Path(p).read_bytes()).hexdigest()

    for key, path in (("index", ptd_core.INDEX_DEFAULT), ("progress", ptd_core.PROGRESS_DEFAULT)):
        if key in _SRC_HASHES:
            _assert(_h(path) == _SRC_HASHES[key], f"{Path(path).name} 哈希已变: {_h(path)}")


# ===========================================================================
# 填充管线集成（原入口，单例缓存）
# ===========================================================================

_fill_cache = None


def _fill_result():
    global _fill_cache
    if _fill_cache is not None:
        return _fill_cache
    tmp = Path(tempfile.mkdtemp(prefix="trial_test_"))
    ws = tmp / "workspace"
    (ws / "scripts").mkdir(parents=True)
    (ws / "desktop-attachments").mkdir(parents=True)

    # 生成一张测试用 PNG（覆盖图例插入路径）
    w = h = 64
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * w for _ in range(h))

    def chunk(tag, data):
        c = tag + data
        return struct_pack(len(data)) + c + zlib_crc(c)

    import struct
    import zlib
    def struct_pack(n):
        return struct.pack(">I", n)
    def zlib_crc(c):
        return struct.pack(">I", zlib.crc32(c))

    fig = ws / "fig.png"
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    fig.write_bytes(png)

    view_entry = next(e for e in _view()["entries"]
                      if e["activity_name"] == "原地运球" and e["sport"] == "篮球")
    method_text = "两腿微屈上体稍前倾，以肘为轴前臂屈伸，用手指和指根触球，球落点在同侧脚外侧前方。"
    stages = [
        {"stage": "导入与示范", "demo_sec": 30, "pause_sec": 0,
         "script": "同学们好，我们先复习原地运球。每人拿一个球，分组散点站位，相互之间拉开距离。看我示范，两腿微屈上体前倾，以肘为轴用手指指根触球，掌心空出。注意看球反弹回来的位置，手指手腕要柔和地用力按拍。",
         "adapted_facts": ["距离", "触球", "前倾", "一个", "分组", "手指", "手腕"], "adapted_note": "口令化改写与队形安排属教学建议，动作要点为教材原文"},
        {"stage": "分解学练", "demo_sec": 25, "pause_sec": 0,
         "script": "跟我做，一二三四。先低运球十下，再高运球十下，手臂放松，用手指和指根按拍球。每做完一组停下来，抖抖手臂再继续。感觉手型不对的同学，把球拿起来先做徒手模仿。",
         "adapted_facts": ["按拍", "放松", "一组"], "adapted_note": "练习节奏与口令为教学安排"},
        {"stage": "纠错与对比", "demo_sec": 15, "pause_sec": 0,
         "script": "我转一圈看，小王做得很稳，球像粘在手上一样。小李你掌心太紧了，掌心空出来再运给我看。对，就是这样，手指先触球再随球送出去。",
         "adapted_facts": ["一圈", "手指"], "adapted_note": "纠错情境为教学建议"},
        {"stage": "巩固运用", "demo_sec": 20, "pause_sec": 0,
         "script": "抬头看我手势，我报几你们就运几下。运球不是拍球，要有迎球缓冲，球才会听话。两人一组互相报数，做完互换角色再练一轮。观察你们的同伴，手型标准的给它点个赞，掌心粘住的帮助提醒一声。",
         "adapted_facts": ["迎球", "缓冲", "报数", "两人", "一组", "一轮", "手型", "帮助"], "adapted_note": "游戏化巩固为教学建议"},
        {"stage": "小结评价", "demo_sec": 15, "pause_sec": 0,
         "script": "这节课我们把原地运球的手感练出来了，谁能说说手指哪里触球。对，指根和手指，掌心要空出来。眼睛不看球也能控住球的同学，请举手，非常棒。下节课我们学行进间运球，把球带着走，到时候今天的手感就派上用场了。好，收球做放松，抖抖手指和手腕。",
         "adapted_facts": ["放松", "手指", "一步", "慢速"], "adapted_note": "衔接下节课与收放器材安排"},
    ]
    draft = {
        "schema": ptd_core.SCHEMA_DRAFT,
        "id": view_entry["id"],
        "record_sha": view_entry["record_sha"],
        "segment": {
            "type": "practice",
            "meta": {
                "学段": "水平三（五至六年级，默认配置）",
                "片段位置": "完整无生试讲基本部分（复习片段）",
                "时长": "",
                "重点": "手指指根触球与按拍节奏",
                "器材": "篮球每人一个，检查球压与场地平整",
                "安全": "检查场地平整与器材完好，分组散点站位间隔一臂距离，听到口令立即停止练习，收球不推搡，注意运动负荷",
                "分层": "基础层低运球，提高层高低运球交替",
                "评价": "看触球手型与控球稳定性，两人互评",
            },
        },
        "config": {"segment_duration_sec": [120, 240], "speech_rate_chars_per_min": 230},
        "render": {
            "sport": "篮球",
            "chapter": "第三章 篮球运动教学内容 | 二、运球",
            "segment_name": "原地运球",
            "difficulty_display": view_entry["index_difficulty"] or "教材未标难度，按入门基础层处理",
            "figure": "图3-2-7 原地低运球、图3-2-8 原地高运球",
            "figure_images": [str(fig)],
            "cta": CYAN_LEGACY,
            "hashtags": "#教师编 #体育教师 #体育试讲 #试讲设计 #一次上岸",
        },
        "fields": {
            "difficulty": (
                {"kind": "index_stars", "display": view_entry["index_difficulty"],
                 "provenance": "textbook"}
                if view_entry["index_difficulty"]
                else {"kind": "index_empty_adapted", "display": "教材未标难度，按入门基础层处理",
                      "provenance": "adapted", "adapted_note": "索引难度为空，不虚构星级"}
            ),
            "method": {"text": method_text, "provenance": "textbook",
                       "evidence": [{"book_file": "人教版教师用书-篮球.md", "line": 1324,
                                     "excerpt": "> 两腿微屈，上体稍前倾，眼看前方；以肘关节为轴，前臂自然屈伸，五指张开，手 指和指根部位触球，通过手腕、手指柔和用力按压球的上方，做出随球、迎球动作。球 的落点控制在运球手同侧脚的外侧前方，反弹高度在膝关节(低运球)(图3-2-7)或 胸腰之间(高运球)(图3-2-8)。"}]},
            "rules": {"text": "降低重心抬头观察，不低头看球，另一侧手臂护球。",
                      "provenance": "adapted", "adapted_facts": ["护球", "重心"],
                      "adapted_note": "规则要点按方法要点改写为可判定表述"},
            "intent": {"text": "建立正确手指触球手型与按压节奏，体会高低运球差异。",
                       "provenance": "adapted", "adapted_facts": ["触球"],
                       "adapted_note": "意图由方法要点归纳"},
            "organization": {"text": "全班分组散点站位，每人一球，间隔一臂距离，教师巡回指导统一口令。",
                             "provenance": "adapted", "adapted_facts": ["一球", "间隔", "距离", "一臂"],
                             "adapted_note": "组织为教学建议"},
            "errors": {"rows": [
                {"error": {"text": "掌心按拍球", "provenance": "textbook",
                           "evidence": [{"book_file": "人教版教师用书-篮球.md", "line": 1348,
                                         "excerpt": "掌心按拍球"}]},
                 "fix": {"text": "体会手指和指根触球，掌心空出，再做一次给我看",
                         "provenance": "adapted", "adapted_facts": ["触球", "一次"],
                         "adapted_note": "纠正后加入再次检查为教学建议"}},
                {"error": {"text": "低头看球", "provenance": "adapted", "adapted_facts": ["看球"],
                           "adapted_note": "常见错误为教学经验"},
                 "fix": {"text": "手势报数游戏引导抬头，观察一次确认改进",
                         "provenance": "adapted", "adapted_facts": ["报数", "一次"],
                         "adapted_note": "纠正后加入再次检查为教学建议"}},
            ]},
        },
        "figures": [],
        "flow": [
            {"stage": st["stage"], "script": st["script"],
             "demo_sec": st["demo_sec"], "pause_sec": st["pause_sec"],
             "provenance": "adapted", "evidence": [], "adapted_facts": st["adapted_facts"],
             "adapted_note": st["adapted_note"]}
            for st in stages
        ],
        "source_view_entry": view_entry,
        "notes": {"human_rewrite_applied": True, "sport": "篮球"},
    }
    # 时长标注按统一口径回填（口播 + 示范停顿），保证与分段合计一致
    br = ptd_core.duration_breakdown(draft)
    draft["segment"]["meta"]["时长"] = (
        f"约{round(br['total_sec'])}秒（口播约{round(br['speech_sec'])}秒"
        f"＋示范停顿约{round(br['demo_pause_sec'])}秒）"
    )
    # 评审记录（与草稿同版本）
    review = {
        "schema": ptd_core.SCHEMA_REVIEW, "id": draft["id"],
        "draft_sha": ptd_core.draft_hash(draft), "reviewer": "agent（内容评审）",
        "checked": {k: True for k in ptd_core.REVIEW_CHECKS},
        "suggestion_notes": [
            {"where": "flow", "why": "口令与队形为教学建议，动作要点保持教材原文"},
            {"where": "fields.errors", "why": "再次检查为教学建议，纠正方法指向教材要点"},
        ],
        "verdict": "pass",
    }
    (ws / "scripts" / "review_trial_daily.json").write_text(
        json.dumps(review, ensure_ascii=False), encoding="utf-8"
    )
    (ws / "scripts" / "pending_trial_daily.json").write_text(
        json.dumps(draft, ensure_ascii=False), encoding="utf-8"
    )

    env = dict(os.environ)
    env["TRIAL_DAILY_WORKSPACE"] = str(ws)
    result = subprocess.run(
        [sys.executable, str(FILL_SCRIPT)],
        capture_output=True, text=True, env=env,
    )
    docx_path = ws / "desktop-attachments" / "2 体育试讲每日一练-帖子内容编辑模板.docx"
    doc = None
    if docx_path.exists():
        doc = Document(str(docx_path))
    _fill_cache = {
        "ws": ws, "result": result, "docx_path": docx_path, "doc": doc,
        "out": result.stdout + result.stderr, "pending": draft,
    }
    return _fill_cache


@case
def fill_pipeline_script_exit_green():
    fr = _fill_result()
    _assert(fr["result"].returncode == 0, f"退出码={fr['result'].returncode}\n{fr['out']}")
    _assert("✅ 全部通过" in fr["out"], "缺成功提示\n" + fr["out"])
    _assert("✅ pending_trial_daily.json 已删除" in fr["out"], "pending 未删除\n" + fr["out"])


@case
def fill_pipeline_docx_structure():
    fr = _fill_result()
    _assert(fr["doc"] is not None, "输出 DOCX 不存在")
    doc = fr["doc"]
    text = "\n".join(p.text or "" for p in doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                text += "\n" + (c.text or "")

    checks = {}
    checks["封面标题"] = "体育试讲设计" in text and "每日一练" in text
    checks["项目标签"] = "【篮球】练习环节" in text
    checks["环节名"] = "原地运球" in text
    checks["活动方法"] = "两腿微屈上体稍前倾" in text
    checks["规则"] = "降低重心抬头观察" in text
    checks["设计意图"] = "建立正确手指触球手型" in text
    checks["组织形式"] = "分组散点站位" in text
    checks["易犯错误表"] = "易犯错误" in text and "纠正方法" in text
    checks["试讲逐字稿"] = "试讲逐字稿" in text
    checks["引流段"] = CYAN_LEGACY in text
    checks["话题标签"] = "#教师编 #体育教师" in text
    img_count = sum(
        len(p._element.findall(".//" + qn("wp:inline"))) for p in doc.paragraphs
    )
    checks["图例图片"] = img_count == 1
    anchors = doc.element.body.findall(".//" + qn("wp:anchor"))
    checks["封面底层图"] = any(a.get("behindDoc") == "1" for a in anchors)
    checks["页眉水印"] = all(
        "PowerPlusWaterMarkObject" in s.header._element.xml for s in doc.sections
    )
    title_sizes = [
        r.font.size.pt for p in doc.paragraphs for r in p.runs
        if r.text.strip() in ("体育试讲设计", "每日一练") and r.font.size
    ]
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    title_sizes += [
                        r.font.size.pt for r in p.runs
                        if r.text.strip() in ("体育试讲设计", "每日一练") and r.font.size
                    ]
    checks["封面标题48pt"] = any(s >= 47 for s in title_sizes)
    titles = [i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "环节拆解"]
    checks["环节拆解另起一页"] = bool(titles) and bool(
        doc.paragraphs[titles[0]].paragraph_format.page_break_before
    )
    body = [p for p in doc.paragraphs if p.text.strip()]
    hs_p = body[-2] if body and body[-1].text.strip() == CYAN_LEGACY else None
    no_page_break = True
    if hs_p is not None:
        prev = hs_p._element.getprevious()
        if prev is not None and prev.tag == qn("w:p"):
            for br in prev.findall(".//" + qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    no_page_break = False
    checks["引流不另起页"] = no_page_break

    failed = [k for k, v in checks.items() if not v]
    _assert(not failed, f"未通过: {failed}")


# ===========================================================================
# 主入口
# ===========================================================================


def main():
    failed = 0
    for name, fn in CASES:
        try:
            fn()
        except Exception as exc:  # noqa: BLE001 - 用例隔离，吞异常会掩盖失败
            failed += 1
            print(f"  ❌ {name}: {exc}")
        else:
            print(f"  ✅ {name}")
    total = len(CASES)
    print(f"\n用例 {total} 个，通过 {total - failed}，失败 {failed}，跳过 0")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
