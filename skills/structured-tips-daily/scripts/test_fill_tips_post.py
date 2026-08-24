#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""test_fill_tips_post.py v2.1 -- 可信测试框架（冻结版）

测试对象：本 skill 的引擎 scripts/fill_tips_post.py（S 引擎），不是项目工作区里的副本。
真实项目 P 只作只读 fixture（模板来源），通过环境变量注入，测试全程只写临时副本。

用法：
    python3 test_fill_tips_post.py                     # 全量（18 用例）
    python3 test_fill_tips_post.py golden atomic       # 按关键字筛选用例

环境变量：
    TIPS_TEST_PROJECT_ROOT  真实项目根（默认 ~/Desktop/AI工作区/01-Projects/自媒体内容库-持续项目/结构化）

沙箱纪律（硬性）：
    所有 IMA 用例一律通过本地 fake stub（_fake_ima_api.cjs）执行，
    并强制覆盖子进程 env IMA_API_PATH，绝无真实网络调用、绝不创建真实笔记。

引擎契约（被测）：
    python3 fill_tips_post.py --project-root <DIR>     # 或 env TIPS_PROJECT_ROOT
    退出码：0=成功  1=环境/模板错误  2=pending JSON 无效（schema/语法/超长）
            3=写入后验证失败（含样式指纹漂移）
            4=pending 清理失败 -> 模板已自动回滚为原字节，pending 保留
    原子性：写入临时文件并验证通过后才替换模板；任何失败模板字节不变；失败保留 pending。

IMA 脚本契约（被测）：
    node upload_to_ima.js <md路径> <标题> [--fresh]
    stdout 最后一行输出结构化 JSON {"status","stage","note_id","media_id","reused","error"}
    退出码：0=成功（含 KB 降级；幂等复用也算 0） 1=用法/文件错误
            2=依赖缺失/IMA 笔记创建失败  3=笔记已创建但 KB 同步失败（重试复用 note_id 不重复建）
"""
import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPTS_DIR.parent
ENGINE = SCRIPTS_DIR / "fill_tips_post.py"
IMA_SCRIPT = SCRIPTS_DIR / "upload_to_ima.js"

REAL_PROJECT = Path(os.environ.get(
    "TIPS_TEST_PROJECT_ROOT",
    str(Path.home() / "Desktop/AI工作区/01-Projects/自媒体内容库-持续项目/结构化"),
))
TEMPLATE_REL = Path("desktop-attachments") / "3 结构化答题技巧-帖子内容编辑模板.docx"
REAL_TEMPLATE = REAL_PROJECT / TEMPLATE_REL

SCHEMA_KEYS = [
    "tip_title", "question_type", "tip_intro",
    "step1", "step2", "step3",
    "case_normal", "case_normal_note", "case_high", "case_high_note",
    "pitfalls_lead", "pitfalls", "tip_takeaway", "hashtags",
]
TITLE_MAX_CHARS = 20
PENDING_REL = Path("scripts") / "pending_tips.json"
SNAPSHOTS_REL = Path("scripts") / "_snapshots_tips"

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        h.update(f.read())
    return h.hexdigest()


def make_workspace(tag: str) -> Path:
    """临时项目根：目录骨架 + 真实模板只读拷贝。绝不写 REAL_PROJECT。"""
    ws = Path(tempfile.mkdtemp(prefix=f"tips_v21_{tag}_"))
    (ws / "scripts").mkdir(parents=True)
    (ws / TEMPLATE_REL.parent).mkdir(parents=True)
    shutil.copy(REAL_TEMPLATE, ws / TEMPLATE_REL)
    return ws


def write_pending(ws: Path, data: dict) -> None:
    (ws / PENDING_REL).write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")


def valid_pending() -> dict:
    return {
        "tip_title": "应急应变别再平均用力，先排序再出招",
        "question_type": "适用题型：应急应变题",
        "tip_intro": "破题角度：考官看的是轻重缓急的排序能力，不是面面俱到。",
        "step1": "第一步：先定优先级，把最紧急、最伤人的事排第一。",
        "step2": "第二步：止损优先，先把事态控制住，再谈调查和追责。",
        "step3": "第三步：处置加善后闭环，给结果也给交代。",
        "case_normal": "普通答法：遇到这种情况，我会先了解情况，再上报领导，然后安抚群众，最后总结经验。",
        "case_normal_note": "点评：平均用力，没有排序，考官看不到你的判断力。",
        "case_high": "高分答法：第一，先疏散人群、切断危险源；第二，立刻上报并同步联络消防医疗；第三，专人对接家属，事后复盘。",
        "case_high_note": "点评：有明确先后顺序，每个动作都有具体抓手。",
        "pitfalls_lead": "避坑提醒：",
        "pitfalls": "别一上来就上报领导当万能开头；别把安抚说得空泛。",
        "tip_takeaway": "说到底，应急应变衡量的不是话术，是你能不能分清轻重。",
        "hashtags": "#公考面试 #结构化面试 #应急应变 #答题技巧 #上岸",
    }


def run_engine(ws: Path) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, str(ENGINE), "--project-root", str(ws)],
        capture_output=True, text=True, timeout=120,
    )


# ── 测试独立实现的样式指纹（与引擎同构但独立编码，防共因失效） ──

def _para_fp(p_el) -> tuple:
    ppr = p_el.find(qn("w:pPr"))
    runs = []
    for r_el in p_el.findall(qn("w:r")):
        rpr = r_el.find(qn("w:rPr"))
        runs.append(rpr.xml if rpr is not None else "<none>")
    return (ppr.xml if ppr is not None else "<none>", tuple(runs))


def style_fingerprint(docx_path: Path) -> list:
    """逐段 (pPr XML, 各 run rPr XML)，含封面文本框内段落。写入前后必须零漂移。"""
    doc = Document(str(docx_path))
    fps = []
    for p in doc.paragraphs:
        el = p._element
        fps.append(_para_fp(el))
        for tx in el.iter(qn("w:txbxContent")):
            for tp in tx.iter(qn("w:p")):
                fps.append(_para_fp(tp))
    return fps


def check_docx_contract(out_docx: Path, pending: dict) -> list:
    """18 段/4 图/样式/内容契约断言（结构真值与模板实测一致）。"""
    errors = []
    doc = Document(str(out_docx))
    paras = doc.paragraphs
    if len(paras) != 18:
        errors.append(f"段数 {len(paras)} != 18")
        return errors
    img_total = sum(len(p._element.findall(".//" + qn("w:drawing"))) for p in paras)
    if img_total != 4:
        errors.append(f"图片总数 {img_total} != 4")
    # 段[0] 封面文本框：2 镜像 + 前缀 + 新标题
    prefixes, titles = [], []
    for tx in paras[0]._element.iter(qn("w:txbxContent")):
        ts = [t.text for t in tx.findall(".//" + qn("w:t"))]
        if len(ts) >= 2:
            prefixes.append(ts[0])
            titles.append(ts[1])
    if prefixes != ["结构化答题技巧：", "结构化答题技巧："]:
        errors.append(f"封面前缀错误: {prefixes}")
    if not titles or not all(t == pending["tip_title"] for t in titles):
        errors.append(f"封面标题未替换: {titles}")
    # 段[7] 怎么答？Heading 2 且无分页
    p7 = paras[7]
    if p7.style.name != "Heading 2" or "怎么答" not in p7.text:
        errors.append(f"段[7] 非怎么答蓝色标题: {p7.text!r}/{p7.style.name}")
    p7_pPr = p7._element.find(qn("w:pPr"))
    if p7_pPr is not None and p7_pPr.find(qn("w:pageBreakBefore")) is not None:
        errors.append("段[7] 不应有 pageBreakBefore")
    # 段[8]/[10] emoji 前缀
    if not paras[8].text.startswith("🙅"):
        errors.append(f"段[8] 缺 🙅 前缀: {paras[8].text[:12]!r}")
    if not paras[10].text.startswith("👍"):
        errors.append(f"段[10] 缺 👍 前缀: {paras[10].text[:12]!r}")
    # 段[16] 引流段样式
    p16, r16 = paras[16], paras[16].runs[0] if paras[16].runs else None
    if r16 is None or not (p16.alignment == 1 and r16.bold is True and str(r16.font.color.rgb) == "85120F"):
        errors.append("段[16] 引流段样式丢失（居中/加粗/#85120F）")
    if "关注我" not in paras[16].text:
        errors.append("段[16] 引流段原文被改动")
    # 内容全命中
    full = "\n".join(p.text or "" for p in paras)
    for key in ["question_type", "tip_intro", "step1", "step3", "case_normal",
                "case_normal_note", "case_high", "case_high_note",
                "pitfalls", "tip_takeaway", "hashtags"]:
        if pending[key] not in full:
            errors.append(f"字段未写入: {key}")
    return errors


# ── IMA 沙箱：fake stub + 强制 env，绝无真实调用 ──

FAKE_STUB_SRC = r"""// 测试专用 fake stub：只写本地日志，绝无网络调用
const fs = require('fs');
const MODE = process.env.FAKE_IMA_MODE || 'ok';
const LOG = process.env.FAKE_IMA_LOG || '/dev/null';
function log(url) { try { fs.appendFileSync(LOG, url + '\n'); } catch (e) {} }
async function imaApi(url, body) {
  log(url);
  if (url === 'openapi/note/v1/import_doc') {
    if (MODE === 'note_fail') return JSON.stringify({ code: 1, msg: 'fake: note create failed' });
    return JSON.stringify({ code: 0, data: { note_id: 'FAKE_NOTE_001' } });
  }
  if (url === 'openapi/wiki/v1/search_knowledge_base') {
    if (MODE === 'kb_search_empty') return JSON.stringify({ code: 0, data: { info_list: [] } });
    return JSON.stringify({ code: 0, data: { info_list: [{ kb_name: '总分总', kb_id: 'KB001' }] } });
  }
  if (url === 'openapi/wiki/v1/search_knowledge') {
    return JSON.stringify({ code: 0, data: { info_list: [{ title: '00_结构化考官思维', media_id: 'FOLDER001' }] } });
  }
  if (url === 'openapi/wiki/v1/add_knowledge') {
    return JSON.stringify({ code: 0, data: { media_id: 'MEDIA001' } });
  }
  return JSON.stringify({ code: -1, msg: 'fake: unexpected url ' + url });
}
module.exports = { imaApi };
"""


def write_fake_ima_api(ws: Path) -> Path:
    stub = ws / "_fake_ima_api.cjs"
    stub.write_text(FAKE_STUB_SRC, encoding="utf-8")
    return stub


def sandbox_env(ws: Path, mode: str = "ok") -> dict:
    """强制 IMA_API_PATH 指向本地 fake stub——这是杜绝真实调用的硬闸。"""
    log = ws / "ima_calls.log"
    return {
        **os.environ,
        "IMA_API_PATH": str(ws / "_fake_ima_api.cjs"),
        "FAKE_IMA_LOG": str(log),
        "FAKE_IMA_MODE": mode,
    }


def run_ima(md: Path, title: str, env: dict, extra_args: list = None) -> subprocess.CompletedProcess:
    args = ["node", str(IMA_SCRIPT), str(md), title] + (extra_args or [])
    return subprocess.run(args, capture_output=True, text=True, timeout=60, env=env)


def last_json(stdout: str) -> dict:
    lines = [ln for ln in stdout.strip().split("\n") if ln.strip().startswith("{")]
    return json.loads(lines[-1]) if lines else {}


def count_calls(log: Path, kw: str) -> int:
    if not log.exists():
        return 0
    return sum(1 for ln in log.read_text(encoding="utf-8").splitlines() if kw in ln)


# ============================ 用例 ============================

def case_golden(ws):
    """TP1 黄金路径：合法 14 字段 -> exit 0，全契约命中，pending 清理，快照生成，样式指纹零漂移。"""
    orig_copy = ws / "_test_orig_tpl.docx"
    shutil.copy(ws / TEMPLATE_REL, orig_copy)  # 写入前留存（此时模板 = 干净模板）
    pending = valid_pending()
    write_pending(ws, pending)
    tpl_sha_before = sha256(ws / TEMPLATE_REL)
    r = run_engine(ws)
    assert r.returncode == 0, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    assert "全部验证通过" in r.stdout, r.stdout
    errs = check_docx_contract(ws / TEMPLATE_REL, pending)
    assert not errs, errs
    assert not (ws / PENDING_REL).exists(), "成功后 pending 应被清理"
    snaps = list((ws / SNAPSHOTS_REL).glob("snapshot_*.docx"))
    assert len(snaps) == 1, f"应生成 1 个快照，实际 {len(snaps)}"
    assert sha256(snaps[0]) == tpl_sha_before, "快照应等于写入前模板"
    # 独立实现复核：产出文档与写入前模板的样式指纹必须完全一致
    assert style_fingerprint(orig_copy) == style_fingerprint(ws / TEMPLATE_REL), \
        "样式指纹漂移（独立实现复核）：字体/字号/颜色/对齐等 pPr/rPr 被改动"


def case_missing_field(ws):
    """缺字段 -> exit 2，模板字节不变，pending 保留。"""
    pending = valid_pending()
    del pending["case_high_note"]
    write_pending(ws, pending)
    before = sha256(ws / TEMPLATE_REL)
    r = run_engine(ws)
    assert r.returncode == 2, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    assert "case_high_note" in r.stdout + r.stderr, "错误信息应指出缺失字段名"
    assert sha256(ws / TEMPLATE_REL) == before, "失败时模板字节必须不变"
    assert (ws / PENDING_REL).exists(), "失败时 pending 应保留"


def case_extra_field(ws):
    """多余第 15 字段 -> exit 2，模板不变。"""
    pending = valid_pending()
    pending["extra_key"] = "多余字段"
    write_pending(ws, pending)
    before = sha256(ws / TEMPLATE_REL)
    r = run_engine(ws)
    assert r.returncode == 2, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    assert "extra_key" in r.stdout + r.stderr, "错误信息应指出多余字段名"
    assert sha256(ws / TEMPLATE_REL) == before


def case_oversized_title(ws):
    """标题超 20 字 -> exit 2，模板不变，写入前拒绝。"""
    pending = valid_pending()
    pending["tip_title"] = "这是一个超过二十个字的超长标题用来测试引擎是否会拒绝写入"
    write_pending(ws, pending)
    before = sha256(ws / TEMPLATE_REL)
    r = run_engine(ws)
    assert r.returncode == 2, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    assert "20" in r.stdout + r.stderr, "错误信息应说明 20 字上限"
    assert sha256(ws / TEMPLATE_REL) == before


def case_bad_json(ws):
    """JSON 语法坏 -> exit 2，模板不变。"""
    (ws / PENDING_REL).write_text('{"tip_title": "坏JSON', encoding="utf-8")
    before = sha256(ws / TEMPLATE_REL)
    r = run_engine(ws)
    assert r.returncode == 2, f"exit={r.returncode}"
    assert sha256(ws / TEMPLATE_REL) == before


def case_corrupted_template(ws):
    """模板损坏 -> exit 1 清晰报错，pending 保留，无 Python traceback。"""
    write_pending(ws, valid_pending())
    (ws / TEMPLATE_REL).write_bytes(b"not a docx file at all")
    r = run_engine(ws)
    assert r.returncode == 1, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    assert "Traceback" not in r.stderr, "损坏模板不应裸抛 traceback"
    assert (ws / PENDING_REL).exists(), "环境失败时 pending 应保留"


def case_no_pending(ws):
    """无 pending -> exit 1，模板不变。"""
    before = sha256(ws / TEMPLATE_REL)
    r = run_engine(ws)
    assert r.returncode == 1, f"exit={r.returncode}"
    assert sha256(ws / TEMPLATE_REL) == before


def case_missing_template(ws):
    """模板不存在 -> exit 1，不崩溃。"""
    write_pending(ws, valid_pending())
    (ws / TEMPLATE_REL).unlink()
    r = run_engine(ws)
    assert r.returncode == 1, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"


def case_env_project_root(ws):
    """env TIPS_PROJECT_ROOT 与 --project-root 等效（可迁移性）。"""
    write_pending(ws, valid_pending())
    r = subprocess.run(
        [sys.executable, str(ENGINE)],
        capture_output=True, text=True, timeout=120,
        env={**os.environ, "TIPS_PROJECT_ROOT": str(ws)},
    )
    assert r.returncode == 0, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"


def case_style_tamper_detected(ws):
    """样式损坏必须被拦截（不再假绿）：人为改字号 -> validate_doc 报指纹漂移；纯文本写入不误报。"""
    import importlib.util
    from docx.shared import Pt
    sys.dont_write_bytecode = True  # 不在 skill 目录留 __pycache__ 噪音
    spec = importlib.util.spec_from_file_location("fill_engine", ENGINE)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)

    pending = valid_pending()
    # 破坏组：正常写入后人为改字号（模拟"文本对了但样式坏了"）
    doc = Document(str(ws / TEMPLATE_REL))
    fp_before = mod.capture_style_fingerprint(doc)
    mod.fill_document(doc, pending)
    doc.paragraphs[5].runs[0].font.size = Pt(99)
    tampered = ws / "_tampered.docx"
    doc.save(str(tampered))
    errs = mod.validate_doc(tampered, pending, fp_before)
    assert any("样式指纹漂移" in e for e in errs), f"样式损坏未被拦截: {errs}"
    assert any("字段未写入" not in e and "样式指纹漂移" in e for e in errs), errs

    # 对照组：仅文本写入、不动样式 -> 验证零错误（不误伤正常路径）
    doc2 = Document(str(ws / TEMPLATE_REL))
    fp2 = mod.capture_style_fingerprint(doc2)
    mod.fill_document(doc2, pending)
    clean = ws / "_clean.docx"
    doc2.save(str(clean))
    errs2 = mod.validate_doc(clean, pending, fp2)
    assert not errs2, f"纯文本写入被误报: {errs2}"


def case_pending_cleanup_failure(ws):
    """pending 清理失败 -> exit 4 且模板回滚原字节（不允许留下\"已变模板\"中间态）。"""
    pending = valid_pending()
    write_pending(ws, pending)
    orig = sha256(ws / TEMPLATE_REL)
    p = ws / PENDING_REL
    if sys.platform == "darwin":
        lock, unlock = ["chflags", "uchg", str(p)], ["chflags", "nouchg", str(p)]
    else:
        lock, unlock = ["chattr", "+i", str(p)], ["chattr", "-i", str(p)]
    try:
        rv = subprocess.run(lock, capture_output=True, text=True, timeout=30)
        if rv.returncode != 0:
            return f"skip: cannot lock file ({(rv.stderr or '').strip()[:60]})"
        r = run_engine(ws)
        assert r.returncode == 4, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
        assert "ROLLBACK" in r.stdout, f"应打印回滚说明: {r.stdout}"
        assert sha256(ws / TEMPLATE_REL) == orig, "exit 4 时模板必须回滚为写入前字节"
        assert p.exists(), "锁定中的 pending 应保留（可重试）"
    finally:
        subprocess.run(unlock, capture_output=True, timeout=30)


def case_ima_usage_error(ws):
    """IMA 无参数 -> exit 1（沙箱 env 强制注入，绝不触网）。"""
    if shutil.which("node") is None:
        return "skip: no node"
    write_fake_ima_api(ws)
    r = subprocess.run(["node", str(IMA_SCRIPT)], capture_output=True, text=True,
                       timeout=60, env=sandbox_env(ws))
    assert r.returncode == 1, f"exit={r.returncode}"
    assert last_json(r.stdout).get("status") == "error", r.stdout


def case_ima_missing_file(ws):
    """IMA md 文件不存在 -> exit 1，不抛未捕获异常。"""
    if shutil.which("node") is None:
        return "skip: no node"
    write_fake_ima_api(ws)
    r = run_ima(ws / "nonexistent.md", "标题", sandbox_env(ws))
    assert r.returncode == 1, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    j = last_json(r.stdout)
    assert j.get("status") == "error", f"应有结构化错误而非裸异常: {r.stdout!r}"


def case_ima_missing_dependency(ws):
    """IMA 依赖缺失 -> exit 2 + 结构化 JSON 输出（失败可诊断）。"""
    if shutil.which("node") is None:
        return "skip: no node"
    md = ws / "note.md"
    md.write_text("# t", encoding="utf-8")
    r = subprocess.run(
        ["node", str(IMA_SCRIPT), str(md), "标题"],
        capture_output=True, text=True, timeout=60,
        env={**os.environ, "IMA_API_PATH": "/nonexistent/ima_api.cjs"},
    )
    assert r.returncode == 2, f"exit={r.returncode}\n{r.stdout}\n{r.stderr}"
    assert last_json(r.stdout).get("status") == "error", f"最后一行应为结构化 JSON: {r.stdout!r}"


def case_ima_note_fail(ws):
    """IMA 笔记创建失败 -> exit 2 + 结构化 JSON；两次失败均无状态文件残留。"""
    if shutil.which("node") is None:
        return "skip: no node"
    write_fake_ima_api(ws)
    md = ws / "note.md"
    md.write_text("# t", encoding="utf-8")
    env = sandbox_env(ws, mode="note_fail")
    r1 = run_ima(md, "标题NF", env)
    assert r1.returncode == 2, f"exit={r1.returncode}\n{r1.stdout}\n{r1.stderr}"
    j1 = last_json(r1.stdout)
    assert j1.get("stage") == "note" and j1.get("status") == "error", j1
    assert not (ws / ".ima_upload_state.json").exists(), "创建失败不应落状态文件"
    r2 = run_ima(md, "标题NF", env)
    assert r2.returncode == 2, f"重试 exit={r2.returncode}"
    assert count_calls(ws / "ima_calls.log", "import_doc") == 2, "两次尝试各调一次 import_doc"


def case_ima_kb_fail_retry_no_dup(ws):
    """KB 同步失败后重试 -> 复用 note_id，绝不重复调 import_doc（核心幂等契约）。"""
    if shutil.which("node") is None:
        return "skip: no node"
    write_fake_ima_api(ws)
    md = ws / "note.md"
    md.write_text("# t", encoding="utf-8")
    env = sandbox_env(ws, mode="kb_search_empty")
    # 第一次：笔记创建成功，KB 搜索为空 -> exit 3（部分成功）
    r1 = run_ima(md, "标题KB", env)
    assert r1.returncode == 3, f"exit={r1.returncode}\n{r1.stdout}\n{r1.stderr}"
    j1 = last_json(r1.stdout)
    assert j1.get("note_id") == "FAKE_NOTE_001", j1
    state = json.loads((ws / ".ima_upload_state.json").read_text(encoding="utf-8"))
    entry = list(state.values())[0]
    assert entry["note_id"] == "FAKE_NOTE_001" and entry["kb_done"] is False, state
    # 第二次重试：复用已建笔记 -> 仍 exit 3，但 import_doc 调用数不增长
    r2 = run_ima(md, "标题KB", env)
    assert r2.returncode == 3, f"重试 exit={r2.returncode}\n{r2.stdout}"
    assert "复用" in r2.stdout, f"重试应说明复用 note_id: {r2.stdout}"
    assert count_calls(ws / "ima_calls.log", "import_doc") == 1, \
        f"重试不得重复建笔记: import_doc 被调 {count_calls(ws / 'ima_calls.log', 'import_doc')} 次"


def case_ima_ok_reused(ws):
    """完整成功后重试 -> exit 0 + reused:true，且零新增 API 调用。"""
    if shutil.which("node") is None:
        return "skip: no node"
    write_fake_ima_api(ws)
    md = ws / "note.md"
    md.write_text("# t", encoding="utf-8")
    env = sandbox_env(ws, mode="ok")
    r1 = run_ima(md, "标题OK", env)
    assert r1.returncode == 0, f"exit={r1.returncode}\n{r1.stdout}\n{r1.stderr}"
    j1 = last_json(r1.stdout)
    assert j1.get("status") == "ok" and j1.get("note_id") == "FAKE_NOTE_001" and not j1.get("reused"), j1
    assert count_calls(ws / "ima_calls.log", "import_doc") == 1
    # 第二次：命中 kb_done:true 状态 -> 零 API 调用直接成功
    r2 = run_ima(md, "标题OK", env)
    assert r2.returncode == 0, f"重试 exit={r2.returncode}\n{r2.stdout}"
    j2 = last_json(r2.stdout)
    assert j2.get("reused") is True and j2.get("status") == "ok", j2
    assert count_calls(ws / "ima_calls.log", "import_doc") == 1, "幂等复用不得重复建笔记"


def case_ima_fresh(ws):
    """--fresh 忽略历史状态强制新建（import_doc 次数 +1）。"""
    if shutil.which("node") is None:
        return "skip: no node"
    write_fake_ima_api(ws)
    md = ws / "note.md"
    md.write_text("# t", encoding="utf-8")
    env = sandbox_env(ws, mode="ok")
    r1 = run_ima(md, "标题FR", env)
    assert r1.returncode == 0, f"exit={r1.returncode}\n{r1.stdout}\n{r1.stderr}"
    r2 = run_ima(md, "标题FR", env, extra_args=["--fresh"])
    assert r2.returncode == 0, f"exit={r2.returncode}\n{r2.stdout}\n{r2.stderr}"
    assert count_calls(ws / "ima_calls.log", "import_doc") == 2, "--fresh 应强制新建一次"


CASES = [
    ("golden", case_golden),
    ("missing_field", case_missing_field),
    ("extra_field", case_extra_field),
    ("oversized_title", case_oversized_title),
    ("bad_json", case_bad_json),
    ("corrupted_template", case_corrupted_template),
    ("no_pending", case_no_pending),
    ("missing_template", case_missing_template),
    ("env_project_root", case_env_project_root),
    ("style_tamper_detected", case_style_tamper_detected),
    ("pending_cleanup_failure", case_pending_cleanup_failure),
    ("ima_usage_error", case_ima_usage_error),
    ("ima_missing_file", case_ima_missing_file),
    ("ima_missing_dependency", case_ima_missing_dependency),
    ("ima_note_fail", case_ima_note_fail),
    ("ima_kb_fail_retry_no_dup", case_ima_kb_fail_retry_no_dup),
    ("ima_ok_reused", case_ima_ok_reused),
    ("ima_fresh", case_ima_fresh),
]


def main() -> int:
    if not HAS_DOCX:
        print("FAIL: 需要 python-docx（pip3 install python-docx）")
        return 1
    if not REAL_TEMPLATE.exists():
        print(f"FAIL: 只读 fixture 模板不存在: {REAL_TEMPLATE}")
        return 1

    filters = sys.argv[1:]
    selected = [(n, f) for n, f in CASES if not filters or any(k in n for k in filters)]
    if not selected:
        print(f"FAIL: 筛选 {filters} 未匹配到任何用例（0 用例不得假绿退出 0）")
        return 1
    real_tpl_sha = sha256(REAL_TEMPLATE)  # 全局安全哨兵：真实项目模板必须零改动

    passed, failed, skipped = 0, [], 0
    for name, fn in selected:
        ws = make_workspace(name)
        try:
            note = fn(ws)
            if note and note.startswith("skip"):
                print(f"  ⏭  {name} ({note})")
                skipped += 1
            else:
                print(f"  ✅ {name}")
                passed += 1
        except AssertionError as e:
            print(f"  ❌ {name}\n      {e}")
            failed.append(name)
        except Exception as e:  # noqa: BLE001
            print(f"  ❌ {name} [异常] {type(e).__name__}: {e}")
            failed.append(name)
        finally:
            shutil.rmtree(ws, ignore_errors=True)

    if sha256(REAL_TEMPLATE) != real_tpl_sha:
        print("🔴 安全哨兵触发：真实项目模板被改动！")
        failed.append("SAFETY_SENTINEL")

    print(f"\n结果: {passed} 通过 / {len(failed)} 失败 / {skipped} 跳过（共 {len(selected)} 例）")
    if failed:
        print(f"失败用例: {failed}")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
