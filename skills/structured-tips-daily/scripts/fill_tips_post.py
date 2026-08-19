#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""fill_tips_post.py v2.1 -- 结构化答题技巧·每日一练 帖子生成引擎

用法:
    python3 fill_tips_post.py [--project-root DIR]
    项目根解析优先级: --project-root > 环境变量 TIPS_PROJECT_ROOT > 脚本所在目录上一级（随项目安装）

项目根内路径布局:
    desktop-attachments/3 结构化答题技巧-帖子内容编辑模板.docx  模板（成功时原子替换）
    scripts/pending_tips.json   待写入内容（成功后删除，任何失败保留）
    scripts/_snapshots_tips/    写入前快照（最多保留 10 个）

pending JSON schema（硬性，恰好 14 个顶层 key，全部非空字符串）:
    tip_title(≤20字) question_type tip_intro step1 step2 step3
    case_normal case_normal_note case_high case_high_note
    pitfalls_lead pitfalls tip_takeaway hashtags

退出码:
    0 成功
    1 环境错误（pending 缺失/模板缺失或损坏）
    2 pending JSON 无效
    3 写入后验证失败（含样式指纹不一致）
    4 模板已替换但 pending 清理失败 -> 模板已自动回滚为原字节，pending 保留

安全设计:
    1) schema 全量校验先于任何写盘
    2) 先写临时文件并验证（结构 + 内容 + 样式指纹），通过后 os.replace 原子替换
    3) 快照在写入前生成；任何失败（含 pending 清理失败）模板字节不变
    4) 样式指纹 = 逐段 pPr/rPr XML 对比（含封面文本框），文本可换、样式零漂移
"""

import argparse
import json
import os
import shutil
import sys
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

TEMPLATE_REL = Path("desktop-attachments") / "3 结构化答题技巧-帖子内容编辑模板.docx"
PENDING_REL = Path("scripts") / "pending_tips.json"
SNAPSHOTS_REL = Path("scripts") / "_snapshots_tips"
MAX_SNAPSHOTS = 10

SCHEMA_KEYS = [
    "tip_title", "question_type", "tip_intro",
    "step1", "step2", "step3",
    "case_normal", "case_normal_note", "case_high", "case_high_note",
    "pitfalls_lead", "pitfalls", "tip_takeaway", "hashtags",
]
TITLE_MAX_CHARS = 20

COVER_PREFIX = "结构化答题技巧："  # 段[0] 文本框前缀（2 镜像同步）
DRAIN_COLOR = "85120F"           # 引流段颜色（段[16] 固定）

# 段位映射（技巧教学型；段[7]="怎么答？"蓝色标题不动；段[16] 引流段不动；段[17] 末图不动）
REPLACE_MAP = [
    (2, "question_type"),
    (3, "tip_intro"),
    (4, "step1"),
    (5, "step2"),
    (6, "step3"),
    (8, "case_normal"),
    (9, "case_normal_note"),
    (10, "case_high"),
    (11, "case_high_note"),
    (12, "pitfalls_lead"),
    (13, "pitfalls"),
    (14, "tip_takeaway"),
    (15, "hashtags"),
]
EMOJI_PREFIX = {8: "🙅\u200d♂️", 10: "👍"}


def fail(code: int, *msgs) -> None:
    for m in msgs:
        print(m)
    sys.exit(code)


def resolve_project_root(cli_value: str) -> Path:
    if cli_value:
        root = Path(cli_value).expanduser()
    elif os.environ.get("TIPS_PROJECT_ROOT"):
        root = Path(os.environ["TIPS_PROJECT_ROOT"]).expanduser()
    else:
        root = Path(__file__).resolve().parent.parent  # 随项目安装: <root>/scripts/本脚本
    if not root.is_dir():
        fail(1, f"[ERROR] 项目根不存在: {root}")
    return root


def load_pending(path: Path) -> dict:
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError as e:
        fail(2, f"[ERROR] pending JSON 不是有效 UTF-8: {e}", "        请以 UTF-8 无 BOM 重新保存")
    try:
        data = json.loads(text)
    except json.JSONDecodeError as e:
        fail(2, f"[ERROR] pending JSON 语法错误: {e}")
    if not isinstance(data, dict):
        fail(2, f"[ERROR] pending JSON 顶层必须是对象，实际: {type(data).__name__}")

    keys = set(data.keys())
    expected = set(SCHEMA_KEYS)
    missing, extra = sorted(expected - keys), sorted(keys - expected)
    if missing:
        fail(2, f"[ERROR] 缺少 {len(missing)} 个必填字段: {', '.join(missing)}",
             f"        恰好需要 {len(SCHEMA_KEYS)} 个顶层 key: {', '.join(SCHEMA_KEYS)}")
    if extra:
        fail(2, f"[ERROR] 发现 {len(extra)} 个多余字段: {', '.join(extra)}",
             f"        恰好需要 {len(SCHEMA_KEYS)} 个顶层 key，多余字段一律拒绝")
    for k in SCHEMA_KEYS:
        v = data[k]
        if not isinstance(v, str) or not v.strip():
            fail(2, f"[ERROR] 字段 {k} 必须为非空字符串，实际: {v!r}")
    if len(data["tip_title"]) > TITLE_MAX_CHARS:
        fail(2, f"[ERROR] tip_title 超长: {len(data['tip_title'])} 字（上限 {TITLE_MAX_CHARS} 字）: {data['tip_title']!r}")
    return data


def take_snapshot(root: Path) -> Path | None:
    snap_dir = root / SNAPSHOTS_REL
    snap_dir.mkdir(parents=True, exist_ok=True)
    tpl = root / TEMPLATE_REL
    if not tpl.exists():
        return None
    ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    snap = snap_dir / f"snapshot_{ts}.docx"
    shutil.copy(tpl, snap)
    for old in sorted(snap_dir.glob("snapshot_*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)[MAX_SNAPSHOTS:]:
        old.unlink()
    return snap


def replace_textbox_question(paragraph, new_text: str) -> int:
    """替换段内所有文本框第 2 个 <w:t>（封面大标题），保留第 1 个前缀。"""
    modified = 0
    for txbx in paragraph._element.iter(qn("w:txbxContent")):
        runs = txbx.findall(".//" + qn("w:t"))
        if len(runs) >= 2:
            runs[1].text = new_text
            modified += 1
        elif len(runs) == 1:
            runs[0].text = new_text
            modified += 1
    return modified


def set_textbox_prefix(paragraph, prefix: str) -> None:
    for txbx in paragraph._element.iter(qn("w:txbxContent")):
        runs = txbx.findall(".//" + qn("w:t"))
        if runs:
            runs[0].text = prefix


def replace_run_text_safely(paragraph, new_text: str) -> None:
    """找第一个纯文本 run（无 drawing）写入，图片 run 完全不动，其余文本 run 清空。"""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    target = None
    for r in runs:
        has_t = bool(r._element.findall(".//" + qn("w:t")))
        has_d = bool(r._element.findall(".//" + qn("w:drawing")))
        if has_t and not has_d:
            target = r
            break
    if target is None:
        paragraph.add_run(new_text)
        return
    target.text = new_text
    for r in runs:
        if r is target:
            continue
        if not r._element.findall(".//" + qn("w:drawing")):
            r.text = ""


def fill_document(doc: Document, content: dict) -> int:
    """按段位映射写入全部 14 字段，返回封面文本框修改数（期望 2）。"""
    set_textbox_prefix(doc.paragraphs[0], COVER_PREFIX)
    modified = replace_textbox_question(doc.paragraphs[0], content["tip_title"])
    for idx, key in REPLACE_MAP:
        text = content[key]
        if idx in EMOJI_PREFIX:
            text = EMOJI_PREFIX[idx] + text
        replace_run_text_safely(doc.paragraphs[idx], text)
    return modified


def capture_style_fingerprint(doc: Document) -> list:
    """逐段捕获样式指纹: 每段 (pPr XML, 各 run rPr XML 元组)，含封面文本框内段落。

    写入只允许改 <w:t> 文本；任何 pPr/rPr 变化（字体/字号/颜色/加粗/对齐等）
    都会造成指纹不一致，据此拦截"文本对了但样式坏了"的假绿输出。
    """
    def para_fp(p_el) -> tuple:
        ppr = p_el.find(qn("w:pPr"))
        run_fps = []
        for r_el in p_el.findall(qn("w:r")):
            rpr = r_el.find(qn("w:rPr"))
            run_fps.append(rpr.xml if rpr is not None else "<none>")
        return (ppr.xml if ppr is not None else "<none>", tuple(run_fps))

    fp = []
    for p in doc.paragraphs:
        el = p._element
        fp.append(para_fp(el))
        for txbx in el.iter(qn("w:txbxContent")):
            for tp in txbx.iter(qn("w:p")):
                fp.append(para_fp(tp))
    return fp


def validate_doc(path: Path, content: dict, fp_before: list) -> list:
    """结构契约 + 内容命中 + 样式指纹三重验证（对临时文件执行，通过才替换模板）。"""
    errors = []
    doc = Document(str(path))
    paras = doc.paragraphs

    if len(paras) != 18:
        errors.append(f"段数异常: {len(paras)} (期望 18)")
        return errors
    img_total = sum(len(p._element.findall(".//" + qn("w:drawing"))) for p in paras)
    if img_total != 4:
        errors.append(f"图片总数: {img_total} (期望 4)")

    # 段[0]: 2 个文本框 + 前缀 + 新标题镜像
    txbx_n = 0
    for txbx in paras[0]._element.iter(qn("w:txbxContent")):
        txbx_n += 1
        ts = txbx.findall(".//" + qn("w:t"))
        if len(ts) < 2:
            errors.append(f"段[0] 文本框结构异常: <w:t> 数量 {len(ts)}")
            continue
        if ts[0].text != COVER_PREFIX:
            errors.append(f"段[0] 前缀被改动: {ts[0].text!r}")
        if ts[1].text != content["tip_title"]:
            errors.append("段[0] 封面标题未生效")
    if txbx_n != 2:
        errors.append(f"段[0] 文本框数: {txbx_n} (期望 2)")

    # 段[7]: 怎么答？Heading 2 且无分页
    p7 = paras[7]
    if p7.style.name != "Heading 2" or "怎么答" not in p7.text:
        errors.append(f"段[7] 应为'怎么答？'蓝色标题: {p7.text!r} / {p7.style.name}")
    pPr = p7._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
        errors.append("段[7] 不应有 pageBreakBefore")

    # 段[16]: 引流段样式
    p16 = paras[16]
    r16 = p16.runs[0] if p16.runs else None
    if r16 is None or not (p16.alignment == 1 and r16.bold is True and str(r16.font.color.rgb) == DRAIN_COLOR):
        errors.append("引流段样式丢失（段[16] 须居中 + 加粗 + #85120F）")
    if "关注我" not in (p16.text or ""):
        errors.append("段[16] 引流段原文被改动")

    # 内容命中: 14 字段全部落盘
    for idx, key in REPLACE_MAP:
        expect = EMOJI_PREFIX.get(idx, "") + content[key]
        if expect not in (paras[idx].text or ""):
            errors.append(f"字段未写入或被改动: {key} (段[{idx}])")

    # 样式指纹: 填充前后逐段 pPr/rPr 必须零漂移
    fp_after = capture_style_fingerprint(doc)
    if len(fp_after) != len(fp_before):
        errors.append(f"样式指纹条目数变化: {len(fp_after)} (期望 {len(fp_before)})")
    else:
        for i, (a, b) in enumerate(zip(fp_after, fp_before)):
            if a != b:
                errors.append(f"样式指纹漂移 @条目{i}: pPr 或 run rPr 与模板不一致（字体/字号/颜色/对齐被改动）")
                break
    return errors


def rollback_template(tpl: Path, original_bytes: bytes) -> bool:
    """把模板字节回滚为写入前内容（原子写回）。"""
    try:
        rb = tpl.with_name(tpl.name + ".rollback.tmp")
        rb.write_bytes(original_bytes)
        os.replace(rb, tpl)
        return True
    except OSError:
        return False


def main() -> None:
    ap = argparse.ArgumentParser(description="结构化答题技巧·每日一练 帖子生成引擎")
    ap.add_argument("--project-root", help="项目根目录（默认: TIPS_PROJECT_ROOT 环境变量，其次脚本上一级）")
    args = ap.parse_args()

    root = resolve_project_root(args.project_root)
    tpl = root / TEMPLATE_REL
    pending = root / PENDING_REL

    if not pending.exists():
        fail(1, f"[ERROR] 待写入文件不存在: {pending}", "        请先把新技巧稿写到 pending_tips.json 再运行本脚本")
    if not tpl.exists():
        fail(1, f"[ERROR] 模板不存在: {tpl}")

    content = load_pending(pending)
    print(f"[1/5] schema 校验通过: {len(content)} 字段 / 标题 {len(content['tip_title'])} 字")

    original_bytes = tpl.read_bytes()  # 任何后续失败以此回滚
    try:
        doc = Document(str(tpl))
    except Exception as e:  # noqa: BLE001 模板损坏/非 docx 一律走环境错误
        fail(1, f"[ERROR] 模板无法解析（已损坏或不是 docx）: {tpl}", f"        {type(e).__name__}: {e}",
             "        模板未被改动；请从快照或备份恢复模板后重试")
    fp_before = capture_style_fingerprint(doc)

    snap = take_snapshot(root)
    print(f"[2/5] 快照备份: {snap.name if snap else '无（模板不存在）'}")

    print("[3/5] 写入 14 字段（段位映射 + 封面双镜像）")
    modified = fill_document(doc, content)

    tmp_fd, tmp_name = tempfile.mkstemp(suffix=".docx", prefix=".fill_tips_tmp_", dir=str(tpl.parent))
    os.close(tmp_fd)
    tmp_path = Path(tmp_name)
    doc.save(str(tmp_path))

    print("[4/5] 验证临时文件（结构 + 内容 + 样式指纹，未动模板）")
    errors = validate_doc(tmp_path, content, fp_before)
    if errors:
        tmp_path.unlink(missing_ok=True)
        print(f"[FAIL] 写入后验证发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"  - {e}")
        print(f"[KEEP] 模板未被替换（字节不变），pending 保留排查: {pending}")
        sys.exit(3)
    if modified != 2:
        print(f"[WARN] 封面文本框修改数 {modified}（期望 2），已验证镜像内容正确，继续")

    os.replace(tmp_path, tpl)
    print("[5/5] 原子替换完成")

    try:
        pending.unlink()
    except OSError as e:
        print(f"[ERROR] pending 清理失败: {e}")
        if rollback_template(tpl, original_bytes):
            print("[ROLLBACK] 模板已回滚为写入前字节（本次写入作废，pending 保留可重试）")
            fail(4, f"[EXIT-4] pending 清理失败，已回滚模板避免\"模板已变 + pending 残留\"的中间态",
                 f"        排查 scripts 目录权限后重跑即可；快照留档: {snap}")
        fail(4, f"[EXIT-4] pending 清理失败且回滚失败，请手工从快照恢复模板: {snap}",
             f"        当前 pending 保留: {pending}")

    print("[OK] 全部验证通过！")
    print("     ✅ 段数 18 / 图片 4 / 引流段样式 / 封面前缀双镜像")
    print("     ✅ 14 字段全部命中（含段[8]🙅 / 段[10]👍 前缀）")
    print("     ✅ 样式指纹零漂移（逐段 pPr/rPr 与模板一致）")
    print(f"[CLEAN] 已清理 pending_tips.json")
    print(f"\n{'=' * 50}\n模板已更新: {tpl}\n{'=' * 50}")


if __name__ == "__main__":
    main()
