# -*- coding: utf-8 -*-
"""隔离回归测试：不读写真实工作副本、pending 或进度文件。"""

from __future__ import annotations

import hashlib
import importlib.util
import io
import json
import os
import shutil
import sys
import tempfile
import unittest
import uuid
from contextlib import redirect_stdout
from pathlib import Path
from unittest import mock

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPT_PATH = Path(__file__).with_name("fill_sports_daily_post.py")
SOURCE_TEMPLATE_FIXTURE = (
    Path(__file__).resolve().parent.parent
    / "模板文件"
    / "1 体育笔试每日一练-帖子内容编辑模板.docx"
)
DOCX_NAME = "1 体育笔试每日一练-帖子内容编辑模板.docx"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


class FillSportsDailyPostTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory(prefix="sports-daily-tests-")
        self.addCleanup(self.temp_dir.cleanup)

        self.workspace = Path(self.temp_dir.name) / "体育教师编"
        self.script_dir = self.workspace / "scripts"
        self.source_dir = self.workspace / "模板文件"
        self.output_dir = self.workspace / "desktop-attachments"
        self.snapshot_dir = self.script_dir / "_snapshots_sports"
        self.pending_path = self.script_dir / "pending_sports_daily.json"
        self.progress_path = self.script_dir / "progress_sports.json"
        self.source_path = self.source_dir / DOCX_NAME
        self.output_path = self.output_dir / DOCX_NAME

        self.script_dir.mkdir(parents=True)
        self.source_dir.mkdir(parents=True)
        shutil.copy2(SOURCE_TEMPLATE_FIXTURE, self.source_path)
        self.source_hash = sha256(self.source_path)
        self.progress_bytes = b'{"last_done":null,"last_date":null,"done":[]}\n'
        self.progress_path.write_bytes(self.progress_bytes)

        module_name = f"fill_sports_daily_post_test_{uuid.uuid4().hex}"
        spec = importlib.util.spec_from_file_location(module_name, SCRIPT_PATH)
        if spec is None or spec.loader is None:
            self.fail(f"无法导入生产脚本：{SCRIPT_PATH}")
        self.module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = self.module
        self.addCleanup(sys.modules.pop, module_name, None)
        spec.loader.exec_module(self.module)

        # 生产函数的路径全部重绑到 TemporaryDirectory。validate_source_template
        # 的默认参数在定义时已绑定，需同步改写 __defaults__ 才能隔离 main()。
        self.module.SCRIPT_DIR = self.script_dir
        self.module.WORKSPACE = self.workspace
        self.module.PROJECT_SCRIPT_DIR = self.script_dir
        self.module.SOURCE_TEMPLATE = self.source_path
        self.module.TEMPLATE_PATH = self.output_path
        self.module.PENDING_JSON = self.pending_path
        self.module.SNAPSHOT_DIR = self.snapshot_dir
        self.module.validate_source_template.__defaults__ = (self.source_path,)
        self.module.fill_cover.__defaults__ = (self.output_path,)
        self.module.fill_body.__defaults__ = (self.output_path,)
        self.module.validate_output.__defaults__ = (self.output_path,)

    def valid_pending(self, table_count: int = 1, duplicate_points: bool = False) -> dict:
        points = [
            "先看动作发生的方向。",
            "屈伸可以结合两骨夹角变化判断。",
            "再看动作与人体正中面的距离。",
            "内收外展可以结合正中面判断。",
        ]
        if duplicate_points:
            points[3] = points[1]
        tables = [
            {
                "header": ["运动", "判断依据"],
                "rows": [["屈", "两骨夹角减小"], ["伸", "两骨夹角增大"]],
            },
            {
                "header": ["运动", "方向", "参照"],
                "rows": [
                    ["内收", "靠近正中面", "标准姿势"],
                    ["外展", "远离正中面", "标准姿势"],
                ],
            },
        ][:table_count]
        return {
            "title": "关节运动基本形式",
            "module": "运动解剖学",
            "question": "下列哪项属于方向相反的一组关节运动？",
            "options": ["屈和伸", "屈和旋转", "伸和环转", "内收和旋转"],
            "answer": "答案：A。屈和伸是一组方向相反的运动。",
            "points": points,
            "tables": tables,
            "exam_tips": [
                "屈伸常结合两骨夹角变化判断。",
                "内收外展要看与正中面的距离。",
                "答题时先锁定动作方向。",
                "组合选项需要逐项核对。",
                "相反运动名称容易交叉混搭。",
            ],
            "cta": "关注我，每天一个体育笔试知识点，帮你一次上岸",
            "hashtags": "#教师编 #体育教师 #体育笔试 #每天一个知识点 #一次上岸",
        }

    def write_pending(self, data) -> bytes:
        payload = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
        self.pending_path.write_bytes(payload)
        return payload

    def run_main_success(self) -> str:
        stream = io.StringIO()
        with redirect_stdout(stream):
            self.module.main()
        return stream.getvalue()

    def run_main_failure(self, expected_code: int = 1) -> str:
        stream = io.StringIO()
        with redirect_stdout(stream):
            with self.assertRaises(SystemExit) as caught:
                self.module.main()
        self.assertEqual(caught.exception.code, expected_code)
        return stream.getvalue()

    def assert_isolated_inputs_unchanged(self) -> None:
        self.assertEqual(sha256(self.source_path), self.source_hash)
        self.assertEqual(self.progress_path.read_bytes(), self.progress_bytes)

    def assert_no_staged_docx(self) -> None:
        staged = list(self.output_dir.glob(".sports_daily_build_*.docx"))
        self.assertEqual(staged, [])

    def assert_table_is_adaptive(self, table) -> None:
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        self.assertIsNotNone(layout)
        self.assertEqual(layout.get(qn("w:type")), "autofit")
        tbl_width = tbl_pr.find(qn("w:tblW"))
        self.assertIsNotNone(tbl_width)
        self.assertEqual(tbl_width.get(qn("w:type")), "auto")
        self.assertEqual(tbl_width.get(qn("w:w")), "0")

        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.trPr
            self.assertIsNotNone(tr_pr)
            height = tr_pr.find(qn("w:trHeight"))
            self.assertIsNotNone(height)
            self.assertEqual(height.get(qn("w:hRule")), "atLeast")
            self.assertEqual(height.get(qn("w:val")), "0")
            has_repeat_header = tr_pr.find(qn("w:tblHeader")) is not None
            self.assertEqual(has_repeat_header, row_index == 0)

            for cell in row.cells:
                tc_pr = cell._tc.tcPr
                self.assertIsNotNone(tc_pr)
                width = tc_pr.find(qn("w:tcW"))
                self.assertIsNotNone(width)
                self.assertEqual(width.get(qn("w:type")), "auto")
                self.assertEqual(width.get(qn("w:w")), "0")
                vertical = tc_pr.find(qn("w:vAlign"))
                self.assertIsNotNone(vertical)
                self.assertEqual(vertical.get(qn("w:val")), "center")

    def test_source_template_preflight_accepts_contract_and_rejects_missing_placeholder(self) -> None:
        self.assertTrue(self.module.validate_source_template())

        broken = self.source_dir / "missing_hashtags.docx"
        shutil.copy2(self.source_path, broken)
        doc = Document(broken)
        hits = 0
        for paragraph in doc.paragraphs:
            for text_node in paragraph._element.findall(".//" + qn("w:t")):
                if text_node.text == "{HASHTAGS}":
                    text_node.text = ""
                    hits += 1
        self.assertEqual(hits, 1)
        doc.save(broken)

        with self.assertRaisesRegex(RuntimeError, r"\{HASHTAGS\} actual=0 expected=1"):
            self.module.validate_source_template(broken)
        self.assert_isolated_inputs_unchanged()

    def test_load_pending_rejects_invalid_answer_colons_and_non_string_values(self) -> None:
        cases = []

        bad_answer = self.valid_pending()
        bad_answer["answer"] = "答案：A。第一句。第二句。"
        cases.append(("answer 格式", bad_answer, "answer 必须匹配"))

        answer_colon = self.valid_pending()
        answer_colon["answer"] = "答案：A。解析：错误。"
        cases.append(("answer 多余冒号", answer_colon, "answer 必须匹配"))

        visible_colon = self.valid_pending()
        visible_colon["points"][0] = "提示：先看动作方向。"
        cases.append(("正文中文冒号", visible_colon, "不得含中文冒号、ASCII 冒号或破折号"))

        ascii_colon = self.valid_pending()
        ascii_colon["module"] = "模块:解剖学"
        cases.append(("正文 ASCII 冒号", ascii_colon, "不得含中文冒号、ASCII 冒号或破折号"))

        non_string_option = self.valid_pending()
        non_string_option["options"][0] = 1
        cases.append(("选项非字符串", non_string_option, "options\\[1\\] 必须是非空字符串"))

        non_string_answer = self.valid_pending()
        non_string_answer["answer"] = 1
        cases.append(("答案非字符串", non_string_answer, "answer 必须是非空字符串"))

        non_string_cell = self.valid_pending()
        non_string_cell["tables"][0]["rows"][0][0] = 1
        cases.append(("单元格非字符串", non_string_cell, "必须是非空字符串"))

        cases.append(("顶层非对象", [], "顶层必须是对象"))

        for label, data, error_pattern in cases:
            with self.subTest(label=label):
                self.write_pending(data)
                with self.assertRaisesRegex(ValueError, error_pattern):
                    self.module.load_pending()

        self.assertFalse(self.output_path.exists())
        self.assert_isolated_inputs_unchanged()

    def test_one_table_succeeds_when_point_text_is_duplicated(self) -> None:
        data = self.valid_pending(table_count=1, duplicate_points=True)
        self.write_pending(data)

        output = self.run_main_success()

        self.assertIn("✅ 全部通过", output)
        self.assertIn("✅ pending_sports_daily.json 已删除", output)
        self.assertTrue(self.output_path.exists())
        self.assertFalse(self.pending_path.exists())
        self.assert_no_staged_docx()

        doc = Document(self.output_path)
        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        previous = table._tbl.getprevious()
        previous_text = "".join(
            node.text or "" for node in previous.findall(".//" + qn("w:t"))
        ).strip()
        self.assertEqual(previous_text, data["points"][1])

        body_sequence = []
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                text = "".join(
                    node.text or "" for node in child.findall(".//" + qn("w:t"))
                ).strip()
                body_sequence.append(("paragraph", text))
            elif child.tag == qn("w:tbl"):
                body_sequence.append(("table", ""))
        duplicate_indices = [
            index
            for index, item in enumerate(body_sequence)
            if item == ("paragraph", data["points"][1])
        ]
        table_index = body_sequence.index(("table", ""))
        self.assertEqual(len(duplicate_indices), 2)
        self.assertEqual(table_index, duplicate_indices[0] + 1)
        self.assertLess(table_index, duplicate_indices[1])
        self.assert_table_is_adaptive(table)
        self.assertEqual(self.module.validate_output(data, self.output_path), [])
        self.assert_isolated_inputs_unchanged()

    def test_two_tables_commit_atomically_delete_pending_and_preserve_anchor_styles(self) -> None:
        data = self.valid_pending(table_count=2)
        self.write_pending(data)

        self.output_dir.mkdir(parents=True)
        old_doc = Document(self.source_path)
        old_doc.core_properties.title = "old-working-copy-sentinel"
        old_doc.save(self.output_path)
        old_hash = sha256(self.output_path)

        with mock.patch.object(self.module.os, "replace", wraps=os.replace) as replace:
            output = self.run_main_success()

        replace.assert_called_once()
        staged_path = Path(replace.call_args.args[0])
        committed_path = Path(replace.call_args.args[1])
        self.assertEqual(staged_path.parent, self.output_dir)
        self.assertTrue(staged_path.name.startswith(".sports_daily_build_"))
        self.assertEqual(committed_path, self.output_path)
        self.assertFalse(staged_path.exists())
        self.assertNotEqual(sha256(self.output_path), old_hash)

        snapshots = list(self.snapshot_dir.glob("snapshot_*.docx"))
        self.assertEqual(len(snapshots), 1)
        self.assertEqual(sha256(snapshots[0]), old_hash)
        self.assertIn("✅ 已原子提交", output)
        self.assertIn("✅ 全部通过", output)
        self.assertIn("✅ pending_sports_daily.json 已删除", output)
        self.assertFalse(self.pending_path.exists())
        self.assert_no_staged_docx()

        doc = Document(self.output_path)
        self.assertEqual(len(doc.tables), 2)
        for table_index, table in enumerate(doc.tables):
            expected_anchor = data["points"][1 if table_index == 0 else 3]
            previous = table._tbl.getprevious()
            previous_text = "".join(
                node.text or "" for node in previous.findall(".//" + qn("w:t"))
            ).strip()
            self.assertEqual(previous_text, expected_anchor)
            self.assert_table_is_adaptive(table)

        self.assertEqual(self.module.validate_output(data, self.output_path), [])
        self.assert_isolated_inputs_unchanged()

    def test_first_run_fill_body_failure_leaves_no_output_and_preserves_pending_bytes(self) -> None:
        pending_bytes = self.write_pending(self.valid_pending())

        with mock.patch.object(
            self.module,
            "fill_body",
            side_effect=RuntimeError("injected fill_body failure"),
        ):
            output = self.run_main_failure()

        self.assertIn("写入正文失败", output)
        self.assertIn("首次运行失败，未留下半成品工作副本", output)
        self.assertIn("[KEEP]", output)
        self.assertFalse(self.output_path.exists())
        self.assertEqual(self.pending_path.read_bytes(), pending_bytes)
        self.assert_no_staged_docx()
        self.assertFalse(self.snapshot_dir.exists())
        self.assert_isolated_inputs_unchanged()

    def test_multi_run_cover_paragraph_passes_preflight_and_output_validation(self) -> None:
        # 模拟 WPS 编辑把封面首行拆成多个 run：预检按段落拼接比对可通过，
        # 填充后的输出校验也必须按同口径通过（v1.5 修复的不一致）
        doc = Document(self.source_path)
        cover = doc.paragraphs[0]
        txbx_list = cover._element.findall(".//" + qn("w:txbxContent"))
        self.assertEqual(len(txbx_list), 2)
        for txbx in txbx_list:
            first_para = txbx.findall(qn("w:p"))[0]
            texts = first_para.findall(".//" + qn("w:t"))
            full = "".join(t.text or "" for t in texts)
            self.assertEqual(full, "每天一个体育笔试知识点")
            texts[0].text = full[:4]
            for extra in texts[1:]:
                extra.text = ""
            new_run = OxmlElement("w:r")
            new_text = OxmlElement("w:t")
            new_text.text = full[4:]
            new_run.append(new_text)
            texts[0].getparent().addnext(new_run)
        doc.save(self.source_path)
        self.source_hash = sha256(self.source_path)

        self.assertTrue(self.module.validate_source_template())

        data = self.valid_pending(table_count=1)
        self.write_pending(data)
        output = self.run_main_success()

        self.assertIn("✅ 全部通过", output)
        self.assertIn("✅ pending_sports_daily.json 已删除", output)
        self.assertEqual(self.module.validate_output(data, self.output_path), [])
        self.assert_no_staged_docx()
        self.assert_isolated_inputs_unchanged()

    def test_validate_exception_preserves_existing_working_copy_hash(self) -> None:
        pending_bytes = self.write_pending(self.valid_pending(table_count=2))
        self.output_dir.mkdir(parents=True)
        old_doc = Document(self.source_path)
        old_doc.core_properties.title = "existing-copy-must-survive"
        old_doc.save(self.output_path)
        old_hash = sha256(self.output_path)

        with mock.patch.object(
            self.module,
            "validate_output",
            side_effect=RuntimeError("injected validate_output failure"),
        ):
            output = self.run_main_failure()

        self.assertIn("自动验证失败", output)
        self.assertIn("旧工作副本未改动", output)
        self.assertEqual(sha256(self.output_path), old_hash)
        self.assertEqual(self.pending_path.read_bytes(), pending_bytes)
        self.assert_no_staged_docx()
        self.assertFalse(self.snapshot_dir.exists())
        self.assert_isolated_inputs_unchanged()


if __name__ == "__main__":
    unittest.main(verbosity=2)
