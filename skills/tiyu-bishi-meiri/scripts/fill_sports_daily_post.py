# -*- coding: utf-8 -*-
"""
fill_sports_daily_post.py v1.5
==============================

「体育笔试每日一练」帖子填充脚本。

v1.5 变更：封面输出校验改为按段落拼接比对（与源模板预检口径一致），
容忍 WPS 编辑造成的单段多 run，避免预检通过、输出校验必败的死循环。

流程：
  1. agent 把新一篇内容写到 scripts/pending_sports_daily.json
  2. 跑本脚本：读取与预检 -> 临时构建 -> 自动验证 -> 快照 -> 原子提交 -> 删 JSON

输入：scripts/pending_sports_daily.json
输出：desktop-attachments/1 体育笔试每日一练-帖子内容编辑模板.docx（原地覆盖）

JSON schema（硬性）：
  title      str   知识点标题
  module     str   模块标签（如：运动解剖学）
  question   str   单选题题干（<=60 字）
  options    list  4 个选项文本（不含 A. 前缀）
  answer     str   答案揭晓（含解析一句话）
  points     list  4 段引导说明（每段 <=120 字，核心信息进表格）
  tables     list  1-2 张表，每张 {"header": [...], "rows": [[...], ...]}
  exam_tips  list  5 条考法提醒
  cta        str   引流文案（须与模板固定引流段一致）
  hashtags   str   话题标签（如：#教师编 #体育笔试）
"""

import json
import os
import re
import shutil
import sys
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


SCRIPT_DIR = Path(__file__).resolve().parent
WORKSPACE = Path(os.environ.get("SPORTS_DAILY_WORKSPACE", SCRIPT_DIR.parent)).expanduser().resolve()
PROJECT_SCRIPT_DIR = WORKSPACE / "scripts"
SOURCE_TEMPLATE = WORKSPACE / "模板文件" / "1 体育笔试每日一练-帖子内容编辑模板.docx"
TEMPLATE_PATH = WORKSPACE / "desktop-attachments" / "1 体育笔试每日一练-帖子内容编辑模板.docx"
PENDING_JSON = PROJECT_SCRIPT_DIR / "pending_sports_daily.json"
SNAPSHOT_DIR = PROJECT_SCRIPT_DIR / "_snapshots_sports"
MAX_SNAPSHOTS = 10

EXPECTED_PARAS = 17
EXPECTED_IMAGES = 2
EXPECTED_TXBX = 2
COVER_TITLE = "每天一个体育笔试知识点"
DRAIN_TEXT = "关注我，每天一个体育笔试知识点，帮你一次上岸"
MAX_QUESTION_LEN = 60
MAX_POINT_LEN = 120
MAX_TITLE_LEN = 20
NAVY = RGBColor(0x0B, 0x32, 0x89)
DARK = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def take_snapshot():
    SNAPSHOT_DIR.mkdir(exist_ok=True)
    if not TEMPLATE_PATH.exists():
        return None
    ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    snap = SNAPSHOT_DIR / f"snapshot_{ts}.docx"
    shutil.copy(TEMPLATE_PATH, snap)
    snaps = sorted(SNAPSHOT_DIR.glob("snapshot_*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    for old in snaps[MAX_SNAPSHOTS:]:
        old.unlink()
    return snap


def _require_text(value, label):
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{label} 必须是非空字符串")
    return value


def _visible_text_values(data):
    values = [(key, data[key]) for key in ("title", "module", "question", "cta", "hashtags")]
    values.extend((f"options[{i}]", value) for i, value in enumerate(data["options"], 1))
    values.extend((f"points[{i}]", value) for i, value in enumerate(data["points"], 1))
    values.extend((f"exam_tips[{i}]", value) for i, value in enumerate(data["exam_tips"], 1))
    for ti, table in enumerate(data["tables"], 1):
        values.extend((f"tables[{ti}].header[{i}]", value) for i, value in enumerate(table["header"], 1))
        for ri, row in enumerate(table["rows"], 1):
            values.extend((f"tables[{ti}].rows[{ri}][{ci}]", value) for ci, value in enumerate(row, 1))
    return values


def load_pending():
    with open(PENDING_JSON, encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError("pending_sports_daily.json 顶层必须是对象")
    required = ["title", "module", "question", "options", "answer", "points", "tables", "exam_tips", "cta", "hashtags"]
    missing = [k for k in required if k not in data]
    if missing:
        raise ValueError(f"JSON 缺字段：{missing}")
    for key in ("title", "module", "question", "answer", "cta", "hashtags"):
        _require_text(data[key], key)
    if len(data["title"]) > MAX_TITLE_LEN:
        raise ValueError(f"title {len(data['title'])} 字，超过硬上限 {MAX_TITLE_LEN} 字")
    if len(data["question"]) > MAX_QUESTION_LEN:
        raise ValueError(f"题干 {len(data['question'])} 字，超过硬上限 {MAX_QUESTION_LEN} 字")
    if not isinstance(data["options"], list) or len(data["options"]) != 4:
        raise ValueError("options 必须是 4 个非空字符串")
    for i, value in enumerate(data["options"], 1):
        _require_text(value, f"options[{i}]")
    if not isinstance(data["points"], list) or len(data["points"]) != 4:
        raise ValueError("points 必须是 4 段非空字符串")
    for i, value in enumerate(data["points"], 1):
        _require_text(value, f"points[{i}]")
    if any(len(str(p)) > MAX_POINT_LEN for p in data["points"]):
        raise ValueError(f"points 每段不能超过 {MAX_POINT_LEN} 字（核心信息应放进表格）")
    tables = data["tables"]
    if not isinstance(tables, list) or not 1 <= len(tables) <= 2:
        raise ValueError("tables 必须是 1-2 张表")
    for ti, spec in enumerate(tables):
        if not isinstance(spec, dict) or "header" not in spec or "rows" not in spec:
            raise ValueError(f"表[{ti + 1}] 必须是 {{'header': [...], 'rows': [[...]]}}")
        header = spec["header"]
        rows = spec["rows"]
        if not isinstance(header, list) or not 2 <= len(header) <= 5:
            raise ValueError(f"表[{ti + 1}] header 必须是 2-5 个非空字符串")
        for hi, value in enumerate(header, 1):
            _require_text(value, f"tables[{ti + 1}].header[{hi}]")
        if not isinstance(rows, list) or not 1 <= len(rows) <= 8:
            raise ValueError(f"表[{ti + 1}] rows 必须是 1-8 行")
        for ri, row in enumerate(rows):
            if not isinstance(row, list) or len(row) != len(header):
                raise ValueError(f"表[{ti + 1}] 第 {ri + 1} 行必须是 {len(header)} 个非空字符串")
            for ci, value in enumerate(row, 1):
                _require_text(value, f"tables[{ti + 1}].rows[{ri + 1}][{ci}]")
    if not isinstance(data["exam_tips"], list) or len(data["exam_tips"]) != 5:
        raise ValueError("exam_tips 必须是 5 条非空字符串")
    for i, value in enumerate(data["exam_tips"], 1):
        _require_text(value, f"exam_tips[{i}]")
    if not re.fullmatch(r"答案：[A-D]。[^：:\r\n。]+。", data["answer"]):
        raise ValueError("answer 必须匹配：答案：A。一句不含冒号的解析。")
    if data["answer"].count("：") != 1 or ":" in data["answer"] or "——" in data["answer"]:
        raise ValueError("answer 只允许固定前缀中的 1 个中文冒号，并禁止 ASCII 冒号与破折号")
    for label, value in _visible_text_values(data):
        if "：" in value or ":" in value or "——" in value:
            raise ValueError(f"{label} 不得含中文冒号、ASCII 冒号或破折号")
    if data["cta"].strip() != DRAIN_TEXT:
        raise ValueError("cta 与模板固定引流段不一致（不许改引流文案）")
    return data


def _real_page_break_before(paragraph):
    el = paragraph.find(qn("w:pPr") + "/" + qn("w:pageBreakBefore"))
    return el is not None and el.get(qn("w:val")) not in {"0", "false", "off"}


def validate_source_template(path=SOURCE_TEMPLATE):
    try:
        doc = Document(path)
    except Exception as exc:
        raise RuntimeError(f"源模板不可读取：{path}；{exc}") from exc

    errors = []
    paras = [p for p in doc.element.body.iterchildren(qn("w:p"))]
    if len(paras) != EXPECTED_PARAS:
        errors.append(f"正文段数 actual={len(paras)} expected={EXPECTED_PARAS}")
    if len(doc.tables) != 0:
        errors.append(f"源模板表格数 actual={len(doc.tables)} expected=0")
    drawings = sum(len(p.findall(".//" + qn("w:drawing"))) for p in paras)
    if drawings != EXPECTED_IMAGES:
        errors.append(f"drawing 数 actual={drawings} expected={EXPECTED_IMAGES}")

    txbx_list = [t for p in paras for t in p.findall(".//" + qn("w:txbxContent"))]
    if len(txbx_list) != EXPECTED_TXBX:
        errors.append(f"文本框镜像数 actual={len(txbx_list)} expected={EXPECTED_TXBX}")
    else:
        expected_cover = [
            COVER_TITLE,
            "【{MODULE}】",
            "{QUESTION}",
            "A．{OPTION_A}",
            "B．{OPTION_B}",
            "C．{OPTION_C}",
            "D．{OPTION_D}",
        ]
        for i, txbx in enumerate(txbx_list, 1):
            lines = ["".join(t.text or "" for t in p.findall(".//" + qn("w:t"))) for p in txbx.findall(qn("w:p"))]
            if lines != expected_cover:
                errors.append(f"封面文本框镜像[{i}] actual={lines!r} expected={expected_cover!r}")

    placeholders = [
        "今日考点｜{TITLE}", "{ANSWER}", "{POINT_1}", "{POINT_2}", "{POINT_3}", "{POINT_4}",
        "{TIP_1}", "{TIP_2}", "{TIP_3}", "{TIP_4}", "{TIP_5}", "{HASHTAGS}",
    ]
    text_nodes = [t.text or "" for p in doc.paragraphs for t in p._element.findall(".//" + qn("w:t"))]
    for placeholder in placeholders:
        count = sum(text == placeholder for text in text_nodes)
        if count != 1:
            errors.append(f"占位符 {placeholder} actual={count} expected=1")
    cta_count = sum(p.text == DRAIN_TEXT for p in doc.paragraphs)
    if cta_count != 1:
        errors.append(f"固定引流段 actual={cta_count} expected=1")

    page_break_paras = [p.text.strip() for p in doc.paragraphs if _real_page_break_before(p._element)]
    if page_break_paras != ["考法提醒"]:
        errors.append(f"pageBreakBefore actual={page_break_paras!r} expected=['考法提醒']")
    run_page_break_count = sum(
        br.get(qn("w:type")) == "page" for br in doc.element.findall(".//" + qn("w:br"))
    )
    if run_page_break_count:
        errors.append(f"w:br type=page actual={run_page_break_count} expected=0")
    if errors:
        raise RuntimeError(f"源模板契约不符：{path}；" + "；".join(errors))
    return True


def set_run_text(paragraph, text):
    runs = paragraph.findall('.//' + qn('w:t'))
    if not runs:
        raise RuntimeError(f"段落无文本 run：{paragraph}")
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def fill_cover(data, path=TEMPLATE_PATH):
    doc = Document(path)
    cover = doc.paragraphs[0]
    txbx_list = cover._element.findall('.//' + qn('w:txbxContent'))
    if len(txbx_list) != EXPECTED_TXBX:
        raise RuntimeError(f"封面文本框数量异常：{len(txbx_list)}")
    lines = [
        COVER_TITLE,
        f"【{data['module']}】",
        data["question"],
        f"A．{data['options'][0]}",
        f"B．{data['options'][1]}",
        f"C．{data['options'][2]}",
        f"D．{data['options'][3]}",
    ]
    for txbx in txbx_list:
        paras = txbx.findall(qn('w:p'))
        if len(paras) != len(lines):
            raise RuntimeError(f"文本框段落数异常：{len(paras)}")
        for p, text in zip(paras, lines):
            set_run_text(p, text)
    doc.save(path)


def replace_placeholder(doc, placeholder, value):
    hit = 0
    for para in doc.paragraphs:
        for t in para._element.findall('.//' + qn('w:t')):
            if t.text == placeholder:
                t.text = value
                hit += 1
    return hit


def style_run(run, size=14, bold=False, color=DARK):
    run.font.name = "微软雅黑"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def fill_cell(cell, text, header=False):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = str(text)
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(str(text))
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0]
    if header:
        style_run(run, size=14, bold=True, color=WHITE)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0B3289')
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = tcPr.find(qn('w:vAlign'))
        if vAlign is not None:
            vAlign.addprevious(shd)
        else:
            tcPr.append(shd)
    else:
        style_run(run, size=14, bold=False, color=DARK)


def build_table(doc, spec):
    header = spec["header"]
    rows = spec["rows"]
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._tbl.tblPr
    for child in list(tblPr):
        tblPr.remove(child)
    # 总宽自动、列宽随内容自适应（autofit），与 WPS 手动调校一致
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), '8EAADB')
        el.set(qn('w:space'), '0')
        borders.append(el)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'autofit')
    cell_mar = OxmlElement('w:tblCellMar')
    for name, w in (('top', '0'), ('left', '108'), ('bottom', '0'), ('right', '108')):
        m = OxmlElement('w:' + name)
        m.set(qn('w:w'), w)
        m.set(qn('w:type'), 'dxa')
        cell_mar.append(m)
    for el in (tblW, jc, borders, layout, cell_mar):
        tblPr.append(el)
    # 行高随内容自适应（至少 0），首行重复表头，行内容水平居中
    for ri, row in enumerate(tbl.rows):
        trPr = row._tr.get_or_add_trPr()
        trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), '0')
        trh.set(qn('w:hRule'), 'atLeast')
        trPr.append(trh)
        jcr = OxmlElement('w:jc')
        jcr.set(qn('w:val'), 'center')
        trPr.append(jcr)
        if ri == 0:
            th = OxmlElement('w:tblHeader')
            trPr.append(th)
    # 单元格列宽自动、内容垂直居中
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:type'), 'auto')
            tcW.set(qn('w:w'), '0')
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
    for j, h in enumerate(header):
        fill_cell(tbl.cell(0, j), h, header=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            fill_cell(tbl.cell(i + 1, j), val, header=False)
    return tbl


def fill_body(data, path=TEMPLATE_PATH):
    doc = Document(path)
    marker_paragraphs = {}
    for marker in ("{POINT_2}", "{POINT_4}"):
        found = [
            p for p in doc.paragraphs
            if any((t.text or "") == marker for t in p._element.findall(".//" + qn("w:t")))
        ]
        if len(found) != 1:
            raise RuntimeError(f"表格锚点 {marker} 命中 {len(found)} 处（期望 1）")
        marker_paragraphs[marker] = found[0]
    pairs = [
        ("今日考点｜{TITLE}", f"今日考点｜{data['title']}"),
        ("{ANSWER}", data["answer"]),
        ("{POINT_1}", data["points"][0]),
        ("{POINT_2}", data["points"][1]),
        ("{POINT_3}", data["points"][2]),
        ("{POINT_4}", data["points"][3]),
        ("{TIP_1}", data["exam_tips"][0]),
        ("{TIP_2}", data["exam_tips"][1]),
        ("{TIP_3}", data["exam_tips"][2]),
        ("{TIP_4}", data["exam_tips"][3]),
        ("{TIP_5}", data["exam_tips"][4]),
        ("{HASHTAGS}", data["hashtags"]),
    ]
    for ph, value in pairs:
        n = replace_placeholder(doc, ph, value)
        if n != 1:
            raise RuntimeError(f"占位符 {ph} 命中 {n} 处（期望 1）")
    anchors = [marker_paragraphs["{POINT_2}"], marker_paragraphs["{POINT_4}"]][:len(data["tables"])]
    for para, spec in zip(anchors, data["tables"]):
        tbl = build_table(doc, spec)
        para._element.addnext(tbl._tbl)
    doc.save(path)


def all_text(doc):
    parts = []
    for para in doc.paragraphs:
        for t in para._element.findall('.//' + qn('w:t')):
            parts.append(t.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def validate_output(data, path=TEMPLATE_PATH):
    doc = Document(path)
    errors = []
    paras = [p for p in doc.element.body.iterchildren(qn('w:p'))]
    if len(paras) != EXPECTED_PARAS:
        errors.append(f"段数异常：{len(paras)}（期望 {EXPECTED_PARAS}）")
    imgs = sum(len(p.findall('.//' + qn('w:drawing'))) for p in paras)
    if imgs != EXPECTED_IMAGES:
        errors.append(f"图片数异常：{imgs}（期望 {EXPECTED_IMAGES}）")
    txbx_list = [t for p in paras for t in p.findall('.//' + qn('w:txbxContent'))]
    if len(txbx_list) != EXPECTED_TXBX:
        errors.append(f"文本框数异常：{len(txbx_list)}（期望 {EXPECTED_TXBX}）")
    else:
        def _cover_lines(txbx):
            return [
                "".join(t.text or "" for t in p.findall(".//" + qn("w:t")))
                for p in txbx.findall(qn("w:p"))
            ]

        t0 = _cover_lines(txbx_list[0])
        t1 = _cover_lines(txbx_list[1])
        if t0 != t1:
            errors.append("封面文本框 2 处镜像不一致")
        if not t0 or t0[0] != COVER_TITLE:
            errors.append(f"封面大标题被改：{t0[0] if t0 else '<empty>'!r}")
        expected_cover = [
            COVER_TITLE,
            f"【{data['module']}】",
            data["question"],
            f"A．{data['options'][0]}",
            f"B．{data['options'][1]}",
            f"C．{data['options'][2]}",
            f"D．{data['options'][3]}",
        ]
        if t0 != expected_cover:
            errors.append(f"封面内容 actual={t0!r} expected={expected_cover!r}")
    page_break_paras = [p.text.strip() for p in doc.paragraphs if _real_page_break_before(p._element)]
    if page_break_paras != ["考法提醒"]:
        errors.append(f"pageBreakBefore actual={page_break_paras!r} expected=['考法提醒']")
    run_page_break_count = sum(
        br.get(qn("w:type")) == "page" for br in doc.element.findall(".//" + qn("w:br"))
    )
    if run_page_break_count:
        errors.append(f"w:br type=page actual={run_page_break_count} expected=0")
    expected_tables = len(data["tables"])
    if len(doc.tables) != expected_tables:
        errors.append(f"表格数异常：{len(doc.tables)}（期望 {expected_tables}）")
    for ti, tbl in enumerate(doc.tables):
        layout = tbl._tbl.tblPr.find(qn('w:tblLayout'))
        if layout is None or layout.get(qn('w:type')) != 'autofit':
            errors.append(f"表[{ti + 1}] 未启用列宽自适应（tblLayout autofit）")
        tblW = tbl._tbl.tblPr.find(qn('w:tblW'))
        if tblW is None or tblW.get(qn('w:type')) != 'auto':
            errors.append(f"表[{ti + 1}] 总宽未设为自动")
        expected_anchor = data["points"][1 if ti == 0 else 3]
        previous = tbl._tbl.getprevious()
        previous_text = "" if previous is None else "".join(
            t.text or "" for t in previous.findall(".//" + qn("w:t"))
        ).strip()
        if previous_text != expected_anchor:
            errors.append(f"表[{ti + 1}] 位置异常：前一段未精确命中预期锚点")
        for ri, row in enumerate(tbl.rows, 1):
            tr_pr = row._tr.trPr
            tr_height = None if tr_pr is None else tr_pr.find(qn("w:trHeight"))
            if (
                tr_height is None
                or tr_height.get(qn("w:hRule")) != "atLeast"
                or tr_height.get(qn("w:val")) != "0"
            ):
                errors.append(f"表[{ti + 1}] 第 {ri} 行未使用自适应行高")
            has_header = tr_pr is not None and tr_pr.find(qn("w:tblHeader")) is not None
            if (ri == 1) != has_header:
                errors.append(f"表[{ti + 1}] 第 {ri} 行重复表头标记异常")
            for ci, cell in enumerate(row.cells, 1):
                tc_pr = cell._tc.tcPr
                tc_width = None if tc_pr is None else tc_pr.find(qn("w:tcW"))
                if (
                    tc_width is None
                    or tc_width.get(qn("w:type")) != "auto"
                    or tc_width.get(qn("w:w")) != "0"
                ):
                    errors.append(f"表[{ti + 1}] 单元格[{ri}][{ci}]列宽未设为自动")
                v_align = None if tc_pr is None else tc_pr.find(qn("w:vAlign"))
                if v_align is None or v_align.get(qn("w:val")) != "center":
                    errors.append(f"表[{ti + 1}] 单元格[{ri}][{ci}]未垂直居中")
    for ti, spec in enumerate(data["tables"]):
        header = spec["header"]
        rows = spec["rows"]
        if ti >= len(doc.tables):
            continue
        tbl = doc.tables[ti]
        if len(tbl.rows) != len(rows) + 1 or len(tbl.columns) != len(header):
            errors.append(f"表[{ti + 1}] 尺寸异常")
            continue
        for j, h in enumerate(header):
            if tbl.cell(0, j).text.strip() != h:
                errors.append(f"表[{ti + 1}] 表头[{j + 1}]不一致")
        for ri, row in enumerate(rows):
            for j, val in enumerate(row):
                if tbl.cell(ri + 1, j).text.strip() != val:
                    errors.append(f"表[{ti + 1}] 单元格[{ri + 1}][{j + 1}]未命中")
    text = all_text(doc)
    leftovers = re.findall(r"\{[A-Z_]+\}", text)
    if leftovers:
        errors.append(f"占位符残留：{sorted(set(leftovers))}")
    cta_count = sum(p.text == DRAIN_TEXT for p in doc.paragraphs)
    if cta_count != 1:
        errors.append(f"固定引流段 actual={cta_count} expected=1")
    for label, value in [
        ("title", data["title"]),
        ("module", data["module"]),
        ("answer", data["answer"]),
        ("hashtags", data["hashtags"]),
    ]:
        if value not in text:
            errors.append(f"关键词未命中：{label}")
    for i, p in enumerate(data["points"], 1):
        if p not in text:
            errors.append(f"解析段[{i}]未命中")
    for i, t in enumerate(data["exam_tips"], 1):
        if t not in text:
            errors.append(f"考法提醒[{i}]未命中")
    return errors


def main():
    temp_path = None
    stage = "读取待写入内容"
    try:
        print("[0/6] 读取待写入内容")
        data = load_pending()
        print(f"      题目：{data['question'][:40]}...（{len(data['question'])} 字）")

        stage = "源模板契约预检"
        print("[1/6] 源模板契约预检")
        validate_source_template()
        print(f"      ✅ 源模板通过：{SOURCE_TEMPLATE}")

        stage = "创建隔离构建副本"
        print("[2/6] 创建隔离构建副本")
        TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        fd, temp_name = tempfile.mkstemp(
            prefix=".sports_daily_build_",
            suffix=".docx",
            dir=TEMPLATE_PATH.parent,
        )
        os.close(fd)
        temp_path = Path(temp_name)
        shutil.copy2(SOURCE_TEMPLATE, temp_path)
        print(f"      临时文件：{temp_path.name}")

        stage = "写入封面文本框"
        print("[3/6] 写入封面文本框")
        fill_cover(data, temp_path)
        print("      封面 2 处镜像已同步")

        stage = "写入正文"
        print("[4/6] 写入正文")
        fill_body(data, temp_path)
        print(f"      正文占位符已填充（表格 {len(data['tables'])} 张）")

        stage = "自动验证"
        print("[5/6] 自动验证")
        errors = validate_output(data, temp_path)
        if errors:
            raise RuntimeError("；".join(errors))
        print("      ✅ 构建验证通过")

        stage = "快照并原子提交"
        print("[6/6] 快照并原子提交")
        snap = take_snapshot()
        print(f"      快照：{snap.resolve() if snap else '首次运行，无旧文件可备份'}")
        os.replace(temp_path, TEMPLATE_PATH)
        temp_path = None
        print(f"      ✅ 已原子提交：{TEMPLATE_PATH}")
        print("      ✅ 全部通过")
    except Exception as exc:
        if temp_path is not None and temp_path.exists():
            try:
                temp_path.unlink()
            except OSError as cleanup_exc:
                print(f"      ⚠️ 临时文件清理失败：{cleanup_exc}")
        print(f"      ❌ {stage}失败：{exc}")
        if TEMPLATE_PATH.exists():
            print("      旧工作副本未改动")
        else:
            print("      首次运行失败，未留下半成品工作副本")
        if PENDING_JSON.exists():
            print("[KEEP] pending_sports_daily.json 已原样保留，供排查")
        else:
            print("[INFO] pending_sports_daily.json 原本不存在，未创建新文件")
        sys.exit(1)

    print("[收尾] 删除中间 JSON")
    try:
        PENDING_JSON.unlink()
    except Exception as exc:
        print(f"      ❌ DOCX 已通过并提交，但 JSON 删除失败：{exc}")
        print("      [KEEP] 请只处理 pending_sports_daily.json，勿重跑填充脚本")
        sys.exit(2)
    print("      ✅ pending_sports_daily.json 已删除")
    print("完成")


if __name__ == "__main__":
    main()
