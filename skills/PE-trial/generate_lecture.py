#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
人教版体育试讲备考讲义生成器（生产版 CLI）

用法:
  python3 generate_lecture.py --base-dir <产物库根> --out <讲义.docx> \
      [--projects 01,02,...] [--only 07-10,06-10] [--font "Songti SC"] \
      [--check] [--dry-run]

退出码:
  0 成功（或 --check/--dry-run 检查通过）
  2 用法/IO 错误（占位路径、目录不存在、不可写、字体不可用）
  3 零项目（选中范围内一个试讲稿都没有）
  4 解析异常（任一 Markdown 解析失败，不吞异常、不留 final）
  5 渲染或自检失败（A4/边距/表宽/列数/字体/目录页码任一不过）

排版契约（references/artifacts-spec.md）:
  A4 21x29.7cm、四边 2.5cm、表总宽 <= 16cm、单表列数 <= 10、
  中文 eastAsia 字体显式声明、目录页码渲染后可验证。
"""

import argparse
import datetime
import glob
import os
import re
import shutil
import subprocess
import sys
import tempfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "scripts"))
import validate_artifacts as va  # noqa: E402  复用结构自检（stdlib only）

# === 常量 ===
PROJECTS = [
    ("01", "篮球", "01篮球"), ("02", "排球", "02排球"), ("03", "足球", "03足球"),
    ("04", "乒乓球", "04乒乓球"), ("05", "羽毛球", "05羽毛球"), ("06", "体操", "06体操"),
    ("07", "田径", "07田径"), ("08", "体能", "08体能"),
    ("09", "健康课程", "09健康课程"), ("10", "武术", "10武术"),
]
INFO_BOX_BG = "DAEEF3"
HEADER_ROW_BG = "D6E4F0"
TOTAL_WIDTH_CM = 16.0          # A4 - 2*2.5cm 边距
MAX_TABLE_COLS = 10            # 验证器硬门：>=11 列即坏
PAGE_W_CM, PAGE_H_CM = 21.0, 29.7
MARGIN_CM = 2.5

# 字体候选（macOS 系统字体文件 -> Word eastAsia 字体名），按优先级
FONT_CANDIDATES = [
    ("Songti SC", "/System/Library/Fonts/Supplemental/Songti.ttc"),
    ("PingFang SC", "/System/Library/Fonts/PingFang.ttc"),
    ("Hiragino Sans GB", "/System/Library/Fonts/Hiragino Sans GB.ttc"),
    ("STHeiti", "/System/Library/Fonts/STHeiti Light.ttc"),
]

PLACEHOLDER_RE = re.compile(r"[<>]|%[A-Za-z]|占位|项目根")


class GenError(Exception):
    def __init__(self, code, msg):
        super().__init__(msg)
        self.code = code


# === 底层排版工具（沿用旧版验证过的实现） ===

def shade_cell(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def mark_header_row(table):
    """首行为真表头的表标记 w:tblHeader（跨页重复表头 + 可访问性审计要求）。"""
    tr = table.rows[0]._tr
    tr.get_or_add_trPr().append(
        parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>'))


def set_table_fixed_layout(table):
    tbl = table._tbl
    for elem in tbl.tblPr.findall(qn('w:tblLayout')):
        tbl.tblPr.remove(elem)
    tbl.tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for j, w in enumerate(widths_cm):
            if j < len(row.cells):
                row.cells[j].width = Cm(w)


def compact_table(table):
    """压紧表格：单元格段落零间距 + 单倍行距，行禁止跨页断裂。

    表格单元格默认继承 Normal 的 space_after=4pt / line_spacing=1.25，
    长表会因此虚高数厘米，把表尾少量行挤成孤页。
    """
    for row in table.rows:
        row._tr.get_or_add_trPr().append(
            parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>'))
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0


def compute_col_widths(rows, n_cols):
    CHAR_W, PADDING, FIT_CAP = 0.375, 0.5, 8
    max_visual_lens = [0] * n_cols
    for row in rows:
        for j in range(min(len(row), n_cols)):
            for line in row[j].split('\n'):
                vl = sum(2 if ord(c) > 127 else 1 for c in line)
                if vl > max_visual_lens[j]:
                    max_visual_lens[j] = vl
    if sum(max_visual_lens) == 0:
        return [TOTAL_WIDTH_CM / n_cols] * n_cols
    min_widths = []
    for vl in max_visual_lens:
        chars = min(vl / 2.0, FIT_CAP)
        min_widths.append(max(chars * CHAR_W + PADDING, 1.5))
    if sum(min_widths) > TOTAL_WIDTH_CM:
        widths = list(min_widths)
        while sum(widths) > TOTAL_WIDTH_CM:
            i = widths.index(max(widths))
            if widths[i] <= 1.5:
                break
            widths[i] = max(1.5, widths[i] - 0.1)
        return widths
    remaining = TOTAL_WIDTH_CM - sum(min_widths)
    total_content = sum(max_visual_lens)
    widths = [min_widths[i] + ((max_visual_lens[i] / total_content) * remaining
                               if total_content else 0) for i in range(n_cols)]
    scale = TOTAL_WIDTH_CM / sum(widths)
    return [w * scale for w in widths]


def set_run_font(run, font, size=None, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size:
        run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_font(cell, font, size=None, bold=False, color=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, font, size, bold, color)


def parse_md_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            rows.append([strip_md_marks(c.strip()) for c in line.split('|')[1:-1]])
    return [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]


MD_INLINE_RE = re.compile(r'\*\*(.+?)\*\*|\[口令\]')
MD_HR_RE = re.compile(r'^[-*_]{3,}$')
MD_SEP_CELL_RE = re.compile(r'^[-:\s]*$')


def strip_md_marks(text):
    """去掉 Markdown 内联标记，只留纯文本（用于标题/表格单元格）。"""
    return re.sub(r'\*\*(.+?)\*\*', r'\1', text).replace('**', '')


def _md_row_cells(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def _is_md_sep_row(line):
    line = line.strip()
    if not line.startswith('|'):
        return False
    cells = _md_row_cells(line)
    return bool(cells) and all(MD_SEP_CELL_RE.match(c) for c in cells)


def clean_md_lines(raw_lines):
    """信息框内容清洗：丢弃水平线与 Markdown 表格的表头/分隔行，数据行转“键：值”。"""
    lines = [l.strip() for l in raw_lines]
    out, i = [], 0
    while i < len(lines):
        s = lines[i]
        if not s or MD_HR_RE.match(s):
            i += 1
            continue
        if s.startswith('|'):
            if _is_md_sep_row(s):
                i += 1
                continue
            if i + 1 < len(lines) and _is_md_sep_row(lines[i + 1]):
                i += 2  # 表头行 + 分隔行一起丢弃
                continue
            cells = [strip_md_marks(c) for c in _md_row_cells(s) if c]
            if cells:
                out.append('：'.join(cells))
            i += 1
            continue
        out.append(strip_md_marks(s))
        i += 1
    return out


def add_paragraph_with_commands(doc, text, font):
    """正文段落渲染：**加粗** 转 run 加粗、[口令] 加粗，不残留 Markdown 标记。"""
    p = doc.add_paragraph()
    pos = 0
    for m in MD_INLINE_RE.finditer(text):
        plain = text[pos:m.start()].replace('**', '')
        if plain:
            set_run_font(p.add_run(plain), font)
        set_run_font(p.add_run(m.group(1) or '[口令]'), font, bold=True)
        pos = m.end()
    tail = text[pos:].replace('**', '')
    if tail or not p.runs:
        set_run_font(p.add_run(tail), font)
    return p


def add_numbered_list(doc, items, font):
    """手动编号列表：避免 python-docx 'List Number' 样式在多个列表间续排。"""
    for idx, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        set_run_font(p.add_run(f"{idx}. {text}"), font)


def small_gap(doc):
    """表后间隔段：小字号零间距，避免把下一段推成孤页。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(" ")
    r.font.size = Pt(4)
    return p


# === 解析（每套独立状态，异常向上抛） ===

def find_subtechs(project_dir):
    subtechs = []
    for script_path in sorted(glob.glob(os.path.join(project_dir, "*_试讲稿_v1.0.md"))):
        name_no_ext = os.path.basename(script_path).replace('_试讲稿_v1.0.md', '')
        parts = name_no_ext.split('_', 3)
        if len(parts) < 4:
            continue
        proj_num, proj_name, subtech_id, subtech_name = parts
        subtechs.append({
            'id': subtech_id, 'name': subtech_name,
            'script_path': script_path,
            'design_path': os.path.join(
                project_dir, f"{proj_num}_{proj_name}_{subtech_id}_{subtech_name}_教学设计_v1.0.md"),
            'selfcheck_path': os.path.join(
                project_dir, f"{proj_num}_{proj_name}_{subtech_id}_{subtech_name}_自检表_v1.0.md"),
        })
    return subtechs


def parse_teaching_design(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    result = {'objectives': '', 'key_points': '', 'difficult_points': '', 'venue_equipment': ''}

    in_design_table = False
    for line in lines:
        stripped = line.strip()
        if '教学设计说明' in stripped and stripped.startswith('#'):
            in_design_table = True
            continue
        if in_design_table:
            if stripped.startswith('#') and not stripped.startswith('|'):
                in_design_table = False
                continue
            if stripped.startswith('|'):
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if len(cells) >= 2:
                    if '教学重点' in cells[0]:
                        result['key_points'] = cells[1]
                    elif '教学难点' in cells[0]:
                        result['difficult_points'] = cells[1]

    in_obj, obj_lines = False, []
    for line in lines:
        stripped = line.strip()
        if '教学目标' in stripped and stripped.startswith('#'):
            in_obj = True
            continue
        if in_obj:
            if stripped.startswith('#') and not stripped.startswith('###') and '教学目标' not in stripped:
                break
            if stripped and not stripped.startswith('###'):
                obj_lines.append(stripped)
    result['objectives'] = '\n'.join(clean_md_lines(obj_lines))

    in_venue, venue_lines = False, []
    for line in lines:
        stripped = line.strip()
        if '场地器材' in stripped and stripped.startswith('#'):
            in_venue = True
            continue
        if in_venue:
            if stripped.startswith('#'):
                break
            if stripped:
                venue_lines.append(stripped)
    result['venue_equipment'] = '\n'.join(clean_md_lines(venue_lines))
    return result


def parse_trial_script(filepath):
    """返回 (title, body_elements, score_anchors)。评分锚点仅收集，不进正文。"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    body_elements, score_anchors, title = [], [], ''
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if (stripped.startswith('项目:') or stripped.startswith('项目：')
                or stripped.startswith('> 项目:') or stripped.startswith('> 项目：')
                or stripped.startswith('> 约束') or stripped.startswith('> 生成时间')):
            i += 1; continue
        if stripped.startswith('> 评分锚点'):
            anchor = stripped.replace('> 评分锚点：', '').replace('> 评分锚点:', '').strip()
            score_anchors.append(anchor)
            i += 1; continue
        if stripped in ('---', '***'):
            i += 1; continue
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped[2:].strip()
            i += 1; continue
        if stripped.startswith('#### '):
            body_elements.append(('heading4', stripped[5:].strip())); i += 1; continue
        if stripped.startswith('### '):
            body_elements.append(('heading3', stripped[4:].strip())); i += 1; continue
        if stripped.startswith('## '):
            body_elements.append(('heading2', stripped[3:].strip())); i += 1; continue
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            rows = parse_md_table(table_lines)
            if rows:
                body_elements.append(('table', rows))
            continue
        if not stripped:
            i += 1; continue
        body_elements.append(('paragraph', stripped))
        i += 1
    return title, body_elements, score_anchors


def parse_selfcheck_anchors(filepath):
    """新体系产物：锚点在自检表 R1 行（旧体系试讲稿内联锚点则不触发此路径）。"""
    anchors = []
    for line in open(filepath, 'r', encoding='utf-8'):
        s = line.strip()
        if s.startswith('|'):
            cells = [c.strip() for c in s.split('|')[1:-1]]
            if cells and va.ID_CELL_RE.match(cells[0]):
                desc = cells[1] if len(cells) > 1 else cells[0]
                verdict = cells[2] if len(cells) > 2 else ''
                anchors.append(f"[{cells[0].split()[0]}] {desc}（{verdict}）" if verdict
                               else f"[{cells[0].split()[0]}] {desc}")
    return anchors


# === 文档构建 ===

def setup_page(section):
    section.page_width = Cm(PAGE_W_CM)     # 显式 A4，不依赖默认 Letter
    section.page_height = Cm(PAGE_H_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)


def setup_styles(doc, font):
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(4)
    for level, size, color in [(1, Pt(18), RGBColor(0x1F, 0x49, 0x7D)),
                               (2, Pt(14), RGBColor(0x2E, 0x74, 0xB5)),
                               (3, Pt(12), RGBColor(0x4F, 0x81, 0xBD)),
                               (4, Pt(11), RGBColor(0x66, 0x66, 0x66))]:
        h = doc.styles[f'Heading {level}']
        h.font.name = font
        h.font.size = size
        h.font.bold = True
        h.font.color.rgb = color
        h.element.rPr.rFonts.set(qn('w:eastAsia'), font)
        h.paragraph_format.space_before = Pt(10 if level <= 2 else 8)
        h.paragraph_format.space_after = Pt(4 if level <= 2 else 2)
        h.paragraph_format.keep_with_next = True   # 标题不与后文分页


def add_front_heading(doc, text, font):
    p = doc.add_paragraph()
    set_run_font(p.add_run(text), font, size=Pt(18), bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_front_subheading(doc, text, font):
    p = doc.add_paragraph()
    set_run_font(p.add_run(text), font, size=Pt(14), bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_cover_page(doc, font, stats):
    for _ in range(6):
        doc.add_paragraph()
    for text, size, color, bold in [
        ('人教版体育试讲备考讲义', Pt(28), RGBColor(0x1F, 0x49, 0x7D), True),
        ('基于人教版《体育与健康》教师用书', Pt(14), RGBColor(0x66, 0x66, 0x66), False),
        (f"v2.0 · {datetime.date.today().strftime('%Y年%m月')}", Pt(12), RGBColor(0x99, 0x99, 0x99), False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), font, size=size, bold=bold, color=color)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"覆盖 {stats['n_projects']} 个运动项目 · {stats['n_subtechs']} 个子技术"),
                 font, size=Pt(12), color=RGBColor(0x99, 0x99, 0x99))
    doc.add_page_break()


def add_usage_instructions(doc, font):
    add_front_heading(doc, '使用说明', font)
    add_front_subheading(doc, '讲义结构', font)
    doc.add_paragraph('每个子技术包含三部分内容，按以下顺序阅读：')
    add_numbered_list(doc, [
        '设计要点卡片（蓝色信息框）-- 教学目标、教学重难点、场地器材，帮助快速建立教学认知框架。',
        '试讲稿正文 -- 可直接照念的完整试讲内容，包含简案大纲、四环节全文、[口令]标记。',
        '评分检查清单 -- 该子技术评分锚点汇总，读完试讲稿后逐条自查。',
    ], font)
    add_front_subheading(doc, '时间与字数（默认画像，诊断值）', font)
    doc.add_paragraph(
        '默认画像：试讲 10 分钟，环节 1+2+6+1，字数 1800-2200 字。以上均为有来源的默认诊断值，'
        '不是全国硬规则；实际以各地 EXAM_PROFILE（考试类型、备课/试讲时长、学段、教材版本）为准。')
    add_front_subheading(doc, '备考建议', font)
    add_numbered_list(doc, [
        '先读设计要点卡片，理解"教什么、为什么这么教"。',
        '通读试讲稿正文，感受试讲节奏和内容密度。',
        '脱稿模拟试讲，用手机录音计时，对照画像时长控制。',
        '对照评分检查清单逐条核对；安全与知识错误为一票否决项，优先排查。',
    ], font)
    doc.add_page_break()


def add_scoring_criteria(doc, font):
    """R1 判定制（替代旧版六维度 5 分制自报评分）。"""
    add_front_heading(doc, '通用评分口径（R1 判定制）', font)
    doc.add_paragraph(
        '本讲义不提供"给自己打分"的分值表。评分采用判定制：每个锚点只有'
        '"通过 / 不通过 / 待演练"三态；安全（VETO-SAFETY）与知识错误（VETO-KNOW）一票否决；'
        '标注 [LIVE] 的项目必须真人演练或录像后才能判"通过"，文本自报视为无效。')
    table = doc.add_table(rows=10, cols=3, style='Table Grid')
    table.autofit = False
    set_table_fixed_layout(table)
    mark_header_row(table)
    for j, h in enumerate(['维度', '代表锚点', '判定方式']):
        cell = table.cell(0, j)
        cell.text = h
        shade_cell(cell, HEADER_ROW_BG)
        set_cell_font(cell, font, bold=True)
    data = [
        ['教材与技术准确', 'A1 技术要领与教材一致', '对照教材页码逐条核对，冲突即 VETO-KNOW'],
        ['教学实施', 'B1 环节结构 / B3 分层递进', '结构完整且时长求和等于画像试讲时长'],
        ['语言教态', 'C1 口令 / C2 讲解', '[LIVE] 项须演练后判定'],
        ['板书', 'D1 板书与讲解同步', '备课提纲与板书设计一致性检查'],
        ['过程评价', 'E1 学练赛评链条', '文本可判，证据须给行号'],
        ['考场节奏', 'F1 字数 / F2 时长求和', '诊断值 ±20% / 硬门 ±0.5 分钟'],
        ['安全组织', 'VETO-SAFETY', '高危项目必须有保护帮助措施，一票否决'],
        ['迁移性', 'G1 可迁移到同类技术', '换项复述测试'],
        ['综合', 'A3a/B2a/C1a/C2a [LIVE]', '待演练，禁止文本自报通过'],
    ]
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.cell(i + 1, j)
            cell.text = val
            set_cell_font(cell, font)
    set_col_widths(table, [3.2, 5.4, 7.4])
    compact_table(table)
    small_gap(doc)
    doc.add_page_break()


def add_toc(doc, font, toc_rows, pages=None):
    """静态目录表。pages=None 时页码列填占位（Pass1），否则填实测页码（Pass2）。"""
    add_front_heading(doc, '目录', font)
    table = doc.add_table(rows=len(toc_rows) + 1, cols=3, style='Table Grid')
    table.autofit = False
    set_table_fixed_layout(table)
    mark_header_row(table)
    for j, h in enumerate(['项目', '子技术数', '页码']):
        cell = table.cell(0, j)
        cell.text = h
        shade_cell(cell, HEADER_ROW_BG)
        set_cell_font(cell, font, bold=True)
    for i, (proj_label, n_sub) in enumerate(toc_rows):
        page = (pages or {}).get(proj_label, '—')
        for j, val in enumerate([proj_label, str(n_sub), str(page)]):
            cell = table.cell(i + 1, j)
            cell.text = val
            set_cell_font(cell, font)
    set_col_widths(table, [9.0, 3.5, 3.5])
    compact_table(table)
    doc.add_page_break()


def add_info_box(doc, design_data, font):
    rows_data = []
    for key, label in [('objectives', '教学目标'), ('key_points', '教学重点'),
                       ('difficult_points', '教学难点'), ('venue_equipment', '场地器材')]:
        if design_data.get(key):
            rows_data.append((label, design_data[key]))
    if not rows_data:
        return
    table = doc.add_table(rows=len(rows_data), cols=2, style='Table Grid')
    table.autofit = False
    set_table_fixed_layout(table)
    for i, (label, content) in enumerate(rows_data):
        cell = table.cell(i, 0)
        cell.text = label
        shade_cell(cell, INFO_BOX_BG)
        set_cell_font(cell, font, bold=True)
        cell = table.cell(i, 1)
        for j, line in enumerate(content.split('\n')):
            if j == 0:
                cell.text = line
            else:
                cell.add_paragraph(line)
        set_cell_font(cell, font)
    set_col_widths(table, [2.5, 13.5])
    compact_table(table)
    small_gap(doc)


def add_table_to_doc(doc, rows, font):
    """列数 > MAX_TABLE_COLS 时按列分片，内容零丢失，每片 ≤ 6 数据列。"""
    for chunk in split_wide_table(rows):
        n_cols = max(len(r) for r in chunk)
        table = doc.add_table(rows=len(chunk), cols=n_cols, style='Table Grid')
        table.autofit = False
        set_table_fixed_layout(table)
        mark_header_row(table)
        for i, row in enumerate(chunk):
            for j in range(n_cols):
                cell = table.cell(i, j)
                if j < len(row):
                    cell.text = row[j]
                set_cell_font(cell, font, bold=(i == 0))
                if i == 0:
                    shade_cell(cell, HEADER_ROW_BG)
        set_col_widths(table, compute_col_widths(chunk, n_cols))
        compact_table(table)
        small_gap(doc)


def split_wide_table(rows):
    n_cols = max(len(r) for r in rows)
    if n_cols <= MAX_TABLE_COLS:
        return [rows]
    header, body = rows[0], rows[1:]
    chunks, data_cols = [], list(range(1, n_cols))
    for k in range(0, len(data_cols), 5):
        idx = [0] + data_cols[k:k + 5]
        chunk = [[(r[j] if j < len(r) else '') for j in idx] for r in [header] + body]
        if k:
            chunk[0] = [f"{chunk[0][j]}（续{k // 5 + 1}）" if j else chunk[0][j] for j in range(len(idx))]
        chunks.append(chunk)
    return chunks


def add_score_checklist(doc, anchors, font):
    if not anchors:
        return
    doc.add_heading('评分检查清单', level=3)
    items = []
    for anchor in anchors:
        parts = re.split(r'\s*[--]\s*', anchor, 1)
        items.append((parts[0].strip(), parts[1].strip()) if len(parts) == 2 else ('', anchor))
    table = doc.add_table(rows=len(items) + 1, cols=2, style='Table Grid')
    table.autofit = False
    set_table_fixed_layout(table)
    mark_header_row(table)
    for j, h in enumerate(['评分项', '得分动作']):
        cell = table.cell(0, j)
        cell.text = h
        shade_cell(cell, HEADER_ROW_BG)
        set_cell_font(cell, font, bold=True)
    for i, (item, action) in enumerate(items):
        cell = table.cell(i + 1, 0)
        cell.text = item
        set_cell_font(cell, font)
        cell = table.cell(i + 1, 1)
        cell.text = action
        set_cell_font(cell, font)
    set_col_widths(table, [3.5, 12.5])
    compact_table(table)


def add_trial_script_body(doc, body_elements, font):
    for elem in body_elements:
        if elem[0] == 'heading2':
            doc.add_heading(strip_md_marks(elem[1]), level=3)
        elif elem[0] == 'heading3':
            doc.add_heading(strip_md_marks(elem[1]), level=4)
        elif elem[0] == 'heading4':
            p = doc.add_paragraph()
            set_run_font(p.add_run(strip_md_marks(elem[1])), font, bold=True)
        elif elem[0] == 'table':
            add_table_to_doc(doc, elem[1], font)
        elif elem[0] == 'paragraph':
            add_paragraph_with_commands(doc, elem[1], font)


def set_section_header(section, project_name, font):
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    p.text = ''
    set_run_font(p.add_run(project_name), font, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run()
    r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = fp.add_run()
    r2._r.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r3 = fp.add_run()
    r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def build_document(selected, font, stats, pages=None):
    """selected: [(proj_num, proj_name, dir_name, [subtech, ...]), ...]"""
    doc = Document()
    setup_styles(doc, font)
    for section in doc.sections:
        setup_page(section)
    add_cover_page(doc, font, stats)
    add_usage_instructions(doc, font)
    add_scoring_criteria(doc, font)
    add_toc(doc, font, stats['toc_rows'], pages)
    for proj_num, proj_name, _, subtechs in selected:
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        setup_page(new_section)                       # 新 section 同样显式 A4
        set_section_header(new_section, f"{proj_num} {proj_name}", font)
        doc.add_heading(f"{proj_num} {proj_name}", level=1)
        for k, sub in enumerate(subtechs):
            if k:                                     # 首个子技术紧跟项目标题，消除近空白页
                doc.add_page_break()
            doc.add_heading(f"{sub['id']} {sub['name']}", level=2)
            if os.path.exists(sub['design_path']):
                add_info_box(doc, parse_teaching_design(sub['design_path']), font)
            _, body_elements, anchors = parse_trial_script(sub['script_path'])
            add_trial_script_body(doc, body_elements, font)
            if not anchors and os.path.exists(sub['selfcheck_path']):   # 新体系锚点回退
                anchors = parse_selfcheck_anchors(sub['selfcheck_path'])
            add_score_checklist(doc, anchors, font)
    return doc


# === 渲染与页码验证 ===

def render_pdf(docx_path, out_dir):
    soffice = shutil.which("soffice")
    if not soffice:
        raise GenError(5, "未找到 soffice，无法渲染验证目录页码；请安装 LibreOffice 或将其加入 PATH")
    r = subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                        out_dir, docx_path], capture_output=True, text=True, timeout=300)
    pdf = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if r.returncode != 0 or not os.path.exists(pdf):
        raise GenError(5, f"soffice 渲染失败: {r.stderr.strip()[:300]}")
    return pdf


def extract_project_pages(pdf, proj_labels):
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        raise GenError(5, "未找到 pdftotext，无法验证目录页码")
    r = subprocess.run([pdftotext, "-layout", pdf, "-"], capture_output=True, text=True, timeout=120)
    if r.returncode != 0:
        raise GenError(5, f"pdftotext 失败: {r.stderr.strip()[:200]}")
    pages_text = r.stdout.split("\f")
    # 目录页指纹：单页同时出现 >=2 个项目标签 => 是目录页，不算章节起始页。
    # 正文各项目页眉含本标签（每 section 独立页眉），跳过目录页后首现页即章节起始页。
    toc_pages = [pno for pno, ptext in enumerate(pages_text, start=1)
                 if sum(1 for l in proj_labels if l in ptext) >= 2]
    last_toc = max(toc_pages) if toc_pages else 0
    if not toc_pages:
        raise GenError(5, "渲染 PDF 中未识别到目录页（无任何单页含>=2个项目标签）")
    found = {}
    for label in proj_labels:
        for pno, ptext in enumerate(pages_text, start=1):
            if pno <= last_toc:
                continue
            if label in ptext:
                found[label] = pno
                break
    missing = [l for l in proj_labels if l not in found]
    if missing:
        raise GenError(5, f"渲染 PDF 中找不到项目标题页: {missing}")
    return found, len([t for t in pages_text if t.strip()])


def selfcheck_docx(path, font, toc_pages_expected=None):
    """复用 validate_artifacts 的 DOCX 结构检查 + 目录页码数字检查。"""
    rep = va.Report(path)
    va.validate_docx(path, rep)
    if not rep.ok:
        detail = "; ".join(f"{e['code']} {e['msg']}" for e in rep.veto + rep.errors)
        raise GenError(5, f"DOCX 结构自检失败: {detail}")
    import zipfile
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    if f'w:eastAsia="{font}"' not in xml:
        raise GenError(5, f"document.xml 缺少 eastAsia 字体声明 {font}")
    if toc_pages_expected:
        for label, page in toc_pages_expected.items():
            if str(page) not in xml:
                raise GenError(5, f"目录页码 {label}->{page} 未写入文档")


# === 主流程 ===

def detect_font(explicit):
    if explicit:
        for name, path in FONT_CANDIDATES:
            if name == explicit and os.path.exists(path):
                return name, path
        raise GenError(2, f"--font {explicit!r} 不在可用中文字体候选内；可用: "
                          f"{[n for n, p in FONT_CANDIDATES if os.path.exists(p)] or '（未探测到系统字体文件）'}")
    for name, path in FONT_CANDIDATES:
        if os.path.exists(path):
            return name, path
    raise GenError(2, "未探测到可用中文字体（候选: " +
                       ", ".join(n for n, _ in FONT_CANDIDATES) + "）；请用 --font 指定")


def discover(base_dir, projects_arg, only_arg):
    wanted = None
    if projects_arg:
        wanted = [p.zfill(2) for p in re.split(r"[，,\s]+", projects_arg) if p]
        for w in wanted:
            if not re.fullmatch(r"\d{1,2}", w) or not (1 <= int(w) <= 10):
                raise GenError(2, f"--projects 含非法项目号: {w!r}（应为 01-10）")
    only = set()
    if only_arg:
        only = {s.strip() for s in re.split(r"[，,\s]+", only_arg) if s.strip()}
        for o in only:
            if not re.fullmatch(r"\d{2}-\d{2}", o):
                raise GenError(2, f"--only 含非法子技术号: {o!r}（应为如 07-10）")
    selected, total = [], 0
    for proj_num, proj_name, dir_name in PROJECTS:
        if wanted and proj_num not in wanted:
            continue
        project_dir = os.path.join(base_dir, dir_name)
        if not os.path.isdir(project_dir):
            raise GenError(2, f"选中的项目目录不存在: {project_dir}")
        subtechs = find_subtechs(project_dir)
        if only:
            subtechs = [s for s in subtechs if s['id'] in only]
        if subtechs:  # 空项目不进入讲义（避免目录与正文出现空章节）
            selected.append((proj_num, proj_name, dir_name, subtechs))
            total += len(subtechs)
    if total == 0:
        raise GenError(3, "零项目：选中范围内未找到任何 *_试讲稿_v1.0.md（拒绝输出空讲义）")
    return selected, total


def parse_all(selected):
    """解析全部源文件；任何异常 -> exit 4（不吞、不留 final）。"""
    parsed = []
    for proj_num, proj_name, _, subtechs in selected:
        items = []
        for sub in subtechs:
            try:
                design = (parse_teaching_design(sub['design_path'])
                          if os.path.exists(sub['design_path']) else None)
                title, body, anchors = parse_trial_script(sub['script_path'])
                if not anchors and os.path.exists(sub['selfcheck_path']):
                    anchors = parse_selfcheck_anchors(sub['selfcheck_path'])
                items.append(dict(sub, design=design, title=title,
                                  body=body, anchors=anchors))
            except Exception as e:
                raise GenError(4, f"解析异常 {sub['script_path']}: {type(e).__name__}: {e}")
        parsed.append((proj_num, proj_name, _, items))
    return parsed


def main(argv=None):
    ap = argparse.ArgumentParser(description="人教版体育试讲备考讲义生成器（生产版）")
    ap.add_argument("--base-dir", required=True, help="产物库根目录（含 01篮球..10武术）")
    ap.add_argument("--out", required=True, help="输出 DOCX 路径")
    ap.add_argument("--projects", default=None, help="项目号过滤，如 06,07,09（默认全部）")
    ap.add_argument("--only", default=None, help="子技术号过滤，如 07-10,06-10,09-04")
    ap.add_argument("--font", default=None, help="覆盖中文字体探测，如 'Songti SC'")
    ap.add_argument("--check", action="store_true",
                    help="完整构建+渲染+自检到临时文件，不写 final，输出报告")
    ap.add_argument("--dry-run", action="store_true",
                    help="只做发现与解析统计，不构建不写盘")
    args = ap.parse_args(argv)

    tmp_files = []
    tmp_dirs = []
    try:
        if PLACEHOLDER_RE.search(args.base_dir) or PLACEHOLDER_RE.search(args.out):
            raise GenError(2, "路径仍为占位符（含 <> / 项目根 / %X），请替换为本机实际路径")
        base_dir = os.path.abspath(os.path.expanduser(args.base_dir))
        out_path = os.path.abspath(os.path.expanduser(args.out))
        if not os.path.isdir(base_dir):
            raise GenError(2, f"--base-dir 不存在或不是目录: {base_dir}")
        out_parent = os.path.dirname(out_path)
        if not os.path.isdir(out_parent):
            raise GenError(2, f"--out 父目录不存在: {out_parent}")
        probe = os.path.join(out_parent, ".pe_trial_write_probe")
        try:
            with open(probe, "w") as f:
                f.write("x")
        except OSError as e:
            raise GenError(2, f"--out 目录不可写: {out_parent}: {e}")
        finally:
            if os.path.exists(probe):
                os.remove(probe)

        font_name, font_path = detect_font(args.font)
        selected, total = discover(base_dir, args.projects, args.only)
        parsed = parse_all(selected)

        n_projects = sum(1 for _, _, _, st in selected if st)
        stats = {
            'n_projects': n_projects,
            'n_subtechs': total,
            'toc_rows': [(f"{p} {n}", len(st)) for p, n, _, st in selected if st],
        }
        print(f"[plan] 项目 {n_projects} 个 / 子技术 {total} 个 / 字体 {font_name} ({font_path})")

        if args.dry_run:
            for p, n, _, items in parsed:
                print(f"  {p} {n}: {len(items)} 个子技术 -> "
                      + ", ".join(f"{s['id']} {s['name']}（锚点 {len(s['anchors'])}）" for s in items))
            print("[dry-run] 校验通过，未写任何文件")
            return 0

        tmp_dir = tempfile.mkdtemp(prefix="pe-lecture-")
        tmp_dirs.append(tmp_dir)
        tmp1 = os.path.join(tmp_dir, "pass1.docx")
        tmp_files.append(tmp1)

        # Pass 1: 构建含占位目录的完整文档 -> 渲染 -> 提取项目起始页
        build_document(parsed, font_name, stats).save(tmp1)
        pdf1 = render_pdf(tmp1, tmp_dir)
        tmp_files.append(pdf1)
        labels = [l for l, _ in stats['toc_rows']]
        pages, n_pages1 = extract_project_pages(pdf1, labels)
        print(f"[pass1] 渲染 {n_pages1} 页；项目起始页: "
              + ", ".join(f"{l}->p{pages[l]}" for l in labels))

        # Pass 2: 回填目录页码重建 -> 渲染 -> 验证页码一致 + 结构自检
        tmp2 = os.path.join(tmp_dir, "pass2.docx")
        tmp_files.append(tmp2)
        build_document(parsed, font_name, stats, pages=pages).save(tmp2)
        pdf2 = render_pdf(tmp2, tmp_dir)
        tmp_files.append(pdf2)
        pages2, n_pages2 = extract_project_pages(pdf2, labels)
        drift = {l: (pages[l], pages2[l]) for l in labels if pages[l] != pages2[l]}
        if drift:
            raise GenError(5, f"目录页码回填后发生漂移（Pass1 vs Pass2 不一致）: {drift}")
        selfcheck_docx(tmp2, font_name, toc_pages_expected=pages)
        print(f"[pass2] 页码零漂移；结构自检通过（A4/{MARGIN_CM}cm 边距/表宽≤{TOTAL_WIDTH_CM}cm/"
              f"列数≤{MAX_TABLE_COLS}/eastAsia={font_name}/目录页码可验证）")

        if args.check:
            print(f"[check] 全绿：{n_pages2} 页，项目 {n_projects}，子技术 {total}；未写 final")
            return 0

        # 正常模式：验证过的临时文件原子替换 final
        os.replace(tmp2, out_path)
        tmp_files.remove(tmp2)
        print(f"[done] 原子替换 -> {out_path}")
        return 0
    except GenError as e:
        print(f"ERROR[{e.code}]: {e}", file=sys.stderr)
        return e.code
    finally:
        for f in tmp_files:
            try:
                os.remove(f)
            except OSError:
                pass
        for d in tmp_dirs:
            shutil.rmtree(d, ignore_errors=True)


if __name__ == '__main__':
    sys.exit(main())
